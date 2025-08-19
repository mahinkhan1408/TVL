# letterhead_bid_module.py
import tkinter as tk
from tkinter import messagebox, filedialog
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
from PIL import Image

class LetterheadBidModule:
    def __init__(self, root):
        self.root = root
        self.root.title("Letterhead Bid Generator")
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.root.geometry(f"{int(screen_width * 0.7)}x{int(screen_height * 0.7)}")

        self.colors = {
            'primary_blue': '#1e3a5f',
            'background': '#f8f9fa',
            'white': '#ffffff',
            'gray_light': '#e9ecef',
            'gray_medium': '#6c757d',
            'gray_dark': '#495057',
            'green': '#28a745',
            'table_header': '#B6D7A8',
            'table_cell': '#D9D9D9',
        }
        self.root.configure(bg=self.colors['background'])

        self.setup_ui()

    def setup_ui(self):
        header_frame = tk.Frame(self.root, bg=self.colors['primary_blue'], height=60)
        header_frame.pack(fill='x', pady=(0, 10))
        header_frame.pack_propagate(False)
        tk.Label(header_frame, text="Generate Letterhead Bid",
                 font=("Arial", 18, "bold"), fg='white',
                 bg=self.colors['primary_blue']).pack(side="left", expand=True, padx=20)

        main_frame = tk.Frame(self.root, bg=self.colors['white'], padx=20, pady=20)
        main_frame.pack(fill='both', expand=True, padx=20, pady=10)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_columnconfigure(1, weight=2)

        row = 0
        tk.Label(main_frame, text="Select Template:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.template_options = ["ESUS Property Management LLC", "RAPID CARE Field Services"]
        self.selected_template = tk.StringVar(self.root)
        self.selected_template.set(self.template_options[0])
        template_menu = tk.OptionMenu(main_frame, self.selected_template, *self.template_options)
        template_menu.config(font=("Arial", 11), bg=self.colors['white'], relief="solid", bd=1, width=25)
        template_menu.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="WO:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.wo_entry = tk.Entry(main_frame, font=("Arial", 11), relief="solid", bd=1)
        self.wo_entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="Property Address:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.address_text = tk.Text(main_frame, font=("Arial", 11), relief="solid", bd=1, height=3)
        self.address_text.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="Date:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.date_entry = tk.Entry(main_frame, font=("Arial", 11), relief="solid", bd=1)
        self.date_entry.insert(0, datetime.now().strftime("%B %d, %Y"))
        self.date_entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="Description:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.desc_text = tk.Text(main_frame, font=("Arial", 11), relief="solid", bd=1, height=5)
        self.desc_text.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="Quantity:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.qty_entry = tk.Entry(main_frame, font=("Arial", 11), relief="solid", bd=1)
        self.qty_entry.insert(0, "1")
        self.qty_entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        tk.Label(main_frame, text="Price:", font=("Arial", 11, "bold"), bg=self.colors['white']).grid(row=row, column=0, sticky="w", pady=5, padx=5)
        self.price_entry = tk.Entry(main_frame, font=("Arial", 11), relief="solid", bd=1)
        self.price_entry.insert(0, "0.00")
        self.price_entry.grid(row=row, column=1, sticky="ew", pady=5, padx=5)

        row += 1
        generate_button = tk.Button(main_frame, text="Generate Document", command=self.generate_document,
                                    font=("Arial", 12, "bold"), bg=self.colors['green'], fg='white',
                                    relief="solid", bd=1, cursor="hand2")
        generate_button.grid(row=row, column=0, columnspan=2, pady=20)

    def set_cell_background(self, cell, color):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def generate_document(self):
        try:
            template_name = self.selected_template.get()
            wo = self.wo_entry.get()
            address = self.address_text.get("1.0", tk.END).strip()
            date = self.date_entry.get()
            description = self.desc_text.get("1.0", tk.END).strip()
            qty_str = self.qty_entry.get().strip()
            price_str = self.price_entry.get().strip().replace('$', '').replace(',', '')
            
            qty = float(qty_str) if qty_str else 0.0
            price = float(price_str) if price_str else 0.0
            total_cost = qty * price

            if not template_name:
                messagebox.showwarning("Warning", "Please select a template.")
                return

            template_path = ""
            if template_name == "ESUS Property Management LLC":
                template_path = "templates/esus_template.docx"
            elif template_name == "RAPID CARE Field Services":
                template_path = "templates/rapid_care_template.docx"

            if not os.path.exists(template_path):
                messagebox.showerror("File Error", f"Template file not found at: {template_path}")
                return

            doc = Document(template_path)
            
            # Find and replace placeholders
            for para in doc.paragraphs:
                if 'WO:' in para.text:
                    para.text = ''
                    para.add_run(f'WO: {wo}').bold = True
                if 'PROPERTY ADDRESS:' in para.text:
                    para.text = ''
                    para.add_run(f'PROPERTY ADDRESS: {address}').bold = True
                if 'DATE:' in para.text:
                    para.text = ''
                    para.add_run(f'DATE: {date}').bold = True

            # Find the table and populate it
            table_found = False
            for table in doc.tables:
                # Assuming the first table with a header "SL" is the correct one
                if table.rows[0].cells[0].text.strip() == 'SL':
                    # Populate the first data row
                    data_row_cells = table.rows[1].cells
                    data_row_cells[0].text = '1'
                    data_row_cells[1].text = description
                    data_row_cells[2].text = str(qty_str)
                    data_row_cells[3].text = f"${price:.2f}"
                    data_row_cells[4].text = f"${total_cost:.2f}"
                    
                    # Populate the Total row
                    total_row_cells = table.rows[2].cells
                    total_row_cells[3].text = 'Total'
                    total_row_cells[3].paragraphs[0].runs[0].font.bold = True
                    total_row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    total_row_cells[4].text = f"${total_cost:.2f}"
                    total_row_cells[4].paragraphs[0].runs[0].font.bold = True
                    
                    table_found = True
                    break

            if not table_found:
                messagebox.showwarning("Warning", "No bid table found in the template.")
                return

            # Save the document
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=f"Bid_{wo or 'Estimate'}_{datetime.now().strftime('%Y%m%d')}.docx"
            )
            
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Success", f"Document generated and saved to:\n{file_path}")
        except ValueError:
            messagebox.showerror("Input Error", "Please enter valid numerical values for QTY and Price.")
        except FileNotFoundError as e:
            messagebox.showerror("File Error", f"Template or logo file not found. Make sure 'templates' folder exists and contains the correct files.")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {e}")