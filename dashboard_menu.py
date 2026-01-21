# dashboard_menu.py
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from bid_writer_module import BidWriterApp
from notice_board_module import NoticeBoardModule
from vendor_price_module import VendorPriceModule
from gc_roof_ce_module import GCRoofCEModule
from letterhead_bid_module import LetterheadBidModule
from approval_module import ApprovalModule
from wo_inspection_module import WOInspectionModule
from theme_manager import theme_manager
import sys
import os
import importlib.util
import webbrowser

# Try to import PIL for image handling
try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    Image = None
    ImageTk = None
# Import Photo Viewer - handle the space in filename
photo_viewer_path = os.path.join(os.path.dirname(__file__), "Photo Viewer.py")
if os.path.exists(photo_viewer_path):
    spec = importlib.util.spec_from_file_location("photo_viewer", photo_viewer_path)
    photo_viewer_module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(photo_viewer_module)
    FastImageViewer = photo_viewer_module.FastImageViewer
else:
    FastImageViewer = None
import time
import re
import threading
from datetime import datetime
import json
import shutil
import tempfile

class DashboardMenu:
    def __init__(self, root, username):
        self.root = root
        self.root.title("Techvengers Universal App - Dashboard")
        
        self.username = username
        self.current_module_instance = None
        self.cached_modules = {}  # Cache module instances for faster switching
        
        # Use theme manager for colors (must be before database init)
        self.colors = theme_manager.get_current_colors()
        
        # Initialize Supabase database (optional - for data storage, not login)
        # Initialize asynchronously to prevent blocking UI
        self.db = None
        self.user_id = None
        self.is_admin_user = False
        
        # Register for theme updates
        theme_manager.register_theme_callback(self.on_theme_changed)
        
        self.root.update_idletasks()
        width = int(self.root.winfo_screenwidth() * 0.9)
        height = int(self.root.winfo_screenheight() * 0.9)
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
        self.root.configure(bg=self.colors['background'])
        
        # Initialize database after UI is set up and main loop is running
        self.root.after(100, lambda: self._init_database_async(username))
        # Full-page content (no left sidebar)
        self.main_content_frame = tk.Frame(self.root, bg=self.colors['background'])
        self.main_content_frame.pack(fill='both', expand=True)
        
        # Simple top navigation row
        top_nav = tk.Frame(self.main_content_frame, bg=self.colors['background'])
        top_nav.pack(fill='x', padx=16, pady=(12, 0))
        
        nav_buttons = [
            ("Dashboard", self.show_dashboard_content),
            ("Bid Writer", self.show_bid_writer_dashboard),
            ("Letterhead Bid", self.open_letterhead_bid),
            ("Notice Board", self.open_notice_board),
            ("Vendor Price", self.open_vendor_price),
            ("GC/Roof CE", self.open_gc_roof_ce),
            ("Settings", self.show_settings),
        ]
        self.nav_btn_refs = {}
        for text, cmd in nav_buttons:
            b = tk.Button(top_nav, text=text, font=("Segoe UI", 10), bg=self.colors['primary_blue'],
                          fg=self.colors['button_text'], relief='flat', cursor='hand2', padx=10, pady=4,
                          command=lambda c=cmd, t=text: self._nav_click_fullpage(c, t))
            b.pack(side='left', padx=(0, 6))
            self.nav_btn_refs[text] = b
        
        # Initial view
        self._nav_click_fullpage(self.show_dashboard_content, "Dashboard")
        
        self.app_data_dir = os.path.join(os.path.expanduser("~"), ".techvengers_bidwriter")
        os.makedirs(self.app_data_dir, exist_ok=True)
    
    def _init_database_async(self, username):
        """Initialize database asynchronously to prevent blocking UI."""
        def init_in_thread():
            try:
                from database_online import OnlineDatabaseManager
                db = OnlineDatabaseManager()
                # Get user_id from database if user exists
                user = db.get_user(username)
                user_id = user['id'] if user else None
                
                # Update in main thread - use after_idle to ensure main loop is running
                def update_db():
                    try:
                        self.db = db
                        self.user_id = user_id
                        # Check if user is admin
                        if db and user_id:
                            self.is_admin_user = db.is_admin(user_id)
                        else:
                            self.is_admin_user = False
                    except Exception as e:
                        print(f"Error updating database reference: {e}")
                
                # Use after_idle instead of after(0) to ensure main loop is ready
                try:
                    self.root.after_idle(update_db)
                except:
                    # Fallback: try after with a small delay
                    self.root.after(100, update_db)
            except Exception as e:
                print(f"Warning: Could not initialize database: {e}")
                def set_none():
                    try:
                        self.db = None
                        self.user_id = None
                    except Exception as e:
                        print(f"Error setting database to None: {e}")
                
                try:
                    self.root.after_idle(set_none)
                except:
                    self.root.after(100, set_none)
        
        # Start initialization in background thread
        thread = threading.Thread(target=init_in_thread, daemon=True)
        thread.start()

    def on_theme_changed(self, theme_name, colors):
        """Called when theme is changed globally."""
        self.colors = colors
        self.apply_theme_to_dashboard()

    def apply_theme_to_dashboard(self):
        """Apply the current theme to all dashboard elements."""
        # Update root and main frames
        self.root.configure(bg=self.colors['background'])
        self.main_content_frame.configure(bg=self.colors['background'])
        # Update top nav button colors
        for btn in getattr(self, 'nav_btn_refs', {}).values():
            btn.configure(bg=self.colors['primary_blue'], fg=self.colors['button_text'])
        
        # Update any other UI elements that might be visible
        self.refresh_current_view()

    def refresh_current_view(self):
        """Refresh the current view to apply theme changes."""
        # Find which navigation button is currently active and refresh that view
        for btn_text, btn in self.nav_btn_refs.items():
            if btn.cget('bg') == self.colors['nav_hover']:
                # Re-trigger the current view
                if btn_text == "Dashboard":
                    self.show_dashboard_content()
                elif btn_text == "Bid Writer":
                    self.show_bid_writer_dashboard()
                elif btn_text == "Settings":
                    self.show_settings()
                break

    def _nav_click_fullpage(self, command, label):
        # Visual feedback for active nav (outline on hover spec also handled per card)
        for name, btn in self.nav_btn_refs.items():
            btn.configure(relief='flat', highlightthickness=0)
        if label in self.nav_btn_refs:
            self.nav_btn_refs[label].configure(relief='groove', highlightthickness=1,
                                               highlightbackground=self.colors.get('nav_hover', self.colors['primary_blue']))
        command()

    def show_placeholder(self, title):
        self.clear_content_frame()
        container = tk.Frame(self.main_content_frame, bg=self.colors['background'])
        container.pack(fill='both', expand=True)
        self._add_back_bar(container, title)
        tk.Label(container, text=f"Content for {title} (Coming Soon)", 
                 font=("Arial", 18, "bold"), fg=self.colors['primary_blue'],
                 bg=self.colors['background']).pack(expand=True)

    def clear_content_frame(self):
        """Clear content frame efficiently."""
        # Use update_idletasks to batch widget operations
        for widget in self.main_content_frame.winfo_children():
            widget.destroy()
        # Force update to prevent visual lag
        self.root.update_idletasks()

    def show_dashboard_content(self):
        self.clear_content_frame()
        root_bg = self.colors['background']
        self.main_content_frame.config(bg=root_bg)

        # --- Hero header (two-tone) ---
        hero = tk.Frame(self.main_content_frame, bg=root_bg)
        hero.pack(fill='x', padx=20, pady=(20, 10))

        hero_wrap = tk.Frame(hero, bg=root_bg)
        hero_wrap.pack(fill='x')

        left_bg = self.colors.get('nav_hover', self.colors['primary_blue'])
        right_bg = self.colors['primary_blue']

        left = tk.Frame(hero_wrap, bg=left_bg, height=120)
        left.pack(side='left', fill='x', expand=True)
        left.pack_propagate(False)

        right = tk.Frame(hero_wrap, bg=right_bg, height=120)
        right.pack(side='left', fill='x', expand=True)
        right.pack_propagate(False)

        # Left: app title / logo
        tk.Label(left, text="📄  BidWrite Pro", font=("Segoe UI", 20, "bold"),
                 bg=left_bg, fg=self.colors['button_text']).pack(anchor='w', padx=20, pady=(25, 0))
        tk.Label(left, text="Craft winning bids with ease", font=("Segoe UI", 11),
                 bg=left_bg, fg=self.colors['button_text']).pack(anchor='w', padx=20, pady=(6, 0))

        # Right: welcome
        tk.Label(right, text=f"Welcome Back, {self.username}!", font=("Segoe UI", 18, "bold"),
                 bg=right_bg, fg=self.colors['button_text']).pack(anchor='w', padx=20, pady=(25, 0))
        tk.Label(right, text="Let's get started on your next winning bid.", font=("Segoe UI", 11),
                 bg=right_bg, fg=self.colors['button_text']).pack(anchor='w', padx=20, pady=(6, 0))

        # --- Cards grid ---
        cards_wrap = tk.Frame(self.main_content_frame, bg=root_bg)
        cards_wrap.pack(fill='both', expand=True, padx=20, pady=(5, 20))

        # Grid container
        grid = tk.Frame(cards_wrap, bg=root_bg)
        grid.pack(fill='both', expand=True)
        for i in range(4):
            grid.grid_columnconfigure(i, weight=1, uniform='col')

        # Card definitions in requested serial
        # New bid, Open project, GC/Roof CE, Price Sheet, Letterheads,
        # Notice Boards, Photo Viewer, Approval, WO Inspection, Settings
        cards = [
            ("New Bid", "Create a new bid", "🆕", self.create_new_bid),
            ("Open Project", "Continue your saved work", "📂", self.show_bid_writer_dashboard),
            ("GC/Roof CE", "GC/Roof change orders", "🏗️", self.open_gc_roof_ce),
            ("Price Sheet", "Client and Vendor", "💲", self.show_price_sheet_page),
            ("Letterheads", "Letterhead bids", "📝", self.open_letterhead_bid),
            ("Notice Boards", "Announcements", "📢", self.open_notice_board),
            ("Photo Viewer", "View and edit images", "🖼️", self.open_photo_viewer),
            ("Approval", "Approval workflow", "✔️", self.open_approval),
            ("WO Inspection", "Work order inspections", "🔍", self.open_wo_inspection),
            ("Settings", "Preferences", "⚙️", self.show_settings),
        ]
        
        # Add Admin card if user is admin (placed before Settings)
        if self.is_admin_user:
            cards.insert(-1, ("Admin", "User management & analytics", "👤", self.open_admin))

        # Create cards in a neat grid (up to 4 columns per row)
        row, col = 0, 0
        for title, subtitle, icon, cmd in cards:
            self._create_dashboard_card(grid, row, col, icon, title, subtitle, cmd)
            col += 1
            if col >= 4:
                col = 0
                row += 1

    def show_bid_writer_dashboard(self):
        self.clear_content_frame()

        bid_writer_dashboard = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        bid_writer_dashboard.pack(fill="both", expand=True)

        # Back to Dashboard bar
        self._add_back_bar(bid_writer_dashboard, "Bid Writer")

        header_frame = tk.Frame(bid_writer_dashboard, bg=self.colors['background'])
        header_frame.pack(fill='x', pady=(0, 20))

        tk.Label(header_frame, text="Bid Writer", font=("Arial", 24, "bold"), fg=self.colors['primary_blue'], bg=self.colors['background']).pack(side='left')

        # Create New Bid Section
        new_frame = tk.Frame(bid_writer_dashboard, bg=self.colors['gray_light'], relief="solid", bd=1, padx=15, pady=15)
        new_frame.pack(fill='x', pady=(0, 20))
        tk.Label(new_frame, text="Create New Bid", font=("Arial", 14, "bold"), bg=self.colors['gray_light'], fg=self.colors['primary_blue']).pack(anchor='w', pady=(0, 10))
        new_bid_button = tk.Button(new_frame, text="New Bid", font=("Arial", 12), bg=self.colors['light_blue'], fg="white", relief="solid", bd=1, cursor="hand2", command=self.create_new_bid)
        new_bid_button.pack(anchor='w', padx=5, pady=5)
        
        # Recent Bids Section
        recent_frame = tk.Frame(bid_writer_dashboard, bg=self.colors['gray_light'], relief="solid", bd=1, padx=15, pady=15)
        recent_frame.pack(fill='both', expand=True)
        recent_header_frame = tk.Frame(recent_frame, bg=self.colors['gray_light'])
        recent_header_frame.pack(fill='x')
        
        # Left side: Title and Search bar
        left_frame = tk.Frame(recent_header_frame, bg=self.colors['gray_light'])
        left_frame.pack(side='left', anchor='w', pady=(0, 10))
        
        tk.Label(left_frame, text="Recent Projects", font=("Arial", 14, "bold"), bg=self.colors['gray_light'], fg=self.colors['primary_blue'], anchor='w', justify='left').pack(side='left', anchor='w', padx=(0, 15))
        
        # Search bar with filter dropdown (on the left)
        search_frame = tk.Frame(left_frame, bg=self.colors['gray_light'])
        search_frame.pack(side='left', pady=(0, 10))
        
        search_label = tk.Label(search_frame, text="Search:", font=("Arial", 10), bg=self.colors['gray_light'])
        search_label.pack(side='left', padx=(0, 5))
        
        self.search_filter = tk.StringVar(value="all")
        filter_menu = tk.OptionMenu(search_frame, self.search_filter, "all", "work_order", "property_address")
        filter_menu.config(font=("Arial", 9), bg=self.colors['white'])
        filter_menu.pack(side='left', padx=(0, 5))
        
        self.search_entry = tk.Entry(search_frame, width=25, font=("Arial", 11), relief="solid", bd=1)
        self.search_entry.pack(side='left', padx=(0, 5))
        self.search_entry.bind("<Return>", lambda e: self.load_recent_bids())

        search_button = tk.Button(search_frame, text="Search", command=lambda: self.load_recent_bids(), font=("Arial", 10), bg=self.colors['light_blue'], fg="white", relief="flat", cursor="hand2")
        search_button.pack(side='left', padx=(0, 5))
        
        clear_search_button = tk.Button(search_frame, text="Clear", command=lambda: [self.search_entry.delete(0, tk.END), self.search_filter.set("all"), self.load_recent_bids()], font=("Arial", 10), bg=self.colors['gray_medium'], fg="white", relief="flat", cursor="hand2")
        clear_search_button.pack(side='left')
        
        # Dictionary to store selected checkboxes {wo_number: checkbox_var}
        self.selected_bids = {}
        
        # Right side: Action buttons (small buttons)
        buttons_frame = tk.Frame(recent_header_frame, bg=self.colors['gray_light'])
        buttons_frame.pack(side='right', anchor='e', pady=(0, 10))
        
        # Delete Selected button (small)
        self.bulk_delete_button = tk.Button(buttons_frame, text="Delete Selected", command=self.delete_selected_bids, font=("Arial", 9), bg='#dc3545', fg='white', relief='flat', cursor="hand2", state='disabled', padx=8, pady=4)
        self.bulk_delete_button.pack(side='right', padx=(5, 0))
        
        # Export Selected button (small)
        self.bulk_export_button = tk.Button(buttons_frame, text="Export Selected", command=self.export_selected_bids, font=("Arial", 9), bg=self.colors['primary_blue'], fg='white', relief='flat', cursor="hand2", state='disabled', padx=8, pady=4)
        self.bulk_export_button.pack(side='right', padx=(0, 0))


        # Scrollable list for recent bids
        recent_canvas = tk.Canvas(recent_frame, bg=self.colors['white'], highlightthickness=0)
        recent_scrollbar = tk.Scrollbar(recent_frame, orient="vertical", command=recent_canvas.yview)
        recent_scrollable_frame = tk.Frame(recent_canvas, bg=self.colors['white'])

        recent_scrollable_frame.bind(
            "<Configure>",
            lambda e: recent_canvas.configure(scrollregion=recent_canvas.bbox("all"))
        )

        def update_canvas_width(event):
            """Update canvas window width to match canvas width for full-width table"""
            canvas_width = event.width
            recent_canvas.itemconfig(recent_canvas.find_all()[0], width=canvas_width)
        
        recent_canvas.create_window((0, 0), window=recent_scrollable_frame, anchor="nw")
        recent_canvas.configure(yscrollcommand=recent_scrollbar.set)
        recent_canvas.bind('<Configure>', update_canvas_width)
        
        recent_scrollbar.pack(side="right", fill="y")
        recent_canvas.pack(side="left", fill="both", expand=True)
        
        self.recent_bids_list = recent_scrollable_frame
        
        # Initialize cache for projects
        self.app_data_dir = os.path.join(os.path.expanduser("~"), ".techvengers_bidwriter")
        os.makedirs(self.app_data_dir, exist_ok=True)
        self.projects_cache_file = os.path.join(self.app_data_dir, "projects_cache.json")
        
        # Initial load of bids (from cache first, then database)
        self.load_recent_bids_async()
        
        # Start frequent auto-refresh (every 30 seconds = 30000 milliseconds)
        self.projects_refresh_interval = 30 * 1000  # 30 seconds
        self.start_projects_auto_refresh()
    
    def load_recent_bids_async(self, search_term=None):
        """Load bids asynchronously - show cache first, then update from database"""
        # Load from cache first for instant display
        cache_bids = self.load_projects_from_cache(search_term)
        if cache_bids:
            self.display_bids(cache_bids, search_term)
            print(f"Loaded {len(cache_bids)} projects from cache (instant display)")
        
        # Then load from database in background and update
        def load_in_thread():
            bids = []
            
            # Try to load from database
            if self.db and self.user_id:
                try:
                    print("Loading projects from database...")
                    # Get search term from entry if not provided
                    if search_term is None:
                        search_term = self.search_entry.get().strip() if hasattr(self, 'search_entry') else ""
                    filter_type = self.search_filter.get() if hasattr(self, 'search_filter') else "all"
                    
                    # Load from Supabase with search - show all projects from all users
                    if search_term:
                        if filter_type == "work_order":
                            bids = self.db.search_bids_by_wo_number(self.user_id, search_term, all_bids=True)
                        elif filter_type == "property_address":
                            bids = self.db.search_bids_by_property_address(self.user_id, search_term, all_bids=True)
                        else:  # "all"
                            # Search both
                            wo_bids = self.db.search_bids_by_wo_number(self.user_id, search_term, all_bids=True)
                            addr_bids = self.db.search_bids_by_property_address(self.user_id, search_term, all_bids=True)
                            # Combine and deduplicate by wo_number
                            bids_dict = {}
                            for bid in wo_bids + addr_bids:
                                bids_dict[bid['wo_number']] = bid
                            bids = list(bids_dict.values())
                    else:
                        bids = self.db.get_user_bids(self.user_id, all_bids=True)
                    
                    # Sort by updated_at descending
                    bids.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
                    
                    print(f"Found {len(bids)} projects in database")
                    
                    # Save to cache after successful load (only if no search term)
                    if not search_term:
                        self.save_projects_to_cache(bids)
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error loading projects from database: {error_msg}")
                    # Fall back to cache if database fails
                    bids = self.load_projects_from_cache(search_term)
                    if bids:
                        print(f"Using cached projects (offline mode): {len(bids)} projects")
            
            # If no database, try cache
            if not bids:
                bids = self.load_projects_from_cache(search_term)
            
            # Update UI in main thread
            def update_ui():
                if bids:
                    self.display_bids(bids, search_term)
                    # Show notification if updated from database
                    if self.db and len(bids) != len(cache_bids if cache_bids else []):
                        print(f"Updated projects: {len(cache_bids if cache_bids else [])} -> {len(bids)}")
            
            self.root.after(0, update_ui)
        
        # Start loading in background thread
        thread = threading.Thread(target=load_in_thread, daemon=True)
        thread.start()
    
    def start_projects_auto_refresh(self):
        """Start frequent auto-refresh for projects list"""
        if hasattr(self.root, 'after'):
            self.root.after(self.projects_refresh_interval, self.projects_auto_refresh_callback)
    
    def projects_auto_refresh_callback(self):
        """Auto-refresh callback for projects list"""
        try:
            # Only refresh if the bid writer dashboard is currently shown
            if hasattr(self, 'recent_bids_list') and self.recent_bids_list.winfo_exists():
                print("Auto-refreshing projects list...")
                self.load_recent_bids_async()
        except Exception as e:
            print(f"Error in projects auto-refresh: {e}")
        
        # Schedule next refresh
        self.start_projects_auto_refresh()
    
    def save_projects_to_cache(self, bids):
        """Save projects list to local cache"""
        try:
            with open(self.projects_cache_file, 'w') as f:
                json.dump(bids, f, indent=2)
            print(f"Saved {len(bids)} projects to cache")
        except Exception as e:
            print(f"Error saving projects to cache: {e}")
    
    def load_projects_from_cache(self, search_term=None):
        """Load projects list from local cache, optionally filtered by search"""
        try:
            if os.path.exists(self.projects_cache_file):
                with open(self.projects_cache_file, 'r') as f:
                    bids = json.load(f)
                
                # Filter by search term if provided
                if search_term and search_term.strip():
                    search_term_lower = search_term.lower()
                    filtered_bids = []
                    for bid in bids:
                        wo_number = bid.get('wo_number', '').lower()
                        property_address = bid.get('property_address', '').lower()
                        if search_term_lower in wo_number or search_term_lower in property_address:
                            filtered_bids.append(bid)
                    bids = filtered_bids
                
                print(f"Loaded {len(bids)} projects from cache")
                return bids
        except Exception as e:
            print(f"Error loading projects from cache: {e}")
        return []
    
    def display_bids(self, bids, search_term=None):
        """Display bids in the UI - calls load_recent_bids with pre-loaded data"""
        # Call load_recent_bids with pre-loaded bids data to skip database query
        self.load_recent_bids(search_term=search_term, bids_data=bids)
    
    def load_recent_bids(self, search_term=None, bids_data=None):
        for widget in self.recent_bids_list.winfo_children():
            widget.destroy()
        
        # Clear selected bids when reloading
        self.selected_bids = {}
        if hasattr(self, 'bulk_delete_button'):
            self.bulk_delete_button.config(state='disabled')
        if hasattr(self, 'bulk_export_button'):
            self.bulk_export_button.config(state='disabled')
        if hasattr(self, 'select_all_var'):
            self.select_all_var.set(False)

        # Create table headers with consistent grid layout
        header_frame = tk.Frame(self.recent_bids_list, bg=self.colors['primary_blue'])
        header_frame.pack(fill='x')
        
        headers = ['Select', 'Work Order', 'Property Address', 'Client Code', 'WO Type', 'Bid Count', 'Created By', 'Last Modified', 'Delete', 'Export']
        
        # Column configuration: added Select column at the beginning
        # Weights: Select (1), Work Order (2), Property Address (4), Client Code (2), WO Type (2), Bid Count (1), Created By (2), Last Modified (2), Delete (1), Export (1)
        column_weights = [1, 2, 4, 2, 2, 1, 2, 2, 1, 1]
        column_minsizes = [50, 120, 180, 100, 120, 80, 120, 140, 70, 70]
        
        # Configure grid columns with weights and minimum sizes
        for i, (weight, minsize) in enumerate(zip(column_weights, column_minsizes)):
            header_frame.grid_columnconfigure(i, weight=weight, minsize=minsize)

        # Select All checkbox in header
        self.select_all_var = tk.BooleanVar()
        select_all_cb = tk.Checkbutton(header_frame, variable=self.select_all_var, command=self.toggle_select_all, bg=self.colors['primary_blue'], activebackground=self.colors['primary_blue'], highlightthickness=0, selectcolor=self.colors['primary_blue'])
        select_all_cb.grid(row=0, column=0, sticky='nsew', padx=8, pady=8)
        
        for i, header in enumerate(headers[1:], start=1):  # Start from 1 to skip Select header
            tk.Label(header_frame, text=header, font=("Arial", 11, "bold"), fg='white', bg=self.colors['primary_blue'], padx=8, pady=8, anchor='w', justify='left').grid(row=0, column=i, sticky='nsw')

        try:
            # Use provided bids_data if available (from cache), otherwise load from database
            if bids_data is not None:
                bids = bids_data
                # Sort by updated_at descending
                bids.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
            elif self.db and self.user_id:
                # Get search term from entry if not provided
                if search_term is None:
                    search_term = self.search_entry.get().strip() if hasattr(self, 'search_entry') else ""
                filter_type = self.search_filter.get() if hasattr(self, 'search_filter') else "all"
                
                # Load from Supabase with search - show all projects from all users
                if search_term:
                    if filter_type == "work_order":
                        bids = self.db.search_bids_by_wo_number(self.user_id, search_term, all_bids=True)
                    elif filter_type == "property_address":
                        bids = self.db.search_bids_by_property_address(self.user_id, search_term, all_bids=True)
                    else:  # "all"
                        # Search both
                        wo_bids = self.db.search_bids_by_wo_number(self.user_id, search_term, all_bids=True)
                        addr_bids = self.db.search_bids_by_property_address(self.user_id, search_term, all_bids=True)
                        # Combine and deduplicate by wo_number
                        bids_dict = {}
                        for bid in wo_bids + addr_bids:
                            bids_dict[bid['wo_number']] = bid
                        bids = list(bids_dict.values())
                else:
                    bids = self.db.get_user_bids(self.user_id, all_bids=True)
                
                # Sort by updated_at descending
                bids.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
                
                # Save to cache after database load (only if no search term)
                if not search_term:
                    self.save_projects_to_cache(bids)
            else:
                bids = []
            
            if not bids:
                tk.Label(self.recent_bids_list, text="No matching projects found." if search_term else "No projects found.", bg=self.colors['white'], fg=self.colors['gray_dark'], font=("Arial", 10, "italic")).pack(padx=10, pady=10)
                return
            
            for bid in bids:
                    wo_number = bid['wo_number']
                    
                    # Load full bid data to get all fields including property_address, client_code, wo_type
                    bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                    
                    # Get data from full bid_data first, then fallback to summary bid
                    property_address = ''
                    client_code = ''
                    wo_type = ''
                    
                    if bid_data:
                        property_address = bid_data.get('property_address') or ''
                        client_code = bid_data.get('client_code') or ''
                        wo_type = bid_data.get('wo_type') or ''
                    
                    # Fallback to summary bid if not found in full bid_data
                    if not property_address:
                        property_address = bid.get('property_address') or ''
                    if not client_code:
                        client_code = bid.get('client_code') or ''
                    if not wo_type:
                        wo_type = bid.get('wo_type') or ''
                    
                    # Parse updated_at timestamp - format to human-readable
                    updated_at = bid.get('updated_at', '')
                    modified_time = ''
                    if updated_at:
                        try:
                            # Try parsing ISO format timestamp
                            if 'T' in updated_at:
                                # ISO format: 2024-01-01T12:00:00+00:00
                                date_str = updated_at.split('T')[0]
                                # Parse date and format as "Dec 24, 2025"
                                dt = datetime.strptime(date_str, '%Y-%m-%d')
                                modified_time = dt.strftime('%b %d, %Y')
                            else:
                                # Try other formats
                                try:
                                    dt = datetime.strptime(updated_at[:10], '%Y-%m-%d')
                                    modified_time = dt.strftime('%b %d, %Y')
                                except:
                                    modified_time = updated_at[:10] if len(updated_at) >= 10 else updated_at
                        except Exception as e:
                            print(f"Error parsing date: {e}")
                            modified_time = updated_at[:10] if len(updated_at) >= 10 else updated_at
                    
                    # Get bid count from loaded bid_data
                    bid_count = 0
                    if bid_data and 'selected_items' in bid_data:
                        for category in bid_data['selected_items']:
                            bid_count += len(bid_data['selected_items'][category])
                    
                    # Create row with proper grid layout - must match header configuration
                    row_frame = tk.Frame(self.recent_bids_list, bg=self.colors['white'])
                    
                    # Use same column configuration as header for perfect alignment
                    column_weights = [1, 2, 4, 2, 2, 1, 2, 2, 1, 1]
                    column_minsizes = [50, 120, 180, 100, 120, 80, 120, 140, 70, 70]
                    
                    for i, (weight, minsize) in enumerate(zip(column_weights, column_minsizes)):
                        row_frame.grid_columnconfigure(i, weight=weight, minsize=minsize)
                    
                    # Get created_by username
                    created_by = bid.get('created_by_username') or ''
                    if not created_by or created_by == '':
                        # Try to get username from user_id if created_by_username is missing
                        bid_user_id = bid.get('user_id')
                        if bid_user_id and self.db:
                            try:
                                user = self.db.get_user_by_id(bid_user_id)
                                if user:
                                    created_by = user.get('username', 'N/A')
                                else:
                                    created_by = 'N/A'
                            except Exception as e:
                                print(f"Error fetching username for user_id {bid_user_id}: {e}")
                                created_by = 'N/A'
                        else:
                            created_by = 'N/A'
                    
                    # Ensure all values are strings and handle empty/None values
                    property_address = str(property_address).strip() if property_address and property_address != 'N/A' else ''
                    client_code = str(client_code).strip() if client_code and client_code != 'N/A' else ''
                    wo_type = str(wo_type).strip() if wo_type and wo_type != 'N/A' else ''
                    bid_count = str(bid_count) if bid_count else '0'
                    created_by = str(created_by).strip() if created_by and created_by != 'N/A' else ''
                    modified_time = str(modified_time).strip() if modified_time and modified_time != 'N/A' else ''
                    
                    # Make all labels clickable
                    def make_label_click_handler(wo):
                        return lambda e: self.open_existing_bid(wo)
                    
                    # Checkbox for selection
                    checkbox_var = tk.BooleanVar()
                    self.selected_bids[wo_number] = checkbox_var
                    checkbox = tk.Checkbutton(row_frame, variable=checkbox_var, command=self.update_bulk_delete_button_state, bg=self.colors['white'], activebackground=self.colors['white'], highlightthickness=0, selectcolor=self.colors['white'])
                    checkbox.grid(row=0, column=0, sticky='nsew', padx=8, pady=8)
                    
                    # Create all cells with consistent styling and proper padding
                    wo_label = tk.Label(row_frame, text=wo_number, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['primary_blue'], anchor='w', cursor="hand2", padx=8, pady=8)
                    wo_label.grid(row=0, column=1, sticky='nsew')
                    wo_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    addr_label = tk.Label(row_frame, text=property_address, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    addr_label.grid(row=0, column=2, sticky='nsew')
                    addr_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    client_label = tk.Label(row_frame, text=client_code, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    client_label.grid(row=0, column=3, sticky='nsew')
                    client_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    wo_type_label = tk.Label(row_frame, text=wo_type, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    wo_type_label.grid(row=0, column=4, sticky='nsew')
                    wo_type_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    count_label = tk.Label(row_frame, text=bid_count, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    count_label.grid(row=0, column=5, sticky='nsew')
                    count_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    created_by_label = tk.Label(row_frame, text=created_by, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['primary_blue'], anchor='w', cursor="hand2", padx=8, pady=8)
                    created_by_label.grid(row=0, column=6, sticky='nsew')
                    created_by_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    time_label = tk.Label(row_frame, text=modified_time, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    time_label.grid(row=0, column=7, sticky='nsew')
                    time_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    # Delete button (smaller size)
                    delete_button = tk.Button(row_frame, text="Delete", command=lambda wo=wo_number: self.delete_bid_state(wo), font=("Arial", 8), bg='#dc3545', fg='white', relief='flat', padx=4, pady=3)
                    delete_button.grid(row=0, column=8, sticky='nsew')

                    # Docs1 and Docs2 buttons container (replaces Export button)
                    docs_frame = tk.Frame(row_frame, bg=self.colors['white'])
                    docs_frame.grid(row=0, column=9, sticky='nsew')
                    
                    # Docs1 button (small)
                    docs1_button = tk.Button(docs_frame, text="Docs1", command=lambda wo=wo_number: self.export_to_docs1(wo), font=("Arial", 7), bg=self.colors['primary_blue'], fg='white', relief='flat', padx=3, pady=2)
                    docs1_button.pack(side='left', fill='both', expand=True, padx=1)
                    
                    # Docs2 button (small)
                    docs2_button = tk.Button(docs_frame, text="Docs2", command=lambda wo=wo_number: self.export_to_docs2(wo), font=("Arial", 7), bg=self.colors['primary_blue'], fg='white', relief='flat', padx=3, pady=2)
                    docs2_button.pack(side='left', fill='both', expand=True, padx=1)
                    
                    # Bind the entire row to open the bid
                    def make_open_handler(wo):
                        return lambda e: self.open_existing_bid(wo)
                    row_frame.bind("<Button-1>", make_open_handler(wo_number))
                    
                    # Pack the row frame
                    row_frame.pack(fill='x', pady=1)
            else:
                # Fallback to local JSON files if database not available
                files = [f for f in os.listdir(self.app_data_dir) if f.startswith("WO_") and f.endswith(".json")]

                # Filter files based on search term
                if search_term:
                    files = [f for f in files if search_term.lower() in f.lower()]

                files.sort(key=lambda f: os.path.getmtime(os.path.join(self.app_data_dir, f)), reverse=True)
                
                if not files:
                    tk.Label(self.recent_bids_list, text="No matching projects found.", bg=self.colors['white'], fg=self.colors['gray_dark'], font=("Arial", 10, "italic")).pack(padx=10, pady=10)
                    return

                for i, file in enumerate(files):
                    wo_number = file.replace("WO_", "").replace(".json", "")
                    file_path = os.path.join(self.app_data_dir, file)
                    modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M')

                    # Load the JSON to get bid count
                    bid_count = 0
                    try:
                        with open(file_path, 'r') as f:
                            state = json.load(f)
                            if 'selected_items' in state:
                                for category in state['selected_items']:
                                    bid_count += len(state['selected_items'][category])
                    except Exception as e:
                        print(f"Error loading bid count for {file}: {e}")
                        bid_count = "N/A"

                    row_frame = tk.Frame(self.recent_bids_list, bg=self.colors['white'])
                    row_frame.pack(fill='x', pady=1)
                    
                    # Use same column configuration as header for perfect alignment
                    column_weights = [1, 2, 4, 2, 2, 1, 2, 2, 1, 1]
                    column_minsizes = [50, 120, 180, 100, 120, 80, 120, 140, 70, 70]
                    
                    for i, (weight, minsize) in enumerate(zip(column_weights, column_minsizes)):
                        row_frame.grid_columnconfigure(i, weight=weight, minsize=minsize)
                    
                    # Checkbox for selection
                    checkbox_var = tk.BooleanVar()
                    self.selected_bids[wo_number] = checkbox_var
                    checkbox = tk.Checkbutton(row_frame, variable=checkbox_var, command=self.update_bulk_delete_button_state, bg=self.colors['white'], activebackground=self.colors['white'], highlightthickness=0, selectcolor=self.colors['white'])
                    checkbox.grid(row=0, column=0, sticky='nsew', padx=8, pady=8)
                    
                    # Bind the entire row to open the bid - use default parameter to capture wo_number correctly
                    def make_open_handler(wo):
                        return lambda e: self.open_existing_bid(wo)
                    row_frame.bind("<Button-1>", make_open_handler(wo_number))

                    # Make all labels clickable
                    def make_label_click_handler(wo):
                        return lambda e: self.open_existing_bid(wo)
                    
                    wo_label = tk.Label(row_frame, text=wo_number, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['primary_blue'], anchor='w', cursor="hand2", padx=8, pady=8)
                    wo_label.grid(row=0, column=1, sticky='nsew')
                    wo_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    # Empty cells for columns that don't exist in local file mode
                    addr_label = tk.Label(row_frame, text="", font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', padx=8, pady=8)
                    addr_label.grid(row=0, column=2, sticky='nsew')
                    
                    client_label = tk.Label(row_frame, text="", font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', padx=8, pady=8)
                    client_label.grid(row=0, column=3, sticky='nsew')
                    
                    wo_type_label = tk.Label(row_frame, text="", font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', padx=8, pady=8)
                    wo_type_label.grid(row=0, column=4, sticky='nsew')
                    
                    count_label = tk.Label(row_frame, text=str(bid_count), font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    count_label.grid(row=0, column=5, sticky='nsew')
                    count_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    created_by_label = tk.Label(row_frame, text="", font=("Arial", 10), bg=self.colors['white'], fg=self.colors['primary_blue'], anchor='w', padx=8, pady=8)
                    created_by_label.grid(row=0, column=6, sticky='nsew')
                    
                    time_label = tk.Label(row_frame, text=modified_time, font=("Arial", 10), bg=self.colors['white'], fg=self.colors['gray_dark'], anchor='w', cursor="hand2", padx=8, pady=8)
                    time_label.grid(row=0, column=7, sticky='nsew')
                    time_label.bind("<Button-1>", make_label_click_handler(wo_number))
                    
                    # Delete button (smaller size)
                    delete_button = tk.Button(row_frame, text="Delete", command=lambda wo=wo_number: self.delete_bid_state(wo), font=("Arial", 8), bg='#dc3545', fg='white', relief='flat', padx=4, pady=3)
                    delete_button.grid(row=0, column=8, sticky='nsew')

                    # Docs1 and Docs2 buttons container (replaces Export button)
                    docs_frame = tk.Frame(row_frame, bg=self.colors['white'])
                    docs_frame.grid(row=0, column=9, sticky='nsew')
                    
                    # Docs1 button (small)
                    docs1_button = tk.Button(docs_frame, text="Docs1", command=lambda wo=wo_number: self.export_to_docs1(wo), font=("Arial", 7), bg=self.colors['primary_blue'], fg='white', relief='flat', padx=3, pady=2)
                    docs1_button.pack(side='left', fill='both', expand=True, padx=1)
                    
                    # Docs2 button (small)
                    docs2_button = tk.Button(docs_frame, text="Docs2", command=lambda wo=wo_number: self.export_to_docs2(wo), font=("Arial", 7), bg=self.colors['primary_blue'], fg='white', relief='flat', padx=3, pady=2)
                    docs2_button.pack(side='left', fill='both', expand=True, padx=1)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load recent projects: {e}")

    # --- Dashboard Card Helper ---
    def _create_dashboard_card(self, parent, row, col, icon, title, subtitle, command):
        card_bg = self.colors['white'] if 'white' in self.colors else '#FFFFFF'
        shadow_color = self.colors.get('shadow_gray', self.colors.get('gray_light', '#D0D0D0'))

        # Container cell (uses grid) to keep shape
        cell = tk.Frame(parent, bg=self.colors['background'])
        cell.grid(row=row, column=col, padx=10, pady=10, sticky='nsew')
        parent.grid_rowconfigure(row, weight=1)
        cell.grid_propagate(True)

        # Shadow (hidden initially)
        shadow = tk.Frame(cell, bg=shadow_color, bd=0, relief='flat')
        shadow.place_forget()

        # Card on top
        card = tk.Frame(cell, bg=card_bg, relief='flat', bd=0,
                        highlightthickness=1, highlightbackground=self.colors['gray_light'])
        card.place(relx=0, rely=0, relwidth=1, relheight=1)

        inner = tk.Frame(card, bg=card_bg)
        inner.pack(fill='both', expand=True, padx=16, pady=14)

        # Icon
        tk.Label(inner, text=icon, font=("Segoe UI", 22), bg=card_bg).pack(anchor='w')
        # Title
        tk.Label(inner, text=title, font=("Segoe UI", 12, "bold"),
                 bg=card_bg, fg=self.colors['primary_blue']).pack(anchor='w', pady=(6, 2))
        # Subtitle
        tk.Label(inner, text=subtitle, font=("Segoe UI", 10),
                 bg=card_bg, fg=self.colors['gray_dark']).pack(anchor='w')

        # Hover effects (outline + shadow) across the whole section
        def on_enter(_e=None):
            # Outline
            card.configure(highlightbackground=self.colors.get('nav_hover', self.colors['primary_blue']),
                           highlightthickness=2)
            # Shadow slightly offset
            try:
                shadow.place(relx=0, rely=0, relwidth=1, relheight=1, x=3, y=3)
                shadow.lift()  # make sure shadow draws under card
                card.lift()
            except Exception:
                pass

        def on_leave(_e=None):
            card.configure(highlightbackground=self.colors['gray_light'], highlightthickness=1)
            try:
                shadow.place_forget()
            except Exception:
                pass

        # Click handler for all card area
        def handle_click(_e=None):
            try:
                command()
            except Exception as ex:
                messagebox.showerror("Action", f"Failed to open: {ex}")

        # Bind recursively so hovering anywhere triggers effects smoothly
        def bind_recursive(widget):
            widget.bind('<Enter>', lambda e: on_enter())
            widget.bind('<Leave>', lambda e: on_leave())
            widget.bind('<Button-1>', handle_click)
            for child in widget.winfo_children():
                bind_recursive(child)

        bind_recursive(cell)

    def toggle_select_all(self):
        """Toggle selection of all checkboxes"""
        select_all = self.select_all_var.get()
        for checkbox_var in self.selected_bids.values():
            checkbox_var.set(select_all)
        self.update_bulk_delete_button_state()
    
    def update_bulk_delete_button_state(self):
        """Update the bulk delete and export buttons state based on selected items"""
        selected_count = sum(1 for var in self.selected_bids.values() if var.get())
        if hasattr(self, 'bulk_delete_button'):
            if selected_count > 0:
                self.bulk_delete_button.config(state='normal', text=f"Delete Selected ({selected_count})")
            else:
                self.bulk_delete_button.config(state='disabled', text="Delete Selected")
        if hasattr(self, 'bulk_export_button'):
            if selected_count > 0:
                self.bulk_export_button.config(state='normal', text=f"Export Selected ({selected_count})")
            else:
                self.bulk_export_button.config(state='disabled', text="Export Selected")
    
    def delete_selected_bids(self):
        """Delete all selected bids"""
        selected_wo_numbers = [wo_number for wo_number, var in self.selected_bids.items() if var.get()]
        
        if not selected_wo_numbers:
            messagebox.showwarning("No Selection", "Please select at least one project to delete.")
            return
        
        count = len(selected_wo_numbers)
        if not messagebox.askyesno("Confirm Delete", f"Are you sure you want to permanently delete {count} selected project(s)?"):
            return
        
        deleted_count = 0
        failed_count = 0
        failed_items = []
        
        try:
            for wo_number in selected_wo_numbers:
                try:
                    if self.db and self.user_id:
                        # Delete from Supabase
                        self.db.delete_bid(wo_number, self.user_id)
                        deleted_count += 1
                    else:
                        # Fallback to local file deletion
                        file_path = os.path.join(self.app_data_dir, f"WO_{wo_number}.json")
                        if os.path.exists(file_path):
                            os.remove(file_path)
                            deleted_count += 1
                        else:
                            failed_count += 1
                            failed_items.append(wo_number)
                except Exception as e:
                    failed_count += 1
                    failed_items.append(wo_number)
                    print(f"Error deleting bid {wo_number}: {e}")
            
            # Show result message
            if deleted_count == count:
                messagebox.showinfo("Success", f"Successfully deleted {deleted_count} project(s).")
            elif deleted_count > 0:
                messagebox.showwarning("Partial Success", f"Deleted {deleted_count} project(s). Failed to delete {failed_count} project(s): {', '.join(failed_items)}")
            else:
                messagebox.showerror("Error", f"Failed to delete all selected projects.")
            
            # Reload the list
            self.load_recent_bids()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete selected bids: {e}")

    def delete_bid_state(self, wo_number):
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to permanently delete the bid for WO# {wo_number}?"):
            try:
                if self.db and self.user_id:
                    # Delete from Supabase
                    self.db.delete_bid(wo_number, self.user_id)
                    messagebox.showinfo("Success", f"Bid for WO# {wo_number} has been deleted.")
                else:
                    # Fallback to local file deletion
                    file_path = os.path.join(self.app_data_dir, f"WO_{wo_number}.json")
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        messagebox.showinfo("Success", f"Bid for WO# {wo_number} has been deleted.")
                    else:
                        messagebox.showerror("Error", f"File not found for WO# {wo_number}.")
                self.load_recent_bids()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete bid: {e}")

    def export_to_docs1(self, wo_number):
        """Generate Docs1 format Word document for a bid"""
        try:
            # Load bid data
            property_address = None
            client_code = None
            wo_type = None
            bid_user_id = self.user_id
            
            if self.db:
                try:
                    bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                    if not bid_data:
                        messagebox.showerror("Error", f"Could not find project with Work Order: {wo_number}")
                        return
                    property_address = bid_data.get('property_address')
                    client_code = bid_data.get('client_code')
                    wo_type = bid_data.get('wo_type')
                    bid_user_id = bid_data.get('user_id', self.user_id)
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load project data: {e}")
                    return
            
            # Create a hidden window for BidWriterApp
            hidden_window = tk.Toplevel(self.root)
            hidden_window.withdraw()  # Hide the window immediately
            
            # Create BidWriterApp instance - it will load state automatically
            bid_writer = BidWriterApp(hidden_window, self.username, wo_number, 
                                     property_address=property_address, 
                                     user_id=bid_user_id, 
                                     client_code=client_code, 
                                     wo_type=wo_type,
                                     on_save_callback=None)
            
            # Wait for state to load (BidWriterApp loads state in __init__ if wo_number_to_load is provided)
            hidden_window.update_idletasks()
            time.sleep(0.3)  # Give it time to load state and initialize UI
            
            # Call save_to_docs1
            bid_writer.save_to_docs1()
            
            # Clean up
            hidden_window.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate Docs1: {e}")
    
    def export_to_docs2(self, wo_number):
        """Generate Docs2 format Word document for a bid"""
        try:
            # Load bid data
            property_address = None
            client_code = None
            wo_type = None
            bid_user_id = self.user_id
            
            if self.db:
                try:
                    bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                    if not bid_data:
                        messagebox.showerror("Error", f"Could not find project with Work Order: {wo_number}")
                        return
                    property_address = bid_data.get('property_address')
                    client_code = bid_data.get('client_code')
                    wo_type = bid_data.get('wo_type')
                    bid_user_id = bid_data.get('user_id', self.user_id)
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load project data: {e}")
                    return
            
            # Create a hidden window for BidWriterApp
            hidden_window = tk.Toplevel(self.root)
            hidden_window.withdraw()  # Hide the window immediately
            
            # Create BidWriterApp instance - it will load state automatically
            bid_writer = BidWriterApp(hidden_window, self.username, wo_number, 
                                     property_address=property_address, 
                                     user_id=bid_user_id, 
                                     client_code=client_code, 
                                     wo_type=wo_type,
                                     on_save_callback=None)
            
            # Wait for state to load (BidWriterApp loads state in __init__ if wo_number_to_load is provided)
            hidden_window.update_idletasks()
            time.sleep(0.3)  # Give it time to load state and initialize UI
            
            # Call save_to_docs2
            bid_writer.save_to_docs2()
            
            # Clean up
            hidden_window.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate Docs2: {e}")
    
    def export_selected_bids(self):
        """Export all selected bids to a single Word document in Docs2 format"""
        selected_wo_numbers = [wo_number for wo_number, var in self.selected_bids.items() if var.get()]
        
        if not selected_wo_numbers:
            messagebox.showwarning("No Selection", "Please select at least one project to export.")
            return
        
        try:
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import tempfile
                use_docx = True
            except ImportError:
                use_docx = False
                messagebox.showinfo("Info", "python-docx not found. Please install it to export documents.\nInstall: pip install python-docx")
                return
            
            if use_docx:
                doc = Document()
                
                # Title
                doc.add_heading("Techvengers Bid Proposal - Multiple Projects", 0)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                
                all_bid_items = []  # List to collect all bids from all projects
                project_info_list = []  # List to store project info for each WO
                
                # Process each selected WO
                for wo_number in selected_wo_numbers:
                    try:
                        # Load bid data
                        property_address = None
                        client_code = None
                        wo_type = None
                        bid_user_id = self.user_id
                        created_at = None
                        updated_at = None
                        processor_name = self.username
                        
                        if self.db:
                            try:
                                bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                                if bid_data:
                                    property_address = bid_data.get('property_address', '')
                                    client_code = bid_data.get('client_code', '')
                                    wo_type = bid_data.get('wo_type', '')
                                    bid_user_id = bid_data.get('user_id', self.user_id)
                                    created_at = bid_data.get('created_at', '')
                                    updated_at = bid_data.get('updated_at', '')
                                    # Get processor name from created_by_username if available
                                    processor_name = bid_data.get('created_by_username', self.username)
                                    
                                    # Store project info
                                    project_info_list.append({
                                        'wo_number': wo_number,
                                        'property_address': property_address or 'N/A',
                                        'client_code': client_code or 'N/A',
                                        'wo_type': wo_type or 'N/A',
                                        'created_at': created_at,
                                        'updated_at': updated_at,
                                        'processor_name': processor_name or 'N/A'
                                    })
                                    
                                    # Create hidden BidWriterApp to extract bids
                                    hidden_window = tk.Toplevel(self.root)
                                    hidden_window.withdraw()
                                    
                                    bid_writer = BidWriterApp(hidden_window, processor_name or self.username, wo_number, 
                                                             property_address=property_address, 
                                                             user_id=bid_user_id, 
                                                             client_code=client_code, 
                                                             wo_type=wo_type,
                                                             on_save_callback=None)
                                    
                                    hidden_window.update_idletasks()
                                    time.sleep(0.3)
                                    
                                    # Extract bids from this project
                                    if hasattr(bid_writer, 'selected_items') and bid_writer.selected_items:
                                        # Process all selected items
                                        conjunction_groups = {}
                                        standalone_bids = []
                                        
                                        for category, category_items in bid_writer.selected_items.items():
                                            for item in category_items.values():
                                                if item.get("selected", False):
                                                    conjunction_key = item.get("conjunction_key", tk.StringVar()).get().strip().upper() if hasattr(item.get("conjunction_key"), 'get') else ""
                                                    if conjunction_key:
                                                        if conjunction_key not in conjunction_groups:
                                                            conjunction_groups[conjunction_key] = []
                                                        conjunction_groups[conjunction_key].append(item)
                                                    else:
                                                        standalone_bids.append(item)
                                        
                                        # Sort and combine
                                        all_items_for_wo = []
                                        for key in sorted(conjunction_groups.keys()):
                                            all_items_for_wo.extend(sorted(conjunction_groups[key], key=lambda x: x.get('instance_info', {}).get('key', '')))
                                        all_items_for_wo.extend(sorted(standalone_bids, key=lambda x: x.get('instance_info', {}).get('key', '')))
                                        
                                        # Extract bid data for each item
                                        for item in all_items_for_wo:
                                            try:
                                                # Get bid text
                                                bid_text, _ = bid_writer._get_item_bid_data(item)
                                                bid_text_without_price = bid_writer._extract_bid_text_without_price(bid_text)
                                                if not bid_text_without_price:
                                                    bid_text_without_price = bid_text
                                                
                                                # Calculate price
                                                total_price = bid_writer._calculate_item_price(item)
                                                
                                                # Get photo
                                                instance_key = item.get('instance_info', {}).get('key', '')
                                                category_name = next((cat for cat, items in bid_writer.selected_items.items() if instance_key in items), '')
                                                photo_key = f"{category_name}_{instance_key}"
                                                photo_data = None
                                                if hasattr(bid_writer, 'item_photos') and photo_key in bid_writer.item_photos:
                                                    photo_data = bid_writer.item_photos[photo_key]
                                                
                                                all_bid_items.append({
                                                    'wo_number': wo_number,
                                                    'text': bid_text_without_price,
                                                    'price': total_price,
                                                    'photo': photo_data
                                                })
                                            except Exception as e:
                                                print(f"Error processing item for WO {wo_number}: {e}")
                                                continue
                                    
                                    hidden_window.destroy()
                            except Exception as e:
                                print(f"Error loading bid for WO {wo_number}: {e}")
                                continue
                    
                    except Exception as e:
                        print(f"Error processing WO {wo_number}: {e}")
                        continue
                
                if not all_bid_items:
                    messagebox.showwarning("No Bids", "No bids found in the selected projects.")
                    return
                
                # Add project information section for each project, followed by its bids
                global_bid_number = 1
                for i, proj_info in enumerate(project_info_list, 1):
                    if i > 1:
                        doc.add_page_break()
                    
                    # Project heading
                    doc.add_heading(f"Project {i}: WO# {proj_info['wo_number']}", level=1)
                    doc.add_paragraph()
                    
                    # Project Information Table
                    info_table = doc.add_table(rows=4, cols=2)
                    info_table.style = 'Table Grid'
                    info_table.columns[0].width = Inches(2.0)
                    info_table.columns[1].width = Inches(5.0)
                    
                    info_table.rows[0].cells[0].paragraphs[0].clear()
                    run1 = info_table.rows[0].cells[0].paragraphs[0].add_run("Work Order (WO):")
                    run1.bold = True
                    info_table.rows[0].cells[1].text = proj_info['wo_number']
                    
                    info_table.rows[1].cells[0].paragraphs[0].clear()
                    run1 = info_table.rows[1].cells[0].paragraphs[0].add_run("Property Address:")
                    run1.bold = True
                    info_table.rows[1].cells[1].text = proj_info['property_address']
                    
                    info_table.rows[2].cells[0].paragraphs[0].clear()
                    run1 = info_table.rows[2].cells[0].paragraphs[0].add_run("Work Order Type:")
                    run1.bold = True
                    info_table.rows[2].cells[1].text = proj_info['wo_type']
                    
                    info_table.rows[3].cells[0].paragraphs[0].clear()
                    run1 = info_table.rows[3].cells[0].paragraphs[0].add_run("Client Code:")
                    run1.bold = True
                    info_table.rows[3].cells[1].text = proj_info['client_code']
                    
                    doc.add_paragraph()
                    
                    # Dates and Processor Table
                    dates_table = doc.add_table(rows=3, cols=2)
                    dates_table.style = 'Table Grid'
                    dates_table.columns[0].width = Inches(2.5)
                    dates_table.columns[1].width = Inches(4.5)
                    
                    dates_table.rows[0].cells[0].paragraphs[0].clear()
                    run1 = dates_table.rows[0].cells[0].paragraphs[0].add_run("Project Creation Date:")
                    run1.bold = True
                    # Format date
                    try:
                        if proj_info['created_at'] and 'T' in str(proj_info['created_at']):
                            date_str = str(proj_info['created_at']).split('T')[0]
                            dt = datetime.strptime(date_str, '%Y-%m-%d')
                            formatted_date = dt.strftime('%b %d, %Y')
                        else:
                            formatted_date = str(proj_info['created_at']) if proj_info['created_at'] else 'N/A'
                    except:
                        formatted_date = str(proj_info['created_at']) if proj_info['created_at'] else 'N/A'
                    dates_table.rows[0].cells[1].text = formatted_date
                    
                    dates_table.rows[1].cells[0].paragraphs[0].clear()
                    run1 = dates_table.rows[1].cells[0].paragraphs[0].add_run("Project Last Modification Date:")
                    run1.bold = True
                    # Format date
                    try:
                        if proj_info['updated_at'] and 'T' in str(proj_info['updated_at']):
                            date_str = str(proj_info['updated_at']).split('T')[0]
                            dt = datetime.strptime(date_str, '%Y-%m-%d')
                            formatted_date = dt.strftime('%b %d, %Y')
                        else:
                            formatted_date = str(proj_info['updated_at']) if proj_info['updated_at'] else 'N/A'
                    except:
                        formatted_date = str(proj_info['updated_at']) if proj_info['updated_at'] else 'N/A'
                    dates_table.rows[1].cells[1].text = formatted_date
                    
                    dates_table.rows[2].cells[0].paragraphs[0].clear()
                    run1 = dates_table.rows[2].cells[0].paragraphs[0].add_run("Processor Name:")
                    run1.bold = True
                    dates_table.rows[2].cells[1].text = proj_info['processor_name']
                    
                    doc.add_paragraph()
                    
                    # Get bids for this WO
                    wo_bids = [bid for bid in all_bid_items if bid['wo_number'] == proj_info['wo_number']]
                    
                    # Summary Section
                    total_items = len(wo_bids)
                    summary_para1 = doc.add_paragraph()
                    summary_para1.add_run("Summary: ").bold = True
                    summary_para1.add_run(f"Total Bid Count: {total_items}")
                    
                    doc.add_paragraph()
                    
                    # Add bids for this project in Docs2 format (text -> price -> photo)
                    for bid_item in wo_bids:
                        # Add bid number and text
                        bid_paragraph = doc.add_paragraph()
                        bid_paragraph.add_run(f"{global_bid_number}. {bid_item['text']}").bold = False
                        
                        # Add price below bid text
                        price_paragraph = doc.add_paragraph()
                        price_run = price_paragraph.add_run(f"Price ${bid_item['price']:.2f}")
                        price_run.bold = True
                        
                        # Add photo below price if available
                        if bid_item['photo'] and bid_item['photo'].get('original'):
                            try:
                                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                                    temp_path = temp_file.name
                                    bid_item['photo']['original'].save(temp_path)
                                
                                photo_paragraph = doc.add_paragraph()
                                photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                photo_paragraph.add_run().add_picture(temp_path, width=Inches(4.0))
                                
                                os.remove(temp_path)
                            except Exception as e:
                                print(f"Error adding image: {e}")
                                error_para = doc.add_paragraph()
                                error_para.add_run("Error loading image").italic = True
                        
                        # Add spacing between items
                        doc.add_paragraph()
                        global_bid_number += 1
                
                doc.add_paragraph()
                footer = doc.add_paragraph('Generated by Techvengers Bid Writer')
                footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_run = footer.runs[0]
                footer_run.italic = True
                
                # Generate filename
                date_str = datetime.now().strftime("%d%b")
                default_filename = f"Bids_MultipleProjects_{date_str}.docx"
                
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
                    initialfile=default_filename,
                    title="Save Combined Bids Document (Docs2 Format)"
                )
                
                if file_path:
                    doc.save(file_path)
                    messagebox.showinfo("Success", f"Combined bids saved successfully to:\n{file_path}")
                    
                    if messagebox.askyesno("Open File", "Would you like to open the saved document?"):
                        try:
                            if os.name == 'nt':
                                os.startfile(file_path)
                            elif sys.platform == 'darwin':
                                os.system(f'open "{file_path}"')
                            else:
                                os.system(f'xdg-open "{file_path}"')
                        except Exception as e:
                            messagebox.showinfo("File Saved", f"Document saved successfully!\nLocation: {file_path}")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export selected bids: {e}")
            import traceback
            traceback.print_exc()

    def create_new_bid(self):
        """Show dialog to get Work Order, Property Address, Client Code, and WO Type, then open bid writer"""
        from tkinter import ttk
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Create New Bid")
        dialog.geometry("580x560")
        dialog.configure(bg=self.colors['background'])
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        
        # Center the dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (580 // 2)
        y = (dialog.winfo_screenheight() // 2) - (600 // 2)
        dialog.geometry(f"580x600+{x}+{y}")
        
        # Header frame with title (blue background)
        header_frame = tk.Frame(dialog, bg=self.colors['primary_blue'], height=65)
        header_frame.pack(fill='x')
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(header_frame, text="Create New Bid", 
                              font=("Segoe UI", 20, "bold"), 
                              bg=self.colors['primary_blue'], fg='white')
        title_label.pack(expand=True, pady=18)
        
        # Form frame with better padding and spacing
        form_frame = tk.Frame(dialog, bg=self.colors['background'])
        form_frame.pack(pady=25, padx=40, fill='both', expand=False)
        
        # Work Order
        wo_label = tk.Label(form_frame, text="Work Order:", 
                           font=("Segoe UI", 11, "bold"),
                           bg=self.colors['background'], 
                           fg=self.colors['text_primary'], anchor='w')
        wo_label.pack(fill='x', pady=(0, 8))
        wo_entry = tk.Entry(form_frame, font=("Segoe UI", 11), 
                           relief="solid", bd=1, 
                           highlightthickness=1,
                           highlightcolor=self.colors['primary_blue'],
                           highlightbackground='#CCCCCC',
                           bg=self.colors['white'], 
                           fg=self.colors['text_primary'])
        wo_entry.pack(fill='x', ipady=8, pady=(0, 18))
        wo_entry.focus()
        
        # Property Address
        address_label = tk.Label(form_frame, text="Property Address:", 
                                font=("Segoe UI", 11, "bold"),
                                bg=self.colors['background'], 
                                fg=self.colors['text_primary'], anchor='w')
        address_label.pack(fill='x', pady=(0, 8))
        address_entry = tk.Entry(form_frame, font=("Segoe UI", 11), 
                                relief="solid", bd=1,
                                highlightthickness=1,
                                highlightcolor=self.colors['primary_blue'],
                                highlightbackground='#CCCCCC',
                                bg=self.colors['white'], 
                                fg=self.colors['text_primary'])
        address_entry.pack(fill='x', ipady=8, pady=(0, 18))
        
        # Client Code
        client_label = tk.Label(form_frame, text="Client Code:", 
                               font=("Segoe UI", 11, "bold"),
                               bg=self.colors['background'], 
                               fg=self.colors['text_primary'], anchor='w')
        client_label.pack(fill='x', pady=(0, 8))
        client_entry = tk.Entry(form_frame, font=("Segoe UI", 11), 
                               relief="solid", bd=1,
                               highlightthickness=1,
                               highlightcolor=self.colors['primary_blue'],
                               highlightbackground='#CCCCCC',
                               bg=self.colors['white'], 
                               fg=self.colors['text_primary'])
        client_entry.pack(fill='x', ipady=8, pady=(0, 18))
        
        # WO Type dropdown
        wo_type_label = tk.Label(form_frame, text="WO Type:", 
                                font=("Segoe UI", 11, "bold"),
                                bg=self.colors['background'], 
                                fg=self.colors['text_primary'], anchor='w')
        wo_type_label.pack(fill='x', pady=(0, 8))
        
        wo_type_options = [
            "Initial REO Service",
            "Initial Sales Clean",
            "Initial Secure",
            "Initial Services",
            "Inspection",
            "Initial Grass Cut",
            "Final Secure",
            "Eviction",
            "Bid Request (VRFY)",
            "Bid Request",
            "Bid Approval"
        ]
        
        wo_type_var = tk.StringVar()
        wo_type_combo = ttk.Combobox(form_frame, textvariable=wo_type_var, 
                                     values=wo_type_options, font=("Segoe UI", 11),
                                     state="readonly")
        wo_type_combo.pack(fill='x', ipady=8, pady=(0, 30))
        wo_type_combo.current(0)  # Set default to first option
        
        # Style the combobox for better appearance
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TCombobox', 
                       fieldbackground='white',
                       borderwidth=1,
                       relief='solid',
                       padding=8)
        style.map('TCombobox',
                 fieldbackground=[('readonly', 'white')],
                 bordercolor=[('focus', self.colors['primary_blue']), ('!focus', '#CCCCCC')])
        
        # Buttons frame with better spacing - ensure it's visible
        buttons_frame = tk.Frame(dialog, bg=self.colors['background'])
        buttons_frame.pack(pady=(20, 25), fill='x', padx=40)
        
        def create_bid():
            wo_number = wo_entry.get().strip()
            property_address = address_entry.get().strip()
            client_code = client_entry.get().strip()
            wo_type = wo_type_var.get().strip()
            
            if not wo_number:
                messagebox.showwarning("Warning", "Please enter a Work Order Number")
                wo_entry.focus()
                return
            
            if not property_address:
                messagebox.showwarning("Warning", "Please enter a Property Address")
                address_entry.focus()
                return
            
            # Create new project in database with all fields
            if self.db and self.user_id:
                try:
                    # Create initial project entry with all 4 fields
                    self.db.create_project(wo_number, self.user_id, property_address, 
                                         client_code, wo_type, self.username)
                except Exception as e:
                    print(f"Warning: Could not create project in database: {e}")
                    # Continue anyway - project will be created when first saved
            
            dialog.destroy()
            new_window = tk.Toplevel(self.root)
            BidWriterApp(new_window, self.username, wo_number_to_load=wo_number, 
                        property_address=property_address, user_id=self.user_id, 
                        client_code=client_code, wo_type=wo_type,
                        on_save_callback=self.load_recent_bids)
        
        def cancel():
            dialog.destroy()
        
        # Cancel button (improved styling)
        cancel_btn = tk.Button(buttons_frame, text="Cancel", 
                              font=("Segoe UI", 11, "bold"),
                              bg=self.colors['gray_medium'], fg='white', 
                              relief="flat",
                              cursor="hand2", 
                              padx=25, pady=10, 
                              command=cancel,
                              activebackground='#999999')
        cancel_btn.pack(side='left', padx=5)
        
        # Create button (improved styling)
        create_btn = tk.Button(buttons_frame, text="Create New Project", 
                              font=("Segoe UI", 11, "bold"),
                              bg=self.colors['primary_blue'], fg='white', 
                              relief="flat",
                              cursor="hand2", 
                              padx=25, pady=10, 
                              command=create_bid,
                              activebackground=self.colors.get('nav_hover', self.colors['primary_blue']))
        create_btn.pack(side='left', padx=5)
        
        # Bind Enter key navigation
        wo_entry.bind("<Return>", lambda e: address_entry.focus())
        address_entry.bind("<Return>", lambda e: client_entry.focus())
        client_entry.bind("<Return>", lambda e: wo_type_combo.focus())
        wo_type_combo.bind("<Return>", lambda e: create_bid())
        dialog.bind("<Escape>", lambda e: cancel())
        
    def open_existing_bid(self, wo_number):
        """Open existing bid and load all project fields if available - can open from any user"""
        property_address = None
        client_code = None
        wo_type = None
        bid_user_id = self.user_id
        
        if self.db:
            try:
                # Load bid from any user (all_users=True)
                bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                if bid_data:
                    property_address = bid_data.get('property_address')
                    client_code = bid_data.get('client_code')
                    wo_type = bid_data.get('wo_type')
                    # Use the bid's user_id if available, otherwise use current user_id
                    bid_user_id = bid_data.get('user_id', self.user_id)
                else:
                    # Bid not found - show error
                    messagebox.showerror("Error", f"Could not find project with Work Order: {wo_number}")
                    return
            except Exception as e:
                error_msg = str(e)
                print(f"Could not load project data: {e}")
                messagebox.showerror("Error", f"Failed to load project:\n{error_msg}")
                return
        
        try:
            new_window = tk.Toplevel(self.root)
            BidWriterApp(new_window, self.username, wo_number, property_address=property_address, 
                        user_id=bid_user_id, client_code=client_code, wo_type=wo_type,
                        on_save_callback=self.load_recent_bids)
        except Exception as e:
            error_msg = str(e)
            print(f"Error opening bid writer: {e}")
            messagebox.showerror("Error", f"Failed to open project:\n{error_msg}")

    def open_letterhead_bid(self): 
        self.clear_content_frame()
        letterhead_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        letterhead_frame.pack(fill="both", expand=True)
        
        # Back to Dashboard bar
        self._add_back_bar(letterhead_frame, "Letterheads")
        
        # Create LetterheadBidModule in the frame
        LetterheadBidModule(letterhead_frame, username=self.username, user_id=self.user_id if hasattr(self, 'user_id') else None)

    def open_notice_board(self):
        """Display notice board in main content frame"""
        self.clear_content_frame()
        notice_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        notice_frame.pack(fill="both", expand=True)
        
        # Back to Dashboard bar
        self._add_back_bar(notice_frame, "Notice Board")
        
        # Create NoticeBoardModule in the frame
        NoticeBoardModule(notice_frame, username=self.username, user_id=self.user_id if hasattr(self, 'user_id') else None, colors=self.colors)

    def show_price_sheet_page(self):
        """Display the Price Sheet page with 3 cards"""
        self.clear_content_frame()
        price_sheet_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        price_sheet_frame.pack(fill="both", expand=True)
        
        # Back to Dashboard bar
        self._add_back_bar(price_sheet_frame, "Price Sheet")
        
        # Header
        header_frame = tk.Frame(price_sheet_frame, bg=self.colors['background'])
        header_frame.pack(fill='x', pady=(0, 20))
        tk.Label(header_frame, text="Price Sheet", font=("Arial", 24, "bold"), 
                fg=self.colors['primary_blue'], bg=self.colors['background']).pack(side='left')
        
        # Cards grid container - similar to dashboard layout
        cards_wrap = tk.Frame(price_sheet_frame, bg=self.colors['background'])
        cards_wrap.pack(fill='both', expand=True, padx=20, pady=(5, 20))
        
        grid = tk.Frame(cards_wrap, bg=self.colors['background'])
        grid.pack(fill='both', expand=True)
        # Use 4 columns to match dashboard card size
        for i in range(4):
            grid.grid_columnconfigure(i, weight=1, uniform='col')
        
        # Define the 3 price sheet cards
        price_sheet_cards = [
            ("Vendor Price", "Vendor pricing information", "💲", self.open_vendor_price_module),
            ("Special Contractor Price", "Special contractor pricing", "👷", self.open_special_contractor_price),
            ("Client Allowables", "Client allowables pricing", "💰", self.open_client_allowables),
        ]
        
        # Create cards in a grid (using 4 columns to match dashboard card size)
        # Cards will be placed in columns 0, 1, 2 to maintain same size as dashboard cards
        row, col = 0, 0
        for title, subtitle, icon, cmd in price_sheet_cards:
            self._create_dashboard_card(grid, row, col, icon, title, subtitle, cmd)
            col += 1
            if col >= 4:
                col = 0
                row += 1
        
        # Add empty rows below to match dashboard layout (3 rows total)
        # This ensures the cards don't expand vertically like they would with just 1 row
        for empty_row in range(1, 3):
            grid.grid_rowconfigure(empty_row, weight=1)

    def open_vendor_price_module(self):
        """Open Vendor Price module in a new window"""
        new_window = tk.Toplevel(self.root)
        VendorPriceModule(new_window)

    def open_special_contractor_price(self):
        """Open Special Contractor Price module in a new window"""
        try:
            from special_contractor_price_module import SpecialContractorPriceModule
            new_window = tk.Toplevel(self.root)
            new_window.title("Special Contractor Prices")
            new_window.geometry("1200x800")
            new_window.configure(bg=self.colors['background'])
            
            # Center the window
            new_window.update_idletasks()
            x = (new_window.winfo_screenwidth() // 2) - (new_window.winfo_width() // 2)
            y = (new_window.winfo_screenheight() // 2) - (new_window.winfo_height() // 2)
            new_window.geometry(f"+{x}+{y}")
            
            # Create module in new window
            SpecialContractorPriceModule(
                new_window,
                username=self.username,
                user_id=self.user_id if hasattr(self, 'user_id') else None,
                colors=self.colors,
                db=self.db
            )
        except ImportError as e:
            messagebox.showerror("Error", f"Failed to import Special Contractor Price module: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open Special Contractor Price module: {e}")

    def open_client_allowables(self):
        """Open Client Allowables module (placeholder)"""
        messagebox.showinfo("Coming Soon", "Client Allowables module is coming soon.")

    def open_vendor_price(self):
        new_window = tk.Toplevel(self.root)
        VendorPriceModule(new_window)

    def open_gc_roof_ce(self):
        new_window = tk.Toplevel(self.root)
        GCRoofCEModule(new_window)

    def open_photo_viewer(self):
        if FastImageViewer is None:
            messagebox.showerror("Error", "Photo Viewer module not found.")
            return
        new_window = tk.Toplevel(self.root)
        FastImageViewer(new_window)

    def open_approval(self):
        self.clear_content_frame()
        approval_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        approval_frame.pack(fill="both", expand=True)
        
        # Back to Dashboard bar
        self._add_back_bar(approval_frame, "Approval Tracking")
        
        # Create ApprovalModule in the frame
        ApprovalModule(approval_frame, username=self.username, user_id=self.user_id if hasattr(self, 'user_id') else None)
    
    def open_wo_inspection(self):
        self.clear_content_frame()
        wo_inspection_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'], padx=20, pady=20)
        wo_inspection_frame.pack(fill="both", expand=True)
        
        # Back to Dashboard bar
        self._add_back_bar(wo_inspection_frame, "WO Inspection")
        
        # Create WOInspectionModule in the frame
        WOInspectionModule(wo_inspection_frame, username=self.username, user_id=self.user_id if hasattr(self, 'user_id') else None)

    def open_admin(self):
        """Open admin module"""
        if not self.db or not self.user_id:
            messagebox.showerror("Error", "Database connection not available.")
            return
        
        if not self.db.is_admin(self.user_id):
            messagebox.showerror("Access Denied", "You do not have admin privileges.")
            return
        
        self.clear_content_frame()
        container = tk.Frame(self.main_content_frame, bg=self.colors['background'])
        container.pack(fill='both', expand=True)
        
        # Back bar
        self._add_back_bar(container, "Admin")
        
        # Admin module
        try:
            from admin_module import AdminModule
            AdminModule(container, self.db, self.user_id, self.username, self.colors)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open admin module: {e}")
            import traceback
            traceback.print_exc()

    def show_settings(self):
        """Show the settings page with theme options."""
        self.clear_content_frame()
        self.main_content_frame.config(bg=self.colors['background'])
        self._add_back_bar(self.main_content_frame, "Settings")
        
        # Settings header
        header_frame = tk.Frame(self.main_content_frame, bg=self.colors['background'])
        header_frame.pack(pady=20, fill='x', padx=40)
        
        tk.Label(header_frame, text="Settings", 
                font=("Arial", 24, "bold"), 
                fg=self.colors['primary_blue'], 
                bg=self.colors['background']).pack(anchor='w')
        
        # Settings content container
        settings_container = tk.Frame(self.main_content_frame, bg=self.colors['background'])
        settings_container.pack(fill='both', expand=True, padx=40, pady=20)
        
        # Theme settings section
        theme_section = tk.Frame(settings_container, bg=self.colors['white'], relief="solid", bd=1)
        theme_section.pack(fill='x', pady=(0, 20))
        
        # Theme section header
        theme_header = tk.Frame(theme_section, bg=self.colors['primary_blue'], height=50)
        theme_header.pack(fill='x')
        theme_header.pack_propagate(False)
        
        tk.Label(theme_header, text="Appearance", 
                font=("Arial", 16, "bold"), 
                fg=self.colors['button_text'], 
                bg=self.colors['primary_blue']).pack(side='left', padx=20, pady=15)
        
        # Theme content
        theme_content = tk.Frame(theme_section, bg=self.colors['white'])
        theme_content.pack(fill='x', padx=30, pady=20)
        
        tk.Label(theme_content, text="Theme:", 
                font=("Arial", 12, "bold"), 
                fg=self.colors['text_primary'], 
                bg=self.colors['white']).pack(anchor='w', pady=(0, 10))
        
        # Theme selection frame
        theme_selection_frame = tk.Frame(theme_content, bg=self.colors['white'])
        theme_selection_frame.pack(fill='x')
        
        # Radio buttons for theme selection
        self.theme_var = tk.StringVar(value=theme_manager.current_theme)
        
        light_frame = tk.Frame(theme_selection_frame, bg=self.colors['white'])
        light_frame.pack(fill='x', pady=5)
        
        light_radio = tk.Radiobutton(light_frame, text="Light Mode", 
                                    variable=self.theme_var, value="light",
                                    font=("Arial", 11), 
                                    bg=self.colors['white'], 
                                    fg=self.colors['text_primary'],
                                    selectcolor=self.colors['white'],
                                    activebackground=self.colors['white'],
                                    command=lambda: self.change_theme("light"))
        light_radio.pack(side='left')
        
        tk.Label(light_frame, text="Perfect for bright environments and daytime use", 
                font=("Arial", 9), 
                fg=self.colors['text_secondary'], 
                bg=self.colors['white']).pack(side='left', padx=(10, 0))
        
        dark_frame = tk.Frame(theme_selection_frame, bg=self.colors['white'])
        dark_frame.pack(fill='x', pady=5)
        
        dark_radio = tk.Radiobutton(dark_frame, text="Dark Mode", 
                                   variable=self.theme_var, value="dark",
                                   font=("Arial", 11), 
                                   bg=self.colors['white'], 
                                   fg=self.colors['text_primary'],
                                   selectcolor=self.colors['white'],
                                   activebackground=self.colors['white'],
                                   command=lambda: self.change_theme("dark"))
        dark_radio.pack(side='left')
        
        tk.Label(dark_frame, text="Easier on the eyes for low-light environments", 
                font=("Arial", 9), 
                fg=self.colors['text_secondary'], 
                bg=self.colors['white']).pack(side='left', padx=(10, 0))
        
        # Additional settings sections can be added here
        # For example: Language, Notifications, etc.
        
        # Info section
        info_section = tk.Frame(settings_container, bg=self.colors['white'], relief="solid", bd=1)
        info_section.pack(fill='x', pady=(0, 20))
        
        info_header = tk.Frame(info_section, bg=self.colors['primary_blue'], height=50)
        info_header.pack(fill='x')
        info_header.pack_propagate(False)
        
        tk.Label(info_header, text="Information", 
                font=("Arial", 16, "bold"), 
                fg=self.colors['button_text'], 
                bg=self.colors['primary_blue']).pack(side='left', padx=20, pady=15)
        
        info_content = tk.Frame(info_section, bg=self.colors['white'])
        info_content.pack(fill='x', padx=30, pady=20)
        
        tk.Label(info_content, text="Techvengers Universal App", 
                font=("Arial", 12, "bold"), 
                fg=self.colors['text_primary'], 
                bg=self.colors['white']).pack(anchor='w')
        
        tk.Label(info_content, text="Version 1.0", 
                font=("Arial", 10), 
                fg=self.colors['text_secondary'], 
                bg=self.colors['white']).pack(anchor='w', pady=(5, 10))
        
        tk.Label(info_content, text="© 2024 Techvengers LTD. All rights reserved.", 
                font=("Arial", 9), 
                fg=self.colors['text_secondary'], 
                bg=self.colors['white']).pack(anchor='w')
        
        # Strategix Developer Section
        # Use a light blue background that complements the theme
        # Light theme: light blue, Dark theme: darker blue-gray
        if self.colors.get('background', '').startswith('#f') or self.colors.get('background') == '#f8f9fa':
            light_blue_bg = '#e8f0f8'  # Light blue for light theme
        else:
            light_blue_bg = '#2a3a4f'  # Darker blue-gray for dark theme
        strategix_section = tk.Frame(settings_container, bg=light_blue_bg, relief="solid", bd=1)
        strategix_section.pack(fill='x', pady=(0, 20))
        
        strategix_content = tk.Frame(strategix_section, bg=light_blue_bg)
        strategix_content.pack(fill='x', padx=30, pady=20)
        
        # Create clickable container for logo and text
        def open_strategix_url(event=None):
            webbrowser.open("https://strategixdigi.com/")
        
        strategix_container = tk.Frame(strategix_content, bg=light_blue_bg, cursor='hand2')
        strategix_container.pack(anchor='center', pady=10)
        
        # Try to load logo image if available
        logo_path = os.path.join(os.path.dirname(__file__), "strategix_logo.png")
        logo_displayed = False
        
        if PIL_AVAILABLE and os.path.exists(logo_path):
            try:
                # Load and resize logo (original size: 856x114 px)
                logo_image = Image.open(logo_path)
                # Resize to appropriate size (max width 400px, max height 60px, maintain aspect ratio)
                # The logo is wide, so we'll use a wider constraint
                logo_image.thumbnail((400, 60), Image.Resampling.LANCZOS)
                logo_photo = ImageTk.PhotoImage(logo_image)
                
                logo_label = tk.Label(strategix_container, image=logo_photo, bg=light_blue_bg, cursor='hand2')
                logo_label.image = logo_photo  # Keep a reference
                logo_label.pack(anchor='center', pady=(0, 10))
                logo_label.bind('<Button-1>', open_strategix_url)
                strategix_container.bind('<Button-1>', open_strategix_url)
                logo_displayed = True
            except Exception as e:
                print(f"Could not load Strategix logo: {e}")
        
        # If logo not displayed, create a simple colored representation
        if not logo_displayed:
            # Create a simple logo representation with colored frames
            logo_frame = tk.Frame(strategix_container, bg=light_blue_bg, cursor='hand2')
            logo_frame.pack(anchor='center', pady=(0, 10))
            
            # Left side - abstract symbol representation (green shapes)
            symbol_frame = tk.Frame(logo_frame, bg=light_blue_bg, cursor='hand2')
            symbol_frame.pack(side='left', padx=(0, 20))
            
            # Create simple representation with colored rectangles (approximation of logo)
            dark_green = "#2d5016"  # Dark green
            light_green = "#7cb342"  # Light green
            
            # Bottom larger shape (dark green)
            shape1 = tk.Frame(symbol_frame, bg=dark_green, width=40, height=25, cursor='hand2')
            shape1.pack(side='left', padx=2)
            shape1.pack_propagate(False)
            
            # Top overlapping shape (light green)
            shape2 = tk.Frame(symbol_frame, bg=light_green, width=35, height=20, cursor='hand2')
            shape2.pack(side='left', padx=2)
            shape2.pack_propagate(False)
            
            # Right side - X letter (light green)
            x_label = tk.Label(logo_frame, text="X", font=("Arial", 32, "bold"), 
                              fg=light_green, bg=light_blue_bg, cursor='hand2')
            x_label.pack(side='left')
            
            # Bind all logo elements
            for widget in [symbol_frame, shape1, shape2, x_label]:
                widget.bind('<Button-1>', open_strategix_url)
            logo_frame.bind('<Button-1>', open_strategix_url)
        
        # "Developed BY Strategix" text
        developed_label = tk.Label(strategix_container, text="Developed BY Strategix", 
                                  font=("Arial", 12, "bold"), 
                                  fg=self.colors['text_primary'], 
                                  bg=light_blue_bg,
                                  cursor='hand2')
        developed_label.pack(anchor='center', pady=(5, 0))
        developed_label.bind('<Button-1>', open_strategix_url)
        
        # Hover effects for the container
        def on_enter(widget):
            # Slightly darker blue on hover
            hover_bg = '#d0e0f0' if light_blue_bg == '#e8f0f8' else '#3a4a5f'
            widget.configure(bg=hover_bg)
        def on_leave(widget):
            widget.configure(bg=light_blue_bg)
        
        # Update hover effects to change background of all child widgets
        def on_enter_all(event=None):
            hover_bg = '#d0e0f0' if light_blue_bg == '#e8f0f8' else '#3a4a5f'
            strategix_container.configure(bg=hover_bg)
            strategix_content.configure(bg=hover_bg)
            strategix_section.configure(bg=hover_bg)
            if logo_displayed and 'logo_label' in locals():
                logo_label.configure(bg=hover_bg)
            if 'developed_label' in locals():
                developed_label.configure(bg=hover_bg)
        
        def on_leave_all(event=None):
            strategix_container.configure(bg=light_blue_bg)
            strategix_content.configure(bg=light_blue_bg)
            strategix_section.configure(bg=light_blue_bg)
            if logo_displayed and 'logo_label' in locals():
                logo_label.configure(bg=light_blue_bg)
            if 'developed_label' in locals():
                developed_label.configure(bg=light_blue_bg)
        
        strategix_container.bind('<Enter>', on_enter_all)
        strategix_container.bind('<Leave>', on_leave_all)
        strategix_content.bind('<Enter>', on_enter_all)
        strategix_content.bind('<Leave>', on_leave_all)

    # --- Helper: Back bar for in-page sections ---
    def _add_back_bar(self, parent, title_text=""):
        bar = tk.Frame(parent, bg=self.colors['background'])
        bar.pack(fill='x', padx=10, pady=(0, 10))
        back_btn = tk.Button(bar, text="← Back", font=("Segoe UI", 10),
                             bg=self.colors['primary_blue'], fg=self.colors['button_text'],
                             relief='flat', cursor='hand2', padx=10, pady=4,
                             command=self.show_dashboard_content)
        back_btn.pack(side='left')
        back_btn.bind('<Enter>', lambda e: back_btn.configure(bg=self.colors.get('nav_hover', self.colors['primary_blue'])))
        back_btn.bind('<Leave>', lambda e: back_btn.configure(bg=self.colors['primary_blue']))
        if title_text:
            tk.Label(bar, text=title_text, font=("Segoe UI", 12, "bold"),
                     bg=self.colors['background'], fg=self.colors['primary_blue']).pack(side='left', padx=10)
        return bar

    def change_theme(self, theme_name):
        """Change the application theme."""
        theme_manager.switch_theme(theme_name)
        messagebox.showinfo("Theme Changed", 
                           f"Theme changed to {theme_name.title()} Mode.\n"
                           f"The new theme has been applied to all open windows.")