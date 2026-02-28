import tkinter as tk
from tkinter import messagebox, filedialog
import os
import time
import sys
from datetime import datetime
import tempfile
import csv
import io
import requests
import json
import threading
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from utils import set_cell_background
from theme_manager import theme_manager
import re

try:
    from PIL import Image, ImageTk, ImageGrab
except ImportError:
    messagebox.showwarning("Missing Module", "PIL/Pillow is required for photo support. Install with: pip install pillow")
    Image = None
    ImageTk = None
    ImageGrab = None

class BidWriterApp:
    def __init__(self, root, username, wo_number_to_load=None, property_address=None, user_id=None, client_code=None, wo_type=None, on_save_callback=None):
        self.root = root
        self.root.title("Preservation Universe - Bid Writer")
        self.property_address = property_address or ""
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.root.geometry(f"{int(screen_width * 0.9)}x{int(screen_height * 0.9)}")
        
        # Use global theme manager
        self.colors = theme_manager.get_current_colors()
        
        # Register for theme updates
        theme_manager.register_theme_callback(self.on_theme_changed)
        
        self.root.configure(bg=self.colors['background'])
        
        # Store username and user_id
        self.username = username
        self.user_id = user_id
        try:
            from database_online import OnlineDatabaseManager
            self.db = OnlineDatabaseManager()
            # If user_id not provided, try to get it from database
            if not self.user_id:
                user = self.db.get_user(username)
                if user:
                    self.user_id = user['id']
        except Exception as e:
            print(f"Warning: Could not initialize database: {e}")
            self.db = None

        self.title_frame = tk.Frame(self.root, bg=self.colors['primary_blue'], height=60)
        self.title_frame.pack(fill='x', pady=(0, 3))
        self.title_frame.pack_propagate(False)
        
        self.username_label = tk.Label(self.title_frame, text=f"{username}", 
                                       font=("Arial", 12, "bold"), fg='white', 
                                       bg=self.colors['primary_blue'])
        self.username_label.pack(side="left", padx=12)
        
        # Search frame in the center
        self.search_frame = tk.Frame(self.title_frame, bg=self.colors['primary_blue'])
        self.search_frame.pack(side="left", expand=True, padx=(12, 0))
        
        # Search functionality
        search_container = tk.Frame(self.search_frame, bg=self.colors['white'], relief="solid", bd=1)
        search_container.pack(anchor="center", pady=4)
        
        self.search_entry = tk.Entry(search_container, font=("Arial", 11), relief="flat", bd=0, 
                                    width=30, bg=self.colors['white'], fg=self.colors['text_primary'])
        self.search_entry.pack(side="left", padx=(8, 5), pady=6)
        self.search_entry.bind("<Return>", lambda e: self.perform_search_with_navigation())
        self.search_entry.bind("<KeyRelease>", lambda e: self.perform_search())
        
        search_button = tk.Button(search_container, text="🔍", font=("Arial", 12), 
                                 bg=self.colors['light_blue'], fg="white", relief="flat", 
                                 cursor="hand2", command=self.perform_search)
        search_button.pack(side="right", padx=(0, 5), pady=2)
        
        # Clear search button
        clear_button = tk.Button(search_container, text="✕", font=("Arial", 10), 
                                bg=self.colors['gray_medium'], fg="white", relief="flat", 
                                cursor="hand2", command=self.clear_search)
        clear_button.pack(side="right", padx=(0, 5), pady=2)

        self.refresh_button = tk.Button(self.title_frame, text="Refresh",
                                       font=("Arial", 10, "bold"), bg=self.colors['light_blue'], 
                                       fg="white", relief="flat", cursor="hand2",
                                       activebackground=self.colors['primary_blue'],
                                       command=self.refresh_bids)
        self.refresh_button.pack(side="right", padx=(0, 12))

        self.categories = {}
        self.all_items = {}
        self.bid_data_url = "https://docs.google.com/spreadsheets/d/1sBPUtZqtoPREX2STfjBIs_kNF4HE4kCvsyloL9oC-tY/gviz/tq?tqx=out:csv&sheet=Sheet1"
        
        # Initialize Custom category as empty (will always appear first)
        self.custom_category_name = "Custom"
        self.all_items[self.custom_category_name] = []
        
        self.selected_items = {}
        self.item_photos = {}
        self.item_instances = {}
        self.current_photo_item = None
        self.active_category_button = None
        self.active_category = None
        self.search_highlights = []  # Store highlighted widgets for clearing
        self.current_search_results = []  # Store current search matches for navigation
        
        # Performance optimization: Debouncing and caching for preview updates
        self._preview_update_jobs = {}  # Track pending update jobs per item
        self._preview_cache = {}  # Cache formatted preview text to avoid redundant updates
        self._debounce_delay = 300  # milliseconds to wait after last keystroke
        self._updating_all_previews = False  # Flag to prevent cascading updates

        self.category_frame = tk.Frame(self.root, bg=self.colors['background'])
        self.category_frame.pack(pady=4, anchor="w", padx=12, fill="x")

        self.main_content_frame = tk.Frame(self.root, bg=self.colors['background'])
        self.main_content_frame.pack(pady=4, fill="both", expand=True, padx=12)

        self.canvas = tk.Canvas(self.main_content_frame, bg=self.colors['white'], 
                                highlightthickness=1, highlightcolor=self.colors['gray_light'])
        
        self.v_scrollbar = tk.Scrollbar(self.main_content_frame, orient="vertical", 
                                        command=self.canvas.yview, bg=self.colors['gray_light'])
        
        self.scrollable_frame = tk.Frame(self.canvas, bg=self.colors['white'])

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.v_scrollbar.set)

        def _on_mousewheel(event):
            """Handle mouse wheel scrolling with cross-platform compatibility."""
            try:
                # Windows and newer Mac versions
                if hasattr(event, 'delta') and event.delta:
                    delta = int(-1 * (event.delta / 120))
                # Older Mac versions or other platforms
                else:
                    # For platforms that don't have delta, use default scroll amount
                    delta = -1 if event.num == 4 else 1
                
                # Scroll the canvas
                self.canvas.yview_scroll(delta, "units")
            except (AttributeError, TypeError):
                # Fallback for any unexpected event format
                pass
        
        def _on_mousewheel_up(event):
            """Handle mouse wheel scroll up (Linux)."""
            self.canvas.yview_scroll(-1, "units")
            
        def _on_mousewheel_down(event):
            """Handle mouse wheel scroll down (Linux)."""
            self.canvas.yview_scroll(1, "units")
        
        # Create a comprehensive mouse wheel binding helper
        def bind_mousewheel_to_widget(widget):
            """Bind mouse wheel events to a widget for cross-platform scrolling."""
            try:
                # Windows and Mac
                widget.bind("<MouseWheel>", _on_mousewheel)
                # Linux scroll up
                widget.bind("<Button-4>", _on_mousewheel_up)
                # Linux scroll down  
                widget.bind("<Button-5>", _on_mousewheel_down)
                # Mac trackpad horizontal scroll (bind but ignore for now)
                widget.bind("<Shift-MouseWheel>", lambda e: None)
            except Exception:
                # If binding fails for any reason, continue silently
                pass
        
        # Store the binding function for later use
        self.bind_mousewheel_to_widget = bind_mousewheel_to_widget
        
        # Bind to canvas and scrollable frame
        bind_mousewheel_to_widget(self.canvas)
        bind_mousewheel_to_widget(self.scrollable_frame)

        self.v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        self.wo_frame = tk.Frame(self.root, bg=self.colors['background'])
        self.wo_frame.pack(pady=(4, 0), padx=12, fill='x')

        self.wo_label = tk.Label(self.wo_frame, text="WO: ", 
                                 font=("Arial", 11, "bold"), bg=self.colors['background'], 
                                 fg=self.colors['primary_blue'])
        self.wo_label.pack(side="left", padx=(0, 5))
        self.wo_entry = tk.Entry(self.wo_frame, font=("Arial", 11), relief="solid", bd=1, width=20, bg=self.colors['gray_light'], fg=self.colors['text_primary'], state='readonly')
        self.wo_entry.pack(side="left", padx=(0, 20))
        if wo_number_to_load:
            self.wo_entry.config(state='normal')
            self.wo_entry.insert(0, wo_number_to_load)
            self.wo_entry.config(state='readonly')
        
        # Property Address field
        self.address_label = tk.Label(self.wo_frame, text="Property Address: ", 
                                      font=("Arial", 11, "bold"), bg=self.colors['background'], 
                                      fg=self.colors['primary_blue'])
        self.address_label.pack(side="left", padx=(0, 5))
        self.address_entry = tk.Entry(self.wo_frame, font=("Arial", 11), relief="solid", bd=1, width=40, bg=self.colors['gray_light'], fg=self.colors['text_primary'], state='readonly')
        self.address_entry.pack(side="left", padx=(0, 20))
        if self.property_address:
            self.address_entry.insert(0, self.property_address)
        
        # Client Code field
        self.client_code_label = tk.Label(self.wo_frame, text="Client Code: ", 
                                         font=("Arial", 11, "bold"), bg=self.colors['background'], 
                                         fg=self.colors['primary_blue'])
        self.client_code_label.pack(side="left", padx=(0, 5))
        self.client_code_entry = tk.Entry(self.wo_frame, font=("Arial", 11), relief="solid", bd=1, width=15, bg=self.colors['gray_light'], fg=self.colors['text_primary'], state='readonly')
        self.client_code_entry.pack(side="left", padx=(0, 20))
        if client_code:
            self.client_code_entry.insert(0, client_code)
        
        # WO Type field
        self.wo_type_label = tk.Label(self.wo_frame, text="WO Type: ", 
                                     font=("Arial", 11, "bold"), bg=self.colors['background'], 
                                     fg=self.colors['primary_blue'])
        self.wo_type_label.pack(side="left", padx=(0, 5))
        self.wo_type_entry = tk.Entry(self.wo_frame, font=("Arial", 11), relief="solid", bd=1, width=20, bg=self.colors['gray_light'], fg=self.colors['text_primary'], state='readonly')
        self.wo_type_entry.pack(side="left", padx=(0, 20))
        if wo_type:
            self.wo_type_entry.insert(0, wo_type)

        self.save_state_button = tk.Button(self.wo_frame, text="Save Project", command=self.save_state,
                                           font=("Arial", 10, "bold"), bg=self.colors['light_blue'], fg="white",
                                           relief="flat", cursor="hand2")
        self.save_state_button.pack(side="left", padx=(0, 10))


        self.buttons_container = tk.Frame(self.root, bg=self.colors['background'])
        self.buttons_container.pack(pady=4)

        self.generate_button = tk.Button(self.buttons_container, text="Generate Bids", command=self.generate_bids,
                                         font=("Arial", 12, "bold"), bg=self.colors['green'], 
                                         fg="white", height=1, width=20, relief="solid", bd=1, cursor="hand2")
        self.generate_button.pack(side="left", padx=(0, 10))

        self.clear_button = tk.Button(self.buttons_container, text="Clear Bids", command=self.clear_bids,
                                         font=("Arial", 12, "bold"), bg="#dc3545",
                                         fg="white", height=1, width=15, relief="solid", bd=1, cursor="hand2")
        self.clear_button.pack(side="left", padx=(0, 10))
        
        self.docs1_button = tk.Button(self.buttons_container, text="Docs1", command=self.save_to_docs1,
                                         font=("Arial", 12, "bold"), bg=self.colors['primary_blue'], 
                                         fg="white", height=1, width=12, relief="solid", bd=1, cursor="hand2")
        self.docs1_button.pack(side="left", padx=(0, 10))
        
        self.docs2_button = tk.Button(self.buttons_container, text="Docs2", command=self.save_to_docs2,
                                         font=("Arial", 12, "bold"), bg=self.colors['primary_blue'], 
                                         fg="white", height=1, width=12, relief="solid", bd=1, cursor="hand2")
        self.docs2_button.pack(side="left", padx=(0, 10))

        self.output_frame = tk.Frame(self.root, bg=self.colors['gray_light'], relief="solid", bd=1)
        self.output_frame.pack(padx=12, pady=(0, 4), fill='both', expand=False)
        
        self.output_header_frame = tk.Frame(self.output_frame, bg=self.colors['gray_light'])
        self.output_header_frame.pack(fill='x', padx=2, pady=(2, 0))

        self.bid_count_label = tk.Label(self.output_header_frame, text="Total Bids: 0",
                                         font=("Arial", 11, "bold"), bg=self.colors['gray_light'],
                                         fg=self.colors['gray_dark'], anchor="w")
        self.bid_count_label.pack(side="left", padx=(8, 0))
        
        # Add icon button to open bids in full page view
        self.view_full_page_button = tk.Button(self.output_header_frame, text="⛶",
                                               font=("Arial", 14), bg=self.colors['light_blue'],
                                               fg="white", relief="flat", cursor="hand2",
                                               activebackground=self.colors['primary_blue'],
                                               command=self.open_full_page_view, width=3, height=1)
        self.view_full_page_button.pack(side="left", padx=(10, 0))
        
        self.output_scrollbar = tk.Scrollbar(self.output_frame)
        self.output_scrollbar.pack(side="right", fill="y")

        self.output_text = tk.Text(self.output_frame, font=("Arial", 11), bg=self.colors['white'], 
                                     wrap=tk.WORD, relief="flat", padx=10, pady=10, height=8,
                                     yscrollcommand=self.output_scrollbar.set)
        self.output_text.pack(fill='both', expand=True, padx=(2, 0), pady=2)
        self.output_scrollbar.config(command=self.output_text.yview)

        self.output_text.images = []

        self.app_data_dir = os.path.join(os.path.expanduser("~"), ".techvengers_bidwriter")
        os.makedirs(self.app_data_dir, exist_ok=True)
        self.root.bind('<Control-v>', self.handle_global_paste)
        self.root.bind('<Control-s>', self.focus_search_bar)

        self.on_save_callback = on_save_callback

        # Load bids asynchronously to prevent blocking UI
        self.load_bids_from_url_async(self.bid_data_url)
        if wo_number_to_load:
            # If loading existing bid, try to load all project fields from database (from any user)
            if self.db:
                try:
                    bid_data = self.db.load_bid(wo_number_to_load, self.user_id, all_users=True)
                    if bid_data:
                        # Load property_address
                        if bid_data.get('property_address'):
                            self.property_address = bid_data['property_address']
                            self.address_entry.config(state='normal')
                            self.address_entry.delete(0, tk.END)
                            self.address_entry.insert(0, self.property_address)
                            self.address_entry.config(state='readonly')
                        
                        # Load client_code
                        if bid_data.get('client_code') and hasattr(self, 'client_code_entry'):
                            self.client_code_entry.config(state='normal')
                            self.client_code_entry.delete(0, tk.END)
                            self.client_code_entry.insert(0, bid_data['client_code'])
                            self.client_code_entry.config(state='readonly')
                        
                        # Load wo_type
                        if bid_data.get('wo_type') and hasattr(self, 'wo_type_entry'):
                            self.wo_type_entry.config(state='normal')
                            self.wo_type_entry.delete(0, tk.END)
                            self.wo_type_entry.insert(0, bid_data['wo_type'])
                            self.wo_type_entry.config(state='readonly')
                        
                        # Load the bid state (selected items, photos, etc.)
                        self.load_state()
                except Exception as e:
                    print(f"Could not load project data: {e}")
                    # Still try to load state if possible
                    if wo_number_to_load:
                        self.load_state()

        self.root.after(300000, self.auto_save)

    def on_theme_changed(self, theme_name, colors):
        """Called when theme is changed globally."""
        self.colors = colors
        self.apply_theme_to_all_widgets()

    def perform_search(self):
        """Perform search and highlight matching categories and items across all categories."""
        search_term = self.search_entry.get().strip().lower()
        self.clear_highlights()
        
        if not search_term:
            return
        
        # Search and highlight categories
        for widget in self.category_frame.winfo_children():
            if isinstance(widget, tk.Button):
                category_text = widget.cget("text").lower()
                if search_term in category_text:
                    original_bg = widget.cget("bg")
                    widget.configure(bg="#FFD700")  # Yellow highlight
                    self.search_highlights.append((widget, original_bg))
        
        # Search and highlight items across ALL categories
        self.highlight_matching_items_all_categories(search_term)

    def highlight_matching_items_all_categories(self, search_term):
        """Highlight matching items across all categories and highlight their category buttons."""
        # Track which categories have matching items
        categories_with_matches = set()
        
        # First, search through the complete all_items dataset to find all matching categories
        if hasattr(self, 'all_items'):
            for category, items_list in self.all_items.items():
                for item_data in items_list:
                    item_name = item_data.get('item_name', '').lower()
                    template_text = item_data.get('template', '').lower()
                    
                    if (search_term in item_name or 
                        search_term in template_text):
                        categories_with_matches.add(category)
        
        # Then, highlight any currently displayed items that match
        if hasattr(self, 'selected_items'):
            for category, category_items in self.selected_items.items():
                for item_key, item_info in category_items.items():
                    if item_info.get("button"):
                        # Check if item name matches search term
                        item_name = item_info.get("original_name", "").lower()
                        template_text = item_info.get("template", "").lower()
                        
                        if (search_term in item_name or 
                            search_term in template_text):
                            
                            # Highlight the item button if it's currently displayed
                            button = item_info["button"]
                            if button and button.winfo_exists():
                                original_bg = button.cget("bg")
                                button.configure(bg="#FFD700")  # Yellow highlight
                                self.search_highlights.append((button, original_bg))
        
        # Highlight category buttons that have matching items (if not already highlighted)
        for widget in self.category_frame.winfo_children():
            if isinstance(widget, tk.Button):
                category_text = widget.cget("text")
                if (category_text in categories_with_matches and 
                    widget.cget("bg") != "#FFD700"):  # Don't re-highlight if already highlighted
                    original_bg = widget.cget("bg")
                    widget.configure(bg="#FFD700")  # Yellow highlight
                    self.search_highlights.append((widget, original_bg))

    def highlight_matching_items(self, search_term):
        """Highlight matching items in the current grid."""
        if not hasattr(self, 'selected_items') or not self.active_category:
            return
            
        category = self.active_category
        if category not in self.selected_items:
            return
            
        for item_key, item_info in self.selected_items[category].items():
            if item_info.get("button"):
                # Check if item name matches search term
                item_name = item_info.get("original_name", "").lower()
                template_text = item_info.get("template", "").lower()
                
                if (search_term in item_name or 
                    search_term in template_text):
                    
                    button = item_info["button"]
                    original_bg = button.cget("bg")
                    button.configure(bg="#FFD700")  # Yellow highlight
                    self.search_highlights.append((button, original_bg))

    def clear_search(self):
        """Clear search entry and highlights."""
        self.search_entry.delete(0, tk.END)
        self.clear_highlights()

    def clear_highlights(self):
        """Clear all search highlights."""
        for widget, original_bg in self.search_highlights:
            try:
                if widget.winfo_exists():
                    widget.configure(bg=original_bg)
            except:
                pass
        self.search_highlights.clear()

    def focus_search_bar(self, event=None):
        """Focus the search bar when CTRL+S is pressed."""
        self.search_entry.focus_set()
        self.search_entry.select_range(0, tk.END)
        return "break"  # Prevent default save dialog

    def perform_search_with_navigation(self):
        """Perform search with highlighting and navigate to first match."""
        self.perform_search()
        self.navigate_to_first_match()

    def navigate_to_first_match(self):
        """Navigate to the first search match found."""
        search_term = self.search_entry.get().strip().lower()
        if not search_term:
            return

        # First, check if any category buttons match
        for widget in self.category_frame.winfo_children():
            if isinstance(widget, tk.Button):
                category_text = widget.cget("text").lower()
                if search_term in category_text:
                    # Click on the first matching category to load its items
                    widget.invoke()
                    return

        # If no category match, look for items and switch to their category
        first_match_category = self.find_first_matching_item_category(search_term)
        if first_match_category:
            # Switch to the category containing the first match
            self.switch_to_category_with_match(first_match_category)
            # Scroll to make the match visible (with a small delay to allow UI to update)
            self.root.after(100, lambda: self.scroll_to_first_item_match(search_term))

    def find_first_matching_item_category(self, search_term):
        """Find the first category that contains a matching item."""
        if not hasattr(self, 'all_items'):
            return None
            
        # Search through all categories to find the first match
        for category, items_list in self.all_items.items():
            for item_data in items_list:
                item_name = item_data.get('item_name', '').lower()
                template_text = item_data.get('template', '').lower()
                
                if search_term in item_name or search_term in template_text:
                    return category
        return None

    def switch_to_category_with_match(self, target_category):
        """Switch to the specified category."""
        for widget in self.category_frame.winfo_children():
            if isinstance(widget, tk.Button) and widget.cget("text") == target_category:
                widget.invoke()
                break

    def scroll_to_first_item_match(self, search_term):
        """Scroll the grid to make the first matching item visible."""
        if not hasattr(self, 'selected_items') or not self.active_category:
            return
            
        category = self.active_category
        if category not in self.selected_items:
            return
        
        # Find the first matching item in the current view
        first_match_row = None
        row_idx = 1  # Start after header row
        
        # Sort items to maintain consistent order
        sorted_items = sorted(self.selected_items[category].items(), 
                            key=lambda x: x[1]['instance_info']['key'])
        
        for item_key, item_info in sorted_items:
            item_name = item_info.get("original_name", "").lower()
            template_text = item_info.get("template", "").lower()
            
            if search_term in item_name or search_term in template_text:
                first_match_row = row_idx
                break
            row_idx += 1
        
        if first_match_row is not None:
            # Scroll the canvas to make the matching row visible
            self.scroll_to_row(first_match_row)

    def scroll_to_row(self, row_number):
        """Scroll the canvas to make the specified row visible."""
        try:
            # Update the scroll region first
            self.canvas.update_idletasks()
            
            # Calculate the approximate y position of the row
            # Assuming each row is about 40-50 pixels high (header + content + padding)
            row_height = 45
            target_y = row_number * row_height
            
            # Get the current scroll region
            scroll_region = self.canvas.cget("scrollregion")
            if scroll_region:
                # Parse the scroll region (format: "x1 y1 x2 y2")
                parts = scroll_region.split()
                if len(parts) == 4:
                    total_height = float(parts[3]) - float(parts[1])
                    if total_height > 0:
                        # Calculate the fraction to scroll to
                        scroll_fraction = min(target_y / total_height, 1.0)
                        # Scroll to position
                        self.canvas.yview_moveto(scroll_fraction)
        except Exception as e:
            # If scrolling fails, continue silently
            pass

    def apply_theme_to_all_widgets(self):
        """Apply current theme to all widgets in the application."""
        # Update root
        self.root.configure(bg=self.colors['background'])
        
        # Update main frames
        self.title_frame.configure(bg=self.colors['primary_blue'])
        self.category_frame.configure(bg=self.colors['background'])
        self.main_content_frame.configure(bg=self.colors['background'])
        self.wo_frame.configure(bg=self.colors['background'])
        self.buttons_container.configure(bg=self.colors['background'])
        self.output_frame.configure(bg=self.colors['gray_light'])
        self.output_header_frame.configure(bg=self.colors['gray_light'])
        
        # Update labels
        self.username_label.configure(bg=self.colors['primary_blue'], fg=self.colors['button_text'])
        self.wo_label.configure(bg=self.colors['background'], fg=self.colors['primary_blue'])
        self.bid_count_label.configure(bg=self.colors['gray_light'], fg=self.colors['gray_dark'])
        
        # Update search elements
        self.search_frame.configure(bg=self.colors['primary_blue'])
        self.search_entry.configure(bg=self.colors['white'], fg=self.colors['text_primary'])
        
        # Update buttons
        self.refresh_button.configure(bg=self.colors['light_blue'], fg=self.colors['button_text'], 
                                     activebackground=self.colors['primary_blue'])
        self.save_state_button.configure(bg=self.colors['light_blue'], fg=self.colors['button_text'])
        self.generate_button.configure(bg=self.colors['green'], fg=self.colors['button_text'])
        self.docs1_button.configure(bg=self.colors['primary_blue'], fg=self.colors['button_text'])
        self.docs2_button.configure(bg=self.colors['primary_blue'], fg=self.colors['button_text'])
        
        # Update text widgets
        self.output_text.configure(bg=self.colors['white'], fg=self.colors['text_primary'])
        self.wo_entry.configure(bg=self.colors['white'], fg=self.colors['text_primary'])
        
        # Update category buttons if they exist
        if hasattr(self, 'category_frame'):
            for widget in self.category_frame.winfo_children():
                if isinstance(widget, tk.Button):
                    if widget == self.active_category_button:
                        widget.configure(bg=self.colors['active_category_color'], fg=self.colors['button_text'])
                    else:
                        widget.configure(bg=self.colors['light_blue'], fg=self.colors['button_text'],
                                       activebackground=self.colors['primary_blue'])
        
        # Update grid if it exists
        self.update_grid_theme()

    def update_grid_theme(self):
        """Update the grid theme if it exists."""
        if hasattr(self, 'scrollable_frame') and self.scrollable_frame.winfo_children():
            # Re-load the current category to apply new theme
            if hasattr(self, 'active_category') and self.active_category:
                self.load_items(self.active_category)



    def auto_save(self):
        try:
            wo_number = self.wo_entry.get().strip()
            if not wo_number:
                # Keep a stable autosave name across this session
                if not hasattr(self, '_autosave_name'):
                    self._autosave_name = f"AutoSave_{int(time.time())}"
                wo_number = self._autosave_name
                self.wo_entry.delete(0, tk.END)
                self.wo_entry.insert(0, wo_number)
            self.save_state(silent=True)
        except Exception:
            # Avoid crashing autosave on transient widget errors
            pass
        finally:
            self.root.after(300000, self.auto_save)

    def update_bid_buttons(self):
        """Update category buttons and apply highlights based on search query."""
        if self.active_category_button:
            self.active_category_button.configure(bg=self.colors['light_blue'])
        
        for widget in self.category_frame.winfo_children():
            widget.destroy()
        
        self.active_category_button = None
        
        # Get list of categories, ensuring Custom is first
        categories = list(self.all_items.keys())
        if self.custom_category_name in categories:
            categories.remove(self.custom_category_name)
            categories.insert(0, self.custom_category_name)
        
        total_categories = len(categories)
        
        # Calculate buttons per row (half rounded up)
        buttons_per_row = (total_categories + 1) // 2
        
        # Create buttons and arrange them in 2 rows using grid
        for idx, category in enumerate(categories):
            bg_color = self.colors['light_blue']

            btn = tk.Button(self.category_frame, text=category, width=20, height=2, 
                             font=("Arial", 12, "bold"), bg=bg_color, 
                             fg='white', relief="flat", cursor="hand2",
                             activebackground=self.colors['primary_blue'], activeforeground='white')

            btn.configure(command=lambda c=category, b=btn: self.load_items_with_highlight(c, b))
            btn.bind("<Enter>", lambda e, b=btn: self.on_hover(b))
            btn.bind("<Leave>", lambda e, b=btn: self.on_leave_button(b))
            
            # Calculate row and column positions
            row = idx // buttons_per_row
            col = idx % buttons_per_row
            
            btn.grid(row=row, column=col, padx=2, pady=2, sticky="ew")
        
        # Configure grid columns to have equal weight
        for col in range(buttons_per_row):
            self.category_frame.grid_columnconfigure(col, weight=1, uniform="category_buttons")

    def on_hover(self, button):
        """Change button color on hover, unless it's the active button."""
        if self.active_category_button and button.cget("text") == self.active_category_button.cget("text"):
            pass
        else:
            button.configure(bg=self.colors['primary_blue'])

    def on_leave_button(self, button):
        """Restore button color on mouse leave, unless it's the active button."""
        if self.active_category_button and button.cget("text") == self.active_category_button.cget("text"):
            pass
        else:
            button.configure(bg=self.colors['light_blue'])
        
    def load_items_with_highlight(self, category, button):
        """Load items and highlight the selected button."""
        if self.active_category_button and self.active_category_button.winfo_exists():
            self.active_category_button.configure(bg=self.colors['light_blue'])
        
        self.active_category_button = button
        self.active_category_button.configure(bg=self.colors['active_category_color'])
        
        self.active_category = category
        self.clear_highlights()  # Clear previous highlights when switching categories
        self.load_items(category)
        
        # Re-apply search if there's a search term
        if hasattr(self, 'search_entry') and self.search_entry.get().strip():
            self.perform_search()

    def refresh_bids(self):
        """Refreshes the bids by reloading from the online URL."""
        self.load_bids_from_url(self.bid_data_url)
        messagebox.showinfo("Refresh Complete", "Bid list has been refreshed successfully.")

    def load_bids_from_url_async(self, url):
        """Loads bid data asynchronously to prevent blocking UI."""
        # Initialize empty categories first, but preserve Custom category
        self.categories = {}
        self.all_items = {}
        # Ensure Custom category is always present
        self.all_items[self.custom_category_name] = []
        
        # Show loading indicator
        loading_label = tk.Label(self.category_frame, text="Loading bids...", 
                                font=("Arial", 11), bg=self.colors['background'],
                                fg=self.colors['gray_dark'])
        loading_label.pack(side="left", padx=5)
        
        def load_in_thread():
            try:
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                
                csv_data = io.StringIO(response.text)
                reader = csv.DictReader(csv_data)
                
                categories = {}
                for row in reader:
                    if 'Category' in row and 'Item' in row and 'Template' in row:
                        category = row['Category']
                        item_name = row['Item']
                        template = row['Template']
                        unit_price = row.get('Unit Price', '0.00')
                        
                        if category not in categories:
                            categories[category] = []
                        categories[category].append({'item_name': item_name, 'template': template, 'unit_price': unit_price})
                
                all_items = categories.copy()
                
                # Update UI in main thread
                def update_ui():
                    loading_label.destroy()
                    self.categories = categories
                    self.all_items = all_items
                    # Ensure Custom category is always present
                    if self.custom_category_name not in self.all_items:
                        self.all_items[self.custom_category_name] = []
                    self.update_bid_buttons()
                    if self.categories:
                        self.active_category = list(self.categories.keys())[0]
                        if self.category_frame.winfo_children():
                            first_button = self.category_frame.winfo_children()[0]
                            self.load_items_with_highlight(self.active_category, first_button)
                
                self.root.after(0, update_ui)
                
            except requests.exceptions.RequestException as e:
                def show_error():
                    try:
                        if loading_label.winfo_exists():
                            loading_label.destroy()
                        messagebox.showwarning("Network Error", f"Could not connect to the online file. Using default bids.\nError: {e}")
                        self.load_default_bids()
                        self.update_bid_buttons()
                        if self.categories:
                            self.active_category = list(self.categories.keys())[0]
                            if self.category_frame.winfo_children():
                                first_button = self.category_frame.winfo_children()[0]
                                self.load_items_with_highlight(self.active_category, first_button)
                    except Exception as err:
                        print(f"Error in error handler: {err}")
                self.root.after(0, show_error)
            except Exception as e:
                def show_error():
                    try:
                        if loading_label.winfo_exists():
                            loading_label.destroy()
                        messagebox.showwarning("Error", f"Failed to read data from online file. Using default bids.\nError: {e}")
                        self.load_default_bids()
                        self.update_bid_buttons()
                        if self.categories:
                            self.active_category = list(self.categories.keys())[0]
                            if self.category_frame.winfo_children():
                                first_button = self.category_frame.winfo_children()[0]
                                self.load_items_with_highlight(self.active_category, first_button)
                    except Exception as err:
                        print(f"Error in error handler: {err}")
                self.root.after(0, show_error)
        
        # Start loading in background thread
        thread = threading.Thread(target=load_in_thread, daemon=True)
        thread.start()
    
    def load_bids_from_url(self, url):
        """Loads bid data from a public CSV file URL (synchronous version for refresh)."""
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            csv_data = io.StringIO(response.text)
            reader = csv.DictReader(csv_data)
            
            self.categories = {}
            for row in reader:
                if 'Category' in row and 'Item' in row and 'Template' in row:
                    category = row['Category']
                    item_name = row['Item']
                    template = row['Template']
                    unit_price = row.get('Unit Price', '0.00')
                    
                    if category not in self.categories:
                        self.categories[category] = []
                    self.categories[category].append({'item_name': item_name, 'template': template, 'unit_price': unit_price})
            
            self.all_items = self.categories.copy()
            # Ensure Custom category is always present
            if self.custom_category_name not in self.all_items:
                self.all_items[self.custom_category_name] = []
            
        except requests.exceptions.RequestException as e:
            messagebox.showwarning("Network Error", f"Could not connect to the online file. Using default bids.\nError: {e}")
            self.load_default_bids()
        except Exception as e:
            messagebox.showwarning("Error", f"Failed to read data from online file. Using default bids.\nError: {e}")
            self.load_default_bids()
        
        self.update_bid_buttons()
        if self.categories:
            self.active_category = list(self.categories.keys())[0]
            first_button = self.category_frame.winfo_children()[0]
            self.load_items_with_highlight(self.active_category, first_button)

    def load_default_bids(self):
        """Loads hardcoded default bids as a fallback."""
        self.categories = {
            "Landscaping": [
                {'item_name': "Trim Shrubs", 'template': "Trim {quantity} LF of shrubs from the {location} of the property. Price includes time, labor and disposal of generated debris. {info}\nPrice: ${total:.2f}", 'unit_price': '15.00'},
                {'item_name': "Trim Tree", 'template': "Trim {quantity} LF from the medium tree 30' tall and 18\" in diameter from {location} of the property. Price includes time, labor and disposal of generated debris. {info}\nPrice: ${total:.2f}", 'unit_price': '125.00'}
            ],
            "Mold": [
                {'item_name': "Antimicrobial", 'template': "Clean & wipe {quantity} SF area of moldy walls in the {location}. Price includes time, labor, required materials and the removal of the generated debris. Must be approved with the Kilz bid. {info}\nPrice: ${total:.2f}", 'unit_price': '0.40'},
                {'item_name': "Kilz", 'template': "Paint 1 coat of Kilz to {quantity} SF area of moldy walls in the {location}. Price includes time, labor, required materials and the removal of the generated debris. Source of the mold: roof leakage. Must be approved with the Kilz bid. {info}\nPrice: ${total:.2f}", 'unit_price': '0.50'}
            ]
        }
        self.all_items = self.categories.copy()
        # Ensure Custom category is always present
        if self.custom_category_name not in self.all_items:
            self.all_items[self.custom_category_name] = []
        
    def load_custom_items(self):
        """Load blank rows for Custom category that can be filled by user"""
        # Preserve preview text content before destroying widgets
        if self.custom_category_name in self.selected_items:
            for item_key, item_info in self.selected_items[self.custom_category_name].items():
                if item_info.get("preview_text"):
                    try:
                        if item_info["preview_text"].winfo_exists():
                            preview_content = item_info["preview_text"].get("1.0", tk.END).strip()
                            # Only preserve if user has actually edited it
                            if preview_content and item_info.get("user_edited", False):
                                item_info["preview_text_content"] = preview_content
                    except:
                        pass
        
        # Clear scrollable frame first
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        
        grid_frame = tk.Frame(self.scrollable_frame, bg=self.colors['white'])
        grid_frame.pack(fill="both", expand=True, padx=4, pady=4)
        self.bind_mousewheel_to_widget(grid_frame)
        
        # Column configurations
        col_configs = [
            (0, 40, 0),     # Key column
            (1, 40, 0),     # Add/Delete column
            (2, 120, 2),    # Item column
            (3, 50, 0),     # Qty column
            (4, 70, 0),     # Unit Price column
            (5, 70, 0),     # Total Price column
            (6, 120, 2),    # Location column
            (7, 150, 2),    # Additional Info column
            (8, 350, 4),    # Live Preview column
            (9, 150, 4)     # Photo column
        ]
        
        for col, min_width, weight in col_configs:
            grid_frame.grid_columnconfigure(col, minsize=min_width, weight=int(weight))
        
        # Headers
        headings = ["Key", "Add", "Item", "Qty", "Unit Price", "Total Price", "Location", "Additional Info", "Live Preview", "Photo"]
        for col, heading in enumerate(headings):
            header_frame = tk.Frame(grid_frame, bg=self.colors['primary_blue'], relief="flat", bd=1)
            header_frame.grid(row=0, column=col, sticky="nsew", padx=1, pady=1)
            label = tk.Label(header_frame, text=heading, font=("Arial", 11, "bold"), 
                             bg=self.colors['primary_blue'], fg='white', anchor="w")
            label.pack(fill="both", expand=True, padx=4, pady=4)
            self.bind_mousewheel_to_widget(header_frame)
            self.bind_mousewheel_to_widget(label)
        
        # Clear existing rows (but keep headers)
        for widget in grid_frame.grid_slaves():
            if int(widget.grid_info()["row"]) > 0:
                widget.destroy()
        
        # Initialize custom instances if not exists
        if self.custom_category_name not in self.item_instances:
            self.item_instances[self.custom_category_name] = {}
        if self.custom_category_name not in self.selected_items:
            self.selected_items[self.custom_category_name] = {}
        
        # Get existing custom instances from saved state or create one if empty
        custom_instances = []
        for item_name in self.item_instances[self.custom_category_name]:
            custom_instances.extend(self.item_instances[self.custom_category_name][item_name])
        
        # Also get instances from selected_items if they exist but aren't in item_instances yet
        for item_key, item_info in self.selected_items[self.custom_category_name].items():
            instance_info = item_info.get('instance_info', {})
            if instance_info:
                # Check if this instance is already in our list
                existing = any(inst.get('key') == instance_info.get('key') for inst in custom_instances)
                if not existing:
                    # Ensure the instance is in item_instances structure
                    item_name = item_info.get('original_name', instance_info.get('display_name', 'Custom Item'))
                    if item_name not in self.item_instances[self.custom_category_name]:
                        self.item_instances[self.custom_category_name][item_name] = []
                    if instance_info not in self.item_instances[self.custom_category_name][item_name]:
                        self.item_instances[self.custom_category_name][item_name].append(instance_info)
                        custom_instances.append(instance_info)
        
        # Initialize counter if needed
        if not hasattr(self, 'custom_item_counter'):
            self.custom_item_counter = 1
        
        # If no custom items exist, create one blank row
        if not custom_instances:
            instance_key = f"custom_item_{self.custom_item_counter}"
            item_name = "Custom Item 1"
            if item_name not in self.item_instances[self.custom_category_name]:
                self.item_instances[self.custom_category_name][item_name] = []
            new_instance = {
                'instance_id': 1,
                'display_name': item_name,
                'key': instance_key,
                'is_custom': True
            }
            self.item_instances[self.custom_category_name][item_name].append(new_instance)
            custom_instances = [new_instance]
            self.custom_item_counter += 1
        
        # Sort instances by key
        custom_instances.sort(key=lambda x: x['key'])
        
        # Create rows for each custom instance
        row_idx = 1
        for instance_info in custom_instances:
            instance_key = instance_info['key']
            item_name = instance_info.get('display_name', 'Custom Item')
            
            # Initialize item info if not exists
            if instance_key not in self.selected_items[self.custom_category_name]:
                self.selected_items[self.custom_category_name][instance_key] = {
                    "selected": False,
                    "template": "{description}",  # Simple template for custom items (no price)
                    "qty": tk.StringVar(value="0"),
                    "unit_price": tk.StringVar(value="0.00"),
                    "location": tk.StringVar(),
                    "add_info": tk.StringVar(),
                    "conjunction_key": tk.StringVar(),
                    "total_price_label": None,
                    "item_name_entry": None,  # Entry for custom item name
                    "checkbox": None,  # Checkbox for selection
                    "preview_text": None,
                    "preview_text_content": "",  # Store preview text content as string
                    "original_name": item_name,
                    "instance_info": instance_info,
                    "photo_frame": None,
                    "photo_label": None,
                    "user_edited": False,
                    "is_custom": True
                }
            else:
                # Preserve existing preview_text_content if it exists
                if "preview_text_content" not in item_info:
                    item_info["preview_text_content"] = ""
            
            item_info = self.selected_items[self.custom_category_name][instance_key]
            
            # Create row cells (similar to regular items but with editable Item field)
            self.create_custom_row(grid_frame, row_idx, item_info, instance_info)
            row_idx += 1
        
        # Reset scroll position
        self.reset_scroll_to_top()
    
    def create_custom_row(self, grid_frame, row_idx, item_info, instance_info):
        """Create a single custom row with editable fields"""
        category = self.custom_category_name
        instance_key = instance_info['key']
        
        # Key cell
        key_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        key_cell.grid(row=row_idx, column=0, sticky="nsew", padx=1, pady=1)
        key_entry = tk.Entry(key_cell, textvariable=item_info["conjunction_key"], font=("Arial", 9), 
                             justify="center", bg=self.colors['white'], fg=self.colors['text_primary'], 
                             relief="flat", bd=0, width=5)
        key_entry.pack(fill="both", expand=True, padx=3, pady=3)
        self.bind_mousewheel_to_widget(key_cell)
        self.bind_mousewheel_to_widget(key_entry)
        
        # Add/Delete cell - always show + button for custom items
        add_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        add_cell.grid(row=row_idx, column=1, sticky="nsew", padx=1, pady=1)
        self.bind_mousewheel_to_widget(add_cell)
        
        def make_add_callback():
            return lambda: self.add_custom_row()
        
        add_btn = tk.Button(add_cell, text="+", font=("Arial", 12, "bold"),
                             bg=self.colors['light_blue'], fg='white', 
                             relief="flat", cursor="hand2",
                             activebackground=self.colors['primary_blue'],
                             command=make_add_callback())
        add_btn.pack(fill="both", expand=True, padx=2, pady=2)
        self.bind_mousewheel_to_widget(add_btn)
        
        # Item cell - Entry field with checkbox for selection
        item_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                             bg=self.colors['selected'] if item_info["selected"] else self.colors['gray_light'])
        item_cell.grid(row=row_idx, column=2, sticky="nsew", padx=1, pady=1)
        
        # Checkbox for selection
        item_checkbox_var = tk.BooleanVar(value=item_info["selected"])
        def toggle_custom_item():
            item_info["selected"] = not item_info["selected"]
            item_checkbox_var.set(item_info["selected"])
            # Update cell background color
            new_bg = self.colors['selected'] if item_info["selected"] else self.colors['gray_light']
            item_cell.configure(bg=new_bg)
            entry_bg = self.colors['selected'] if item_info["selected"] else self.colors['white']
            item_name_entry.configure(bg=entry_bg)
            checkbox_bg = new_bg
            item_checkbox.configure(bg=checkbox_bg, activebackground=checkbox_bg)
            self.update_all_previews()
        
        item_checkbox = tk.Checkbutton(item_cell, variable=item_checkbox_var,
                                       command=toggle_custom_item,
                                       bg=self.colors['selected'] if item_info["selected"] else self.colors['gray_light'],
                                       activebackground=self.colors['selected'] if item_info["selected"] else self.colors['gray_light'],
                                       highlightthickness=0)
        item_checkbox.pack(side="left", padx=(3, 3))
        item_info["checkbox"] = item_checkbox
        item_info["checkbox_var"] = item_checkbox_var
        
        # Entry field for item name
        item_name_var = tk.StringVar(value=item_info.get("original_name", "Custom Item"))
        item_name_entry = tk.Entry(item_cell, textvariable=item_name_var, font=("Arial", 9),
                                    bg=self.colors['white'] if not item_info["selected"] else self.colors['selected'],
                                    fg=self.colors['text_primary'],
                                    relief="flat", bd=0)
        item_name_entry.pack(side="left", fill="both", expand=True, padx=(0, 3), pady=3)
        
        # Store reference and update on change (but don't update preview)
        def update_item_name(*args):
            new_name = item_name_var.get()
            item_info["original_name"] = new_name
            instance_info['display_name'] = new_name
            # Don't update preview - item name should not appear in preview
        
        item_name_var.trace_add("write", update_item_name)
        item_info["item_name_entry"] = item_name_entry
        self.bind_mousewheel_to_widget(item_cell)
        self.bind_mousewheel_to_widget(item_name_entry)
        self.bind_mousewheel_to_widget(item_checkbox)
        
        # Qty cell
        qty_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        qty_cell.grid(row=row_idx, column=3, sticky="nsew", padx=1, pady=1)
        qty_entry = tk.Entry(qty_cell, textvariable=item_info["qty"], font=("Arial", 9), 
                             justify="center", bg=self.colors['white'], fg=self.colors['text_primary'],
                             relief="flat", bd=0)
        qty_entry.pack(fill="both", expand=True, padx=3, pady=3)
        self.bind_mousewheel_to_widget(qty_cell)
        self.bind_mousewheel_to_widget(qty_entry)
        
        # Unit Price cell
        price_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        price_cell.grid(row=row_idx, column=4, sticky="nsew", padx=1, pady=1)
        unit_price_entry = tk.Entry(price_cell, textvariable=item_info["unit_price"], 
                                     font=("Arial", 9), justify="center", 
                                     bg=self.colors['white'], fg=self.colors['text_primary'], 
                                     relief="flat", bd=0)
        unit_price_entry.pack(fill="both", expand=True, padx=3, pady=3)
        self.bind_mousewheel_to_widget(price_cell)
        self.bind_mousewheel_to_widget(unit_price_entry)
        
        # Total Price cell
        total_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        total_cell.grid(row=row_idx, column=5, sticky="nsew", padx=1, pady=1)
        total_label = tk.Label(total_cell, text="0.00", font=("Arial", 9, "bold"), 
                                bg=self.colors['background'], fg=self.colors['text_primary'],
                                justify="center")
        total_label.pack(fill="both", expand=True, padx=3, pady=3)
        item_info["total_price_label"] = total_label
        self.bind_mousewheel_to_widget(total_cell)
        self.bind_mousewheel_to_widget(total_label)
        
        # Location cell
        location_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        location_cell.grid(row=row_idx, column=6, sticky="nsew", padx=1, pady=1)
        location_entry = tk.Entry(location_cell, textvariable=item_info["location"], 
                                   font=("Arial", 9), bg=self.colors['white'], 
                                   fg=self.colors['text_primary'], relief="flat", bd=0)
        location_entry.pack(fill="both", expand=True, padx=3, pady=3)
        self.bind_mousewheel_to_widget(location_cell)
        self.bind_mousewheel_to_widget(location_entry)
        
        # Additional Info cell
        add_info_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        add_info_cell.grid(row=row_idx, column=7, sticky="nsew", padx=1, pady=1)
        add_info_entry = tk.Entry(add_info_cell, textvariable=item_info["add_info"], 
                                   font=("Arial", 9), bg=self.colors['white'], 
                                   fg=self.colors['text_primary'], relief="flat", bd=0)
        add_info_entry.pack(fill="both", expand=True, padx=3, pady=3)
        self.bind_mousewheel_to_widget(add_info_cell)
        self.bind_mousewheel_to_widget(add_info_entry)
        
        # Live Preview cell
        preview_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        preview_cell.grid(row=row_idx, column=8, sticky="nsew", padx=1, pady=1)
        preview_cell.grid_propagate(False)
        preview_cell.configure(width=350)
        
        preview_text = tk.Text(preview_cell, font=("Arial", 9), width=45,
                               bg=self.colors['preview_bg'], fg=self.colors['text_primary'],
                               relief="flat", bd=0, wrap=tk.WORD, height=6,
                               state=tk.NORMAL)
        preview_text.pack(fill="both", expand=True, padx=3, pady=3)
        preview_text.bind("<KeyRelease>", lambda e, item=item_info: self.on_preview_text_change(item))
        item_info["preview_text"] = preview_text
        
        # Restore preview text content if it exists (from previous view or saved state)
        # Only restore if user_edited is True (meaning it was manually edited)
        # If user_edited is False, the content will be auto-generated by update_total_and_preview
        preview_content = item_info.get("preview_text_content", "")
        has_restored_content = False
        if preview_content and item_info.get("user_edited", False):
            preview_text.insert("1.0", preview_content)
            has_restored_content = True
            # Don't change user_edited flag - it's already set correctly
        
        self.bind_mousewheel_to_widget(preview_cell)
        self.bind_mousewheel_to_widget(preview_text)
        
        # Photo cell
        photo_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
        photo_cell.grid(row=row_idx, column=9, sticky="nsew", padx=1, pady=1)
        
        photo_frame = tk.Frame(photo_cell, bg=self.colors['white'], relief="flat", 
                               bd=1, height=100)
        photo_frame.pack(fill="both", expand=True, padx=3, pady=3)
        photo_frame.pack_propagate(False)
        
        photo_label = tk.Label(photo_frame, text="Click to Select Photo", 
                               font=("Arial", 8), fg=self.colors['gray_medium'],
                               bg=self.colors['white'], cursor="hand2")
        photo_label.pack(fill="both", expand=True)
        
        photo_buttons_frame = tk.Frame(photo_frame, bg=self.colors['white'])
        photo_buttons_frame.pack(side="bottom", fill="x", padx=3, pady=3)
        
        paste_btn = tk.Button(photo_buttons_frame, text="Paste (Ctrl+V)",
                               font=("Arial", 7), bg=self.colors['light_blue'],
                               fg="white", relief="flat", cursor="hand2",
                               command=lambda: self.handle_paste(category, instance_key))
        paste_btn.pack(side="bottom", pady=(2, 0))
        
        self.bind_mousewheel_to_widget(photo_cell)
        self.bind_mousewheel_to_widget(photo_frame)
        self.bind_mousewheel_to_widget(photo_label)
        self.bind_mousewheel_to_widget(photo_buttons_frame)
        self.bind_mousewheel_to_widget(paste_btn)
        
        photo_frame.bind("<Enter>", lambda e, cell=photo_frame: self.on_enter(cell))
        photo_frame.bind("<Leave>", lambda e, cell=photo_frame: self.on_leave(cell))
        
        item_info["photo_frame"] = photo_frame
        item_info["photo_label"] = photo_label
        
        photo_key = f"{category}_{instance_key}"
        if photo_key in self.item_photos and self.item_photos[photo_key]:
            self.load_photo_display(category, instance_key)
        
        def make_photo_callbacks():
            return {
                'click': lambda e: self.select_photo(category, instance_key),
            }
        
        callbacks = make_photo_callbacks()
        photo_label.bind("<Button-1>", callbacks['click'])
        photo_frame.bind("<Button-1>", callbacks['click'])
        
        def on_focus_in(event):
            self.current_photo_item = (category, instance_key)
        
        photo_frame.bind("<FocusIn>", on_focus_in)
        photo_label.bind("<FocusIn>", on_focus_in)
        
        # Bind trace for updates (for custom items, don't auto-update preview if user has edited)
        def safe_update_preview(item_ref):
            if item_ref.get("is_custom", False) and item_ref.get("user_edited", False):
                # Only update total, not preview for custom items with user edits
                try:
                    q_str = item_ref["qty"].get().strip().replace(",", "")
                    p_str = item_ref["unit_price"].get().strip().replace(",", "")
                    q = float(q_str) if q_str else 0.0
                    p = float(p_str) if p_str else 0.0
                    total = round(q * p, 2)
                    if item_ref["total_price_label"] and item_ref["total_price_label"].winfo_exists():
                        item_ref["total_price_label"].config(text=f"{total:.2f}")
                except:
                    pass
            else:
                self.update_total_and_preview(item_ref)
        
        item_info["qty"].trace_add("write", lambda *_args, i=item_info: safe_update_preview(i))
        item_info["unit_price"].trace_add("write", lambda *_args, i=item_info: safe_update_preview(i))
        item_info["location"].trace_add("write", lambda *_args, i=item_info: safe_update_preview(i))
        item_info["add_info"].trace_add("write", lambda *_args, i=item_info: safe_update_preview(i))
        item_info["conjunction_key"].trace_add("write", lambda *_args, i=item_info: self.update_all_previews())
        
        # Initial update - ALWAYS ensure preview is populated for custom items
        # Calculate total price first
        try:
            q_str = item_info["qty"].get().strip().replace(",", "")
            p_str = item_info["unit_price"].get().strip().replace(",", "")
            q = float(q_str) if q_str else 0.0
            p = float(p_str) if p_str else 0.0
            total = round(q * p, 2)
            if item_info["total_price_label"].winfo_exists():
                item_info["total_price_label"].config(text=f"{total:.2f}")
        except:
            total = 0.0
            try:
                if item_info["total_price_label"].winfo_exists():
                    item_info["total_price_label"].config(text="0.00")
            except:
                pass
        
        # Always ensure preview is populated
        # If user_edited and content was restored, keep it
        # Otherwise, generate preview immediately (no debounce for initial load)
        if not (item_info.get("user_edited", False) and has_restored_content):
            # Generate preview immediately for initial load (no debounce)
            self.update_live_preview(item_info, total)
    
    def add_custom_row(self):
        """Add a new blank custom row"""
        if not hasattr(self, 'custom_item_counter'):
            self.custom_item_counter = 1
        
        # Find next available counter
        existing_keys = []
        if self.custom_category_name in self.item_instances:
            for item_list in self.item_instances[self.custom_category_name].values():
                for inst in item_list:
                    existing_keys.append(inst['key'])
        
        while f"custom_item_{self.custom_item_counter}" in existing_keys:
            self.custom_item_counter += 1
        
        instance_key = f"custom_item_{self.custom_item_counter}"
        item_name = f"Custom Item {self.custom_item_counter}"
        
        if self.custom_category_name not in self.item_instances:
            self.item_instances[self.custom_category_name] = {}
        if item_name not in self.item_instances[self.custom_category_name]:
            self.item_instances[self.custom_category_name][item_name] = []
        
        new_instance = {
            'instance_id': 1,
            'display_name': item_name,
            'key': instance_key,
            'is_custom': True
        }
        self.item_instances[self.custom_category_name][item_name].append(new_instance)
        self.custom_item_counter += 1
        
        # Reload custom items to show new row
        self.load_custom_items()
    
    def load_items(self, category):
        # Cancel any pending debounced preview updates for items in this category
        # This prevents stale updates from interfering when widgets are recreated
        if category in self.selected_items:
            for item_key, item_info in self.selected_items[category].items():
                item_id = id(item_info)
                if item_id in self._preview_update_jobs:
                    try:
                        self.root.after_cancel(self._preview_update_jobs[item_id])
                        del self._preview_update_jobs[item_id]
                    except:
                        pass
        
        # Preserve preview text content before destroying widgets (for all categories)
        # Only preserve content if user has actually edited it (user_edited = True)
        # Auto-generated content will be regenerated, so we don't need to preserve it
        if category in self.selected_items:
            for item_key, item_info in self.selected_items[category].items():
                if item_info.get("preview_text") and item_info.get("user_edited", False):
                    try:
                        if item_info["preview_text"].winfo_exists():
                            preview_content = item_info["preview_text"].get("1.0", tk.END).strip()
                            if preview_content:
                                item_info["preview_text_content"] = preview_content
                    except:
                        pass
        
        # Clear scrollable frame (destroy all widgets)
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()

        grid_frame = tk.Frame(self.scrollable_frame, bg=self.colors['white'])
        grid_frame.pack(fill="both", expand=True, padx=4, pady=4)
        
        # Bind mouse wheel events to the grid frame
        self.bind_mousewheel_to_widget(grid_frame)
        
        # This list defines the properties for each column
        # minsize is a fixed minimum size in pixels.
        # weight determines how the column expands to fill extra space.
        # Higher weight means more expansion.
        col_configs = [
            (0, 40, 0),     # Key column (fixed width)
            (1, 40, 0),     # Add/Delete column (fixed width)
            (2, 120, 2),    # Item column (expands)
            (3, 50, 0),     # Qty column (fixed width)
            (4, 70, 0),     # Unit Price column (fixed width)
            (5, 70, 0),     # Total Price column (fixed width)
            (6, 120, 2),    # Location column (expands)
            (7, 150, 2),    # Additional Info column (expands)
            (8, 350, 4),    # Live Preview column (wider - increased from 200 to 350, weight from 2 to 4)

            (9, 150, 4)     # Photo column (much wider - ensures photos are visible)
        ]
        
        for col, min_width, weight in col_configs:
            grid_frame.grid_columnconfigure(col, minsize=min_width, weight=int(weight))

        headings = ["Key", "Add", "Item", "Qty", "Unit Price", "Total Price", "Location", "Additional Info", "Live Preview", "Photo"]
        
        for col, heading in enumerate(headings):
            header_frame = tk.Frame(grid_frame, bg=self.colors['primary_blue'], relief="flat", bd=1)
            header_frame.grid(row=0, column=col, sticky="nsew", padx=1, pady=1)
            
            label = tk.Label(header_frame, text=heading, font=("Arial", 11, "bold"), 
                             bg=self.colors['primary_blue'], fg='white', anchor="w")
            label.pack(fill="both", expand=True, padx=4, pady=4)
            
            # Bind mouse wheel events to header elements
            self.bind_mousewheel_to_widget(header_frame)
            self.bind_mousewheel_to_widget(label)

        if category not in self.selected_items:
            self.selected_items[category] = {}
        
        if category not in self.item_instances:
            self.item_instances[category] = {}

        # Handle Custom category separately - show blank rows
        if category == self.custom_category_name:
            self.load_custom_items()
            return

        if category and category in self.all_items:
            for item_data in self.all_items[category]:
                item_name = item_data['item_name']
                if item_name not in self.item_instances[category] or not self.item_instances[category][item_name]:
                    self.item_instances[category][item_name] = [{
                        'instance_id': 1,
                        'display_name': item_name,
                        'key': f"{item_name}_1"
                    }]

        row_idx = 1
        if category and category in self.all_items:
            all_instances_for_category = []
            for item_name in self.item_instances[category]:
                all_instances_for_category.extend(self.item_instances[category][item_name])
            
            all_instances_for_category.sort(key=lambda x: x['key'])

            for instance_info in all_instances_for_category:
                instance_key = instance_info['key']
                display_name = instance_info['display_name']
                original_name = re.sub(r'#.*', '', display_name).strip()
                item_data = next((item for item in self.all_items[category] if item['item_name'] == original_name), None)
                if not item_data: continue

                if instance_key not in self.selected_items[category]:
                    self.selected_items[category][instance_key] = {
                        "selected": False,
                        "template": item_data['template'],
                        "qty": tk.StringVar(value="0"),
                        "unit_price": tk.StringVar(value=item_data['unit_price']),
                        "location": tk.StringVar(),
                        "add_info": tk.StringVar(),
                        "conjunction_key": tk.StringVar(),
                        "total_price_label": None,
                        "button": None,
                        "preview_text": None,
                        "preview_text_content": "",  # Store preview text content as string
                        "original_name": original_name,
                        "instance_info": instance_info,
                        "photo_frame": None,
                        "photo_label": None,
                        "user_edited": False
                    }
                else:
                    # Preserve existing preview_text_content if it exists
                    if "preview_text_content" not in self.selected_items[category][instance_key]:
                        self.selected_items[category][instance_key]["preview_text_content"] = ""
                    # Ensure template is preserved for existing items
                    if "template" not in self.selected_items[category][instance_key] or not self.selected_items[category][instance_key]["template"]:
                        # Template is missing, restore it from item_data
                        self.selected_items[category][instance_key]["template"] = item_data.get('template', '')

                item_info = self.selected_items[category][instance_key]
                
                key_cell = tk.Frame(grid_frame, bd=1, relief="solid", bg=self.colors['gray_light'])
                key_cell.grid(row=row_idx, column=0, sticky="nsew", padx=1, pady=1)
                key_entry = tk.Entry(key_cell, textvariable=item_info["conjunction_key"], font=("Arial", 9), justify="center", bg=self.colors['white'], fg=self.colors['text_primary'], relief="flat", bd=0, width=5)
                key_entry.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind mouse wheel events to key cell elements
                self.bind_mousewheel_to_widget(key_cell)
                self.bind_mousewheel_to_widget(key_entry)

                add_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                    bg=self.colors['gray_light'])
                add_cell.grid(row=row_idx, column=1, sticky="nsew", padx=1, pady=1)
                
                # Bind mouse wheel events to add cell
                self.bind_mousewheel_to_widget(add_cell)
                
                if instance_info['instance_id'] == 1:
                    def make_add_callback(cat=category, name=original_name):
                        return lambda: self.add_item_instance(cat, name)
                    
                    add_btn = tk.Button(add_cell, text="+", font=("Arial", 12, "bold"),
                                         bg=self.colors['light_blue'], fg='white', 
                                         relief="flat", cursor="hand2",
                                         activebackground=self.colors['primary_blue'],
                                         command=make_add_callback())
                    add_btn.pack(fill="both", expand=True, padx=2, pady=2)
                    self.bind_mousewheel_to_widget(add_btn)
                else:
                    def make_delete_callback(cat=category, orig_name=original_name, inst_key=instance_key):
                        return lambda: self.delete_item_instance(cat, orig_name, inst_key)
                    
                    delete_btn = tk.Button(add_cell, text="X", font=("Arial", 10, "bold"),
                                             bg='#dc3545', fg='white',
                                             relief="flat", cursor="hand2",
                                             activebackground='#c82333',
                                             command=make_delete_callback())
                    delete_btn.pack(fill="both", expand=True, padx=2, pady=2)
                    self.bind_mousewheel_to_widget(delete_btn)

                item_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                     bg=self.colors['gray_light'])
                item_cell.grid(row=row_idx, column=2, sticky="nsew", padx=1, pady=1)
                
                def make_toggle_callback(cat=category, name=instance_key):
                    return lambda: self.toggle_item(cat, name)

                btn = tk.Button(item_cell, text=display_name, font=("Arial", 9),
                                 bg=self.colors['selected'] if item_info["selected"] else self.colors['white'],
                                 fg=self.colors['text_primary'], anchor="w", relief="flat", cursor="hand2",
                                 activebackground=self.colors['selected'],
                                 command=make_toggle_callback())
                btn.pack(fill="both", expand=True, padx=3, pady=3)
                item_info["button"] = btn
                
                # Bind mouse wheel events to item cell elements
                self.bind_mousewheel_to_widget(item_cell)
                self.bind_mousewheel_to_widget(btn)

                qty_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                     bg=self.colors['gray_light'])
                qty_cell.grid(row=row_idx, column=3, sticky="nsew", padx=1, pady=1)
                qty_entry = tk.Entry(qty_cell, textvariable=item_info["qty"], font=("Arial", 9), 
                                     justify="center", bg=self.colors['white'], fg=self.colors['text_primary'],
                                     relief="flat", bd=0)
                qty_entry.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind mouse wheel events to qty cell elements
                self.bind_mousewheel_to_widget(qty_cell)
                self.bind_mousewheel_to_widget(qty_entry)

                price_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                      bg=self.colors['gray_light'])
                price_cell.grid(row=row_idx, column=4, sticky="nsew", padx=1, pady=1)
                unit_price_entry = tk.Entry(price_cell, textvariable=item_info["unit_price"], 
                                             font=("Arial", 9), justify="center", 
                                             bg=self.colors['white'], fg=self.colors['text_primary'], relief="flat", bd=0)
                unit_price_entry.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind mouse wheel events to price cell elements
                self.bind_mousewheel_to_widget(price_cell)
                self.bind_mousewheel_to_widget(unit_price_entry)

                total_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                      bg=self.colors['gray_light'])
                total_cell.grid(row=row_idx, column=5, sticky="nsew", padx=1, pady=1)
                total_label = tk.Label(total_cell, text="0.00", font=("Arial", 9, "bold"), 
                                        bg=self.colors['background'], fg=self.colors['text_primary'],
                                        justify="center")
                total_label.pack(fill="both", expand=True, padx=3, pady=3)
                item_info["total_price_label"] = total_label
                
                # Bind mouse wheel events to total cell elements
                self.bind_mousewheel_to_widget(total_cell)
                self.bind_mousewheel_to_widget(total_label)

                location_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                         bg=self.colors['gray_light'])
                location_cell.grid(row=row_idx, column=6, sticky="nsew", padx=1, pady=1)
                location_entry = tk.Entry(location_cell, textvariable=item_info["location"], 
                                         font=("Arial", 9), bg=self.colors['white'], fg=self.colors['text_primary'],
                                         relief="flat", bd=0)
                location_entry.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind mouse wheel events to location cell elements
                self.bind_mousewheel_to_widget(location_cell)
                self.bind_mousewheel_to_widget(location_entry)

                info_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                     bg=self.colors['gray_light'])
                info_cell.grid(row=row_idx, column=7, sticky="nsew", padx=1, pady=1)
                add_info_entry = tk.Entry(info_cell, textvariable=item_info["add_info"], 
                                         font=("Arial", 9), bg=self.colors['white'], fg=self.colors['text_primary'],
                                         relief="flat", bd=0)
                add_info_entry.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind mouse wheel events to info cell elements
                self.bind_mousewheel_to_widget(info_cell)
                self.bind_mousewheel_to_widget(add_info_entry)

                preview_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                         bg=self.colors['gray_light'])
                preview_cell.grid(row=row_idx, column=8, sticky="nsew", padx=1, pady=1)
                # Keep this column compact: do not let children expand the cell's size
                preview_cell.grid_propagate(False)
                preview_cell.configure(width=350)  # Increased from 80 to 350 to match column width

                preview_text = tk.Text(preview_cell, font=("Arial", 9), width=45,  # Increased from 12 to 45 for wider text
                                       bg=self.colors['preview_bg'], fg=self.colors['text_primary'],
                                       relief="flat", bd=0, wrap=tk.WORD, height=6,
                                       state=tk.NORMAL)
                preview_text.pack(fill="both", expand=True, padx=3, pady=3)
                
                # Bind text changes to update the generated bids
                preview_text.bind("<KeyRelease>", lambda e, item=item_info: self.on_preview_text_change(item))
                
                item_info["preview_text"] = preview_text
                
                # Restore preview text content if it exists (from previous view or saved state)
                # Only restore if user_edited is True (meaning it was manually edited)
                # If user_edited is False, the content will be auto-generated by update_total_and_preview
                preview_content = item_info.get("preview_text_content", "")
                has_restored_content = False
                if preview_content and item_info.get("user_edited", False):
                    preview_text.insert("1.0", preview_content)
                    has_restored_content = True
                    # Don't change user_edited flag - it's already set correctly
                
                # Bind mouse wheel events to preview cell elements
                self.bind_mousewheel_to_widget(preview_cell)
                self.bind_mousewheel_to_widget(preview_text)

                photo_cell = tk.Frame(grid_frame, bd=1, relief="solid", 
                                      bg=self.colors['gray_light'])
                photo_cell.grid(row=row_idx, column=9, sticky="nsew", padx=1, pady=1)
                
                photo_frame = tk.Frame(photo_cell, bg=self.colors['white'], relief="flat", 
                                       bd=1, height=100)
                photo_frame.pack(fill="both", expand=True, padx=3, pady=3)
                photo_frame.pack_propagate(False)
                
                photo_label = tk.Label(photo_frame, text="Click to Select Photo", 
                                       font=("Arial", 8), fg=self.colors['gray_medium'],
                                       bg=self.colors['white'], cursor="hand2")
                photo_label.pack(fill="both", expand=True)
                
                photo_buttons_frame = tk.Frame(photo_frame, bg=self.colors['white'])
                photo_buttons_frame.pack(side="bottom", fill="x", padx=3, pady=3)
                
                paste_btn = tk.Button(photo_buttons_frame, text="Paste (Ctrl+V)",
                                       font=("Arial", 7), bg=self.colors['light_blue'],
                                       fg="white", relief="flat", cursor="hand2",
                                       command=lambda c=category, k=instance_key: self.handle_paste(c, k))
                paste_btn.pack(side="bottom", pady=(2, 0))
                
                # Bind mouse wheel events to photo cell elements
                self.bind_mousewheel_to_widget(photo_cell)
                self.bind_mousewheel_to_widget(photo_frame)
                self.bind_mousewheel_to_widget(photo_label)
                self.bind_mousewheel_to_widget(photo_buttons_frame)
                self.bind_mousewheel_to_widget(paste_btn)
                
                photo_frame.bind("<Enter>", lambda e, cell=photo_frame: self.on_enter(cell))
                photo_frame.bind("<Leave>", lambda e, cell=photo_frame: self.on_leave(cell))
                
                item_info["photo_frame"] = photo_frame
                item_info["photo_label"] = photo_label
                
                photo_key = f"{category}_{instance_key}"
                if photo_key in self.item_photos and self.item_photos[photo_key]:
                    self.load_photo_display(category, instance_key)
                
                def make_photo_callbacks(cat=category, key=instance_key):
                    return {
                        'click': lambda e: self.select_photo(cat, key),
                    }
                
                callbacks = make_photo_callbacks()
                photo_label.bind("<Button-1>", callbacks['click'])
                photo_frame.bind("<Button-1>", callbacks['click'])
                
                def on_focus_in(event, cat=category, key=instance_key):
                    self.current_photo_item = (cat, key)
                
                photo_frame.bind("<FocusIn>", on_focus_in)
                photo_label.bind("<FocusIn>", on_focus_in)

                item_info["qty"].trace_add("write", lambda *_args, i=item_info: self.update_total_and_preview(i))
                item_info["unit_price"].trace_add("write", lambda *_args, i=item_info: self.update_total_and_preview(i))
                item_info["location"].trace_add("write", lambda *_args, i=item_info: self.update_total_and_preview(i))
                item_info["add_info"].trace_add("write", lambda *_args, i=item_info: self.update_total_and_preview(i))
                item_info["conjunction_key"].trace_add("write", lambda *_args, i=item_info: self.update_all_previews())
                
                # Initial update - ALWAYS ensure preview is populated
                # Calculate total price first
                try:
                    q_str = item_info["qty"].get().strip().replace(",", "")
                    p_str = item_info["unit_price"].get().strip().replace(",", "")
                    q = float(q_str) if q_str else 0.0
                    p = float(p_str) if p_str else 0.0
                    total = round(q * p, 2)
                    if item_info["total_price_label"].winfo_exists():
                        item_info["total_price_label"].config(text=f"{total:.2f}")
                except:
                    total = 0.0
                    try:
                        if item_info["total_price_label"].winfo_exists():
                            item_info["total_price_label"].config(text="0.00")
                    except:
                        pass
                
                # Always ensure preview is populated for ALL items
                # Check if preview widget exists and is ready before generating
                if item_info.get("preview_text") and item_info["preview_text"].winfo_exists():
                    # Only skip if user has edited AND content was successfully restored
                    if not (item_info.get("user_edited", False) and has_restored_content):
                        # Ensure template exists before generating preview
                        if not item_info.get("template"):
                            # Restore template from item_data
                            item_original_name = item_info.get("original_name", "")
                            if not item_original_name:
                                # Fallback to original_name from outer scope if available
                                item_original_name = original_name
                            if category in self.all_items and item_original_name:
                                item_data = next((item for item in self.all_items[category] if item['item_name'] == item_original_name), None)
                                if item_data and item_data.get('template'):
                                    item_info["template"] = item_data['template']
                        
                        # Generate preview immediately for initial load (no debounce)
                        # Use a direct call that bypasses some checks for initial load
                        try:
                            self._generate_preview_direct(item_info, total)
                            # Verify preview was actually populated - if not, retry
                            preview_check = item_info["preview_text"].get("1.0", tk.END).strip()
                            if not preview_check:
                                # Preview is still empty, retry after a short delay
                                self.root.after(100, lambda i=item_info, t=total: self._ensure_preview_populated(i, t))
                        except Exception as e:
                            # If update fails, try again after a short delay to ensure widget is ready
                            print(f"Preview update failed for {instance_key}, retrying: {e}")
                            self.root.after(100, lambda i=item_info, t=total: self._ensure_preview_populated(i, t))
                
                row_idx += 1
        
        # Reset scroll position to top when switching categories
        self.reset_scroll_to_top()
        
    def reset_scroll_to_top(self):
        """Reset the canvas scroll position to the top."""
        def _perform_scroll_reset():
            try:
                # Update the canvas to ensure the scrollregion is properly set
                self.canvas.update_idletasks()
                # Move the view to the top (0.0 = top, 1.0 = bottom)
                self.canvas.yview_moveto(0.0)
            except Exception:
                # If scrolling fails for any reason, continue silently
                pass
        
        # Use a small delay to ensure the UI has fully updated before scrolling
        self.root.after(50, _perform_scroll_reset)
        
    def save_state(self, silent=False):
        """Saves the current state to Supabase (or local JSON file as fallback)."""
        wo_number = self.wo_entry.get().strip()
        if not wo_number:
            if not silent:
                messagebox.showwarning("Warning", "Please enter a Work Order Number to save the state.")
            return

        state = {
            "selected_items": {},
            "item_photos": {}
        }

        for category, items in self.selected_items.items():
            state["selected_items"][category] = {}
            for item_key, item_data in items.items():
                # Save preview text content if it exists (for custom items especially)
                preview_text_content = ""
                if item_data.get("preview_text"):
                    try:
                        if item_data["preview_text"].winfo_exists():
                            preview_text_content = item_data["preview_text"].get("1.0", tk.END).strip()
                    except:
                        pass
                # Also check if preview_text_content was already saved as a string
                if not preview_text_content and item_data.get("preview_text_content"):
                    preview_text_content = item_data["preview_text_content"]
                
                state["selected_items"][category][item_key] = {
                    "selected": item_data["selected"],
                    "template": item_data["template"],
                    "qty": item_data["qty"].get(),
                    "unit_price": item_data["unit_price"].get(),
                    "location": item_data["location"].get(),
                    "add_info": item_data["add_info"].get(),
                    "original_name": item_data["original_name"],
                    "instance_info": item_data["instance_info"],
                    "conjunction_key": item_data["conjunction_key"].get(),
                    "preview_text_content": preview_text_content,  # Save preview text content
                    "user_edited": item_data.get("user_edited", False),  # Save user_edited flag
                    "is_custom": item_data.get("is_custom", False)  # Save is_custom flag
                }
        
        # Save photo data - preserve the full dictionary structure (path, original image, etc.)
        # But for saving to database, we only need the path (the upload function will handle the rest)
        print(f"\n[save_state] Saving photos. self.item_photos has {len(self.item_photos)} entries")
        print(f"[save_state] Photo keys: {list(self.item_photos.keys())}")
        
        for photo_key, photo_data in self.item_photos.items():
            print(f"[save_state] Processing photo_key: {photo_key}, type: {type(photo_data)}")
            if photo_data:
                # Store the full photo_data dict if it exists, or just the path if it's a string
                if isinstance(photo_data, dict):
                    print(f"  photo_data dict keys: {list(photo_data.keys())}")
                    # Debug: print the actual path value
                    path_value = photo_data.get('path', None)
                    print(f"  path value: {repr(path_value)} (type: {type(path_value)})")
                    
                    # Keep the path and any other metadata, but not the PIL Image object (it's not JSON serializable)
                    photo_dict = {}
                    
                    # Check if path exists and is not empty
                    if 'path' in photo_data:
                        path_value = photo_data['path']
                        if path_value and str(path_value).strip():  # Check if path is not None and not empty
                            photo_dict['path'] = str(path_value).strip()
                            print(f"  ✅ Found valid path: {photo_dict['path']}")
                        else:
                            print(f"  ⚠️ Path exists but is empty/None: {repr(path_value)}")
                    
                    if 'url' in photo_data and photo_data.get('url'):
                        photo_dict['url'] = photo_data['url']
                        print(f"  ✅ Found URL: {photo_data['url']}")
                    if 'storage_path' in photo_data and photo_data.get('storage_path'):
                        photo_dict['storage_path'] = photo_data['storage_path']
                        print(f"  ✅ Found storage_path: {photo_data['storage_path']}")
                    if 'original_path' in photo_data and photo_data.get('original_path'):
                        photo_dict['original_path'] = photo_data['original_path']
                        print(f"  ✅ Found original_path: {photo_data['original_path']}")
                    
                    if photo_dict:  # Only add if we have at least the path
                        state["item_photos"][photo_key] = photo_dict
                        print(f"  ✅ Added to state: {photo_key} -> {photo_dict}")
                    else:
                        print(f"  ⚠️ No valid path/URL/storage_path found, skipping")
                        # Even if path is empty, try to preserve the dict structure for debugging
                        print(f"  📋 Full photo_data: {photo_data}")
                elif isinstance(photo_data, str):
                    # If it's just a string (path), wrap it in a dict
                    state["item_photos"][photo_key] = {'path': photo_data}
                    print(f"  ✅ String path converted: {photo_data}")
                else:
                    print(f"  ⚠️ Unknown photo_data type: {type(photo_data)}")
            else:
                print(f"  ⚠️ photo_data is empty/None for {photo_key}")
        
        print(f"[save_state] Final state['item_photos'] has {len(state['item_photos'])} entries: {list(state['item_photos'].keys())}")
        
        # Get property address from entry field
        property_address = self.address_entry.get().strip() if hasattr(self, 'address_entry') else self.property_address or ""
        
        # Get username - ensure it's always available
        username = self.username if hasattr(self, 'username') and self.username else None
        if not username and self.db and self.user_id:
            # Try to get username from database if not stored
            try:
                user = self.db.get_user_by_id(self.user_id)
                if user:
                    username = user.get('username')
            except Exception as e:
                print(f"Could not fetch username: {e}")
        
        # Try to save to Supabase first
        if self.db and self.user_id:
            try:
                # Get client_code and wo_type from entry fields if they exist
                client_code = None
                wo_type = None
                if hasattr(self, 'client_code_entry'):
                    client_code = self.client_code_entry.get().strip() or None
                if hasattr(self, 'wo_type_entry'):
                    wo_type = self.wo_type_entry.get().strip() or None
                
                # Save to database with all project fields and username
                self.db.save_bid(wo_number, self.user_id, state["selected_items"], state["item_photos"], 
                               property_address, username, client_code, wo_type)
                # Update stored property_address
                self.property_address = property_address
                if not silent:
                    try:
                        messagebox.showinfo("Success", f"Project saved successfully to cloud!")
                    except Exception:
                        pass
                # Fire callback and return - don't save locally if DB save succeeded
                self._safe_post_save_callback()
                return
            except Exception as e:
                if not silent:
                    messagebox.showerror("Error", f"Failed to save to database: {e}\nTrying local backup...")
                # Fall through to local save
                print(f"Database save failed: {e}")
                # Don't disable DB - keep trying
        
        # Fallback to local JSON file (or if DB not available)
        if not self.db or not self.user_id:
            state_file_path = os.path.join(self.app_data_dir, f"WO_{wo_number}.json")
            # Include property_address in local save if available
            if property_address:
                state["property_address"] = property_address
            try:
                with open(state_file_path, 'w') as f:
                    json.dump(state, f, indent=4)
                if not silent:
                    try:
                        messagebox.showinfo("Success", f"State saved locally to:\n{os.path.basename(state_file_path)}")
                    except Exception:
                        pass
            except Exception as e:
                if not silent:
                    messagebox.showerror("Error", f"Failed to save state: {e}")
                return

        # Fire post-save callback safely (if dashboard is in a state to handle it)
        if callable(self.on_save_callback):
            try:
                # Defer to event loop to avoid running during widget teardown
                self.root.after(0, self._safe_post_save_callback)
            except Exception:
                pass

    def _safe_post_save_callback(self):
        try:
            if callable(self.on_save_callback):
                self.on_save_callback()
        except Exception:
            # Ignore UI errors from other windows (e.g., when user navigated away)
            pass

    def load_state(self):
        """Loads a saved state from Supabase (or local JSON file as fallback) and populates the UI."""
        wo_number = self.wo_entry.get().strip()
        if not wo_number:
            messagebox.showwarning("Warning", "Please enter a Work Order Number to load the state.")
            return
        
        state = None
        
        # Try to load from Supabase first - can load from any user
        if self.db:
            try:
                bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                if bid_data:
                    state = {
                        "selected_items": bid_data.get("selected_items", {}),
                        "item_photos": bid_data.get("item_photos", {})
                    }
                    # Load property address if available
                    if bid_data.get('property_address') and hasattr(self, 'address_entry'):
                        self.property_address = bid_data['property_address']
                        self.address_entry.delete(0, tk.END)
                        self.address_entry.insert(0, self.property_address)
                    
                    # Load client_code if available
                    if bid_data.get('client_code') and hasattr(self, 'client_code_entry'):
                        self.client_code_entry.delete(0, tk.END)
                        self.client_code_entry.insert(0, bid_data['client_code'])
                    
                    # Load wo_type if available
                    if bid_data.get('wo_type') and hasattr(self, 'wo_type_entry'):
                        self.wo_type_entry.delete(0, tk.END)
                        self.wo_type_entry.insert(0, bid_data['wo_type'])
            except Exception as e:
                print(f"Error loading from database: {e}")
                # Fall through to local load
        
        # Fallback to local JSON file
        if not state:
            state_file_path = os.path.join(self.app_data_dir, f"WO_{wo_number}.json")
            if os.path.exists(state_file_path):
                try:
                    with open(state_file_path, 'r') as f:
                        state = json.load(f)
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to load state: {e}")
                    return
            else:
                messagebox.showerror("Error", f"No saved state found for WO '{wo_number}'.")
                return
        
        try:
            
            self.selected_items = {}
            self.item_photos = {}
            self.item_instances = {}
            # Ensure Custom category is initialized
            if self.custom_category_name not in self.item_instances:
                self.item_instances[self.custom_category_name] = {}
            if self.custom_category_name not in self.selected_items:
                self.selected_items[self.custom_category_name] = {}

            for category, items in state.get("selected_items", {}).items():
                self.selected_items[category] = {}
                self.item_instances[category] = {}
                
                for item_key, item_data in items.items():
                    original_name = item_data.get('original_name', 'Unknown')
                    instance_info = item_data.get('instance_info', {})
                    
                    # For custom items, use item_key as item_name if original_name is missing
                    if category == self.custom_category_name and not original_name:
                        # Try to extract from instance_info or use a default
                        original_name = instance_info.get('display_name', 'Custom Item')
                    
                    if original_name not in self.item_instances[category]:
                        self.item_instances[category][original_name] = []
                    self.item_instances[category][original_name].append(instance_info)

                    item_data["qty"] = tk.StringVar(value=item_data.get("qty", "0"))
                    item_data["unit_price"] = tk.StringVar(value=item_data.get("unit_price", "0.00"))
                    item_data["location"] = tk.StringVar(value=item_data.get("location", ""))
                    item_data["add_info"] = tk.StringVar(value=item_data.get("add_info", ""))
                    item_data["conjunction_key"] = tk.StringVar(value=item_data.get("conjunction_key", ""))
                    # Ensure is_custom flag is preserved
                    if category == self.custom_category_name:
                        item_data["is_custom"] = True
                    # Preserve preview_text_content and user_edited from saved state
                    if "preview_text_content" not in item_data:
                        item_data["preview_text_content"] = item_data.get("preview_text_content", "")
                    if "user_edited" not in item_data:
                        item_data["user_edited"] = item_data.get("user_edited", False)
                    self.selected_items[category][item_key] = item_data
            
            # Load photos from saved state
            for photo_key, photo_data in state.get("item_photos", {}).items():
                # Handle both dict format (new) and string path format (old)
                if isinstance(photo_data, dict):
                    photo_path = photo_data.get('path') or photo_data.get('original_path') or photo_data.get('url')
                else:
                    photo_path = photo_data  # Old format: just a string path
                
                if photo_path and os.path.exists(str(photo_path)) and Image:
                    try:
                        image = Image.open(photo_path)
                        self.item_photos[photo_key] = {'original': image, 'path': str(photo_path)}
                    except Exception as e:
                        print(f"Error loading photo {photo_key} from {photo_path}: {e}")
                elif photo_path and Image:
                    # Path might be a URL or storage path
                    print(f"Photo path exists but file not found: {photo_path}")
                    # Store the reference anyway for later download
                    self.item_photos[photo_key] = {'path': str(photo_path)}
            
            self.update_bid_buttons()
            if self.active_category:
                self.load_items(self.active_category)
            
            # State loaded silently - no dialog box

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load state: {e}")
    
    def select_photo(self, category, item_key):
        file_path = filedialog.askopenfilename(
            title="Select Photo",
            filetypes=[
                ("Image files", "*.jpg *.jpeg *.png *.gif *.bmp"),
                ("All files", "*.*")
            ]
        )
        
        if file_path:
            self.load_photo(file_path, category, item_key)
    
    def on_enter(self, widget):
        """Highlight the frame on mouse enter."""
        widget.configure(bg=self.colors['gray_light'])

    def on_leave(self, widget):
        """Reset the frame color on mouse leave."""
        widget.configure(bg=self.colors['white'])
        
    def load_photo(self, file_path, category, item_key):
        if Image is None:
            messagebox.showerror("Error", "PIL/Pillow is required for photo support. Install with: pip install pillow")
            return
        
        if not file_path or not file_path.strip():
            print(f"[load_photo] ⚠️ Empty file_path provided for {category}_{item_key}")
            return
            
        try:
            # Normalize the path and ensure it exists
            file_path = os.path.normpath(file_path.strip())
            if not os.path.exists(file_path):
                messagebox.showerror("Error", f"Photo file not found:\n{file_path}")
                return
            
            image = Image.open(file_path)
            
            photo_key = f"{category}_{item_key}"
            self.item_photos[photo_key] = {
                'original': image,
                'path': file_path  # Store the normalized, validated path
            }
            
            print(f"[load_photo] ✅ Loaded photo: {photo_key} from {file_path}")
            self.load_photo_display(category, item_key)
                
        except Exception as e:
            print(f"[load_photo] ❌ Error loading photo: {e}")
            messagebox.showerror("Error", f"Failed to load image:\n{str(e)}")
            
    def load_photo_display(self, category, item_key):
        """Displays a photo in the photo bucket from a stored image object."""
        if Image is None: return
        
        photo_key = f"{category}_{item_key}"
        if photo_key not in self.item_photos or not self.item_photos[photo_key]:
            return
            
        try:
            image_copy = self.item_photos[photo_key]['original'].copy()
            image_copy.thumbnail((180, 100), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(image_copy)
            
            item_info = self.selected_items[category][item_key]
            if item_info["photo_label"]:
                item_info["photo_label"].configure(image=photo, text="")
                item_info["photo_label"].image = photo
                
                remove_btn = tk.Button(item_info["photo_frame"], text="✕", 
                                        font=("Arial", 8, "bold"), bg='red', fg='white',
                                        command=lambda: self.remove_photo(category, item_key))
                remove_btn.place(relx=1.0, rely=0.0, anchor='ne', width=20, height=20)
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to display image:\n{str(e)}")
    
    def remove_photo(self, category, item_key):
        photo_key = f"{category}_{item_key}"
        if photo_key in self.item_photos:
            del self.item_photos[photo_key]
        
        item_info = self.selected_items[category][item_key]
        if item_info["photo_label"]:
            item_info["photo_label"].configure(image="", text="Click to Select Photo")
            item_info["photo_label"].image = None
        
        for widget in item_info["photo_frame"].winfo_children():
            if isinstance(widget, tk.Button) and widget.cget("text") == "✕":
                widget.destroy()
    
    def handle_paste(self, category, item_key):
        if Image is None or ImageGrab is None:
            messagebox.showerror("Error", "PIL/Pillow is required for photo support. Install with: pip install pillow")
            return
            
        try:
            image = ImageGrab.grabclipboard()
            
            if image and isinstance(image, Image.Image):
                # Save pasted image to a temporary file so we have a path
                photo_key = f"{category}_{item_key}"
                
                # Create a temporary file for the pasted image
                temp_dir = os.path.join(self.app_data_dir, "temp_photos")
                os.makedirs(temp_dir, exist_ok=True)
                
                # Create a unique filename (sanitize photo_key for filename)
                import time
                safe_key = photo_key.replace(' ', '_').replace('/', '_').replace('\\', '_')
                temp_filename = f"pasted_{safe_key}_{int(time.time())}.png"
                temp_path = os.path.join(temp_dir, temp_filename)
                
                # Save the image
                image.save(temp_path, 'PNG')
                
                print(f"[handle_paste] ✅ Saved pasted image to: {temp_path}")
                
                self.item_photos[photo_key] = {
                    'original': image,
                    'path': temp_path  # Now we have a valid path!
                }
                
                self.load_photo_display(category, item_key)
            else:
                messagebox.showinfo("Paste", "No image found in clipboard")
        except:
            messagebox.showinfo("Paste", "No image found in clipboard")
    
    def handle_global_paste(self, event):
        if hasattr(self, 'current_photo_item') and self.current_photo_item:
            category, item_key = self.current_photo_item
            self.handle_paste(category, item_key)
    
    def add_item_instance(self, category, item_name):
        item_data = next((item for item in self.all_items[category] if item['item_name'] == item_name), None)
        if not item_data: return

        if category not in self.item_instances:
            self.item_instances[category] = {}
        if item_name not in self.item_instances[category]:
            self.item_instances[category][item_name] = []
        
        existing_numbers = [inst['instance_id'] for inst in self.item_instances[category][item_name]]
        next_number = max(existing_numbers) + 1 if existing_numbers else 2
        
        new_instance = {
            'instance_id': next_number,
            'display_name': f"{item_name} #{next_number}",
            'key': f"{item_name}_{next_number}"
        }
        
        self.item_instances[category][item_name].append(new_instance)
        
        self.load_items(category)

    def delete_item_instance(self, category, original_name, instance_key):
        if category in self.item_instances and original_name in self.item_instances[category]:
            self.item_instances[category][original_name] = [
                inst for inst in self.item_instances[category][original_name] 
                if inst['key'] != instance_key
            ]
            
            if instance_key in self.selected_items[category]:
                del self.selected_items[category][instance_key]
            
            photo_key = f"{category}_{instance_key}"
            if photo_key in self.item_photos:
                del self.item_photos[photo_key]
        
        self.load_items(category)

    def _ensure_preview_populated(self, item_info, total_price, retry_count=0):
        """Ensure preview is populated - retry mechanism for initial load."""
        if not item_info.get("preview_text"):
            return
        try:
            if not item_info["preview_text"].winfo_exists():
                return
            # Check if preview is empty
            current_preview = item_info["preview_text"].get("1.0", tk.END).strip()
            if not current_preview and not item_info.get("user_edited", False):
                # Preview is empty and not user-edited, generate it
                # Ensure template exists before trying to generate
                if not item_info.get("template"):
                    # Try to restore template from item_data
                    original_name = item_info.get("original_name", "")
                    instance_info = item_info.get("instance_info", {})
                    category = None
                    # Find which category this item belongs to
                    for cat, items in self.selected_items.items():
                        if item_info in items.values():
                            category = cat
                            break
                    if category and category in self.all_items:
                        item_data = next((i for i in self.all_items[category] if i['item_name'] == original_name), None)
                        if item_data and item_data.get('template'):
                            item_info["template"] = item_data['template']
                
                # Try to generate preview
                self.update_live_preview(item_info, total_price)
                
                # Verify it was populated, retry if needed (max 3 retries)
                new_preview = item_info["preview_text"].get("1.0", tk.END).strip()
                if not new_preview and retry_count < 3:
                    # Still empty, retry after a longer delay
                    self.root.after(200 * (retry_count + 1), lambda: self._ensure_preview_populated(item_info, total_price, retry_count + 1))
        except Exception as e:
            print(f"Error ensuring preview populated: {e}")
            # Retry on error if we haven't exceeded max retries
            if retry_count < 3:
                self.root.after(200 * (retry_count + 1), lambda: self._ensure_preview_populated(item_info, total_price, retry_count + 1))
    
    def update_total_and_preview(self, item_info):
        """Update total price immediately (lightweight), but debounce preview update."""
        try:
            q_str = item_info["qty"].get().strip().replace(",", "")
            p_str = item_info["unit_price"].get().strip().replace(",", "")
            q = float(q_str) if q_str else 0.0
            p = float(p_str) if p_str else 0.0
            total = round(q * p, 2)
            
            # Update total price immediately (lightweight operation)
            if item_info["total_price_label"].winfo_exists():
                item_info["total_price_label"].config(text=f"{total:.2f}")
        except ValueError:
            total = 0.0
            try:
                if item_info["total_price_label"].winfo_exists():
                    item_info["total_price_label"].config(text="0.00")
            except:
                pass
        except:
            # Widget might be destroyed
            total = 0.0

        # Debounce the preview update (heavy operation)
        item_key = id(item_info)  # Use item ID as key
        
        # Cancel any pending update for this item
        if item_key in self._preview_update_jobs:
            self.root.after_cancel(self._preview_update_jobs[item_key])
        
        # Schedule a new update after debounce delay
        def delayed_update():
            if item_key in self._preview_update_jobs:
                del self._preview_update_jobs[item_key]
            self.update_live_preview(item_info, total)
        
        job_id = self.root.after(self._debounce_delay, delayed_update)
        self._preview_update_jobs[item_key] = job_id

    def _generate_preview_direct(self, item, total_price):
        """Generate preview directly without early return checks - for initial load only."""
        if not item.get("preview_text") or not item["preview_text"].winfo_exists():
            return
        
        # Get template - ensure it exists
        template = item.get("template", "")
        if not template:
            return
        
        # Get current values
        qty = item["qty"].get().strip() or "0"
        location = item["location"].get().strip() or "N/A"
        add_info = item["add_info"].get().strip()
        conjunction_key = item["conjunction_key"].get().strip().upper()
        
        # Extract unit_price for regular items
        unit_price = item.get("unit_price")
        if unit_price:
            unit_price = unit_price.get().strip().replace(",", "") if hasattr(unit_price, "get") else str(unit_price).strip().replace(",", "")
        else:
            unit_price = "0"
        
        # Handle zero or empty quantity
        try:
            qty_float = float(qty) if qty else 0.0
        except ValueError:
            qty_float = 0.0
        
        try:
            unit_price_float = float(unit_price) if unit_price else 0.0
        except ValueError:
            unit_price_float = 0.0
        
        calculated_total_price = qty_float * unit_price_float
        
        # Handle custom items
        if item.get("is_custom", False):
            description = add_info if add_info else ""
            try:
                bid_text = template.format(
                    description=description,
                    quantity=qty,
                    location=location,
                    info=add_info
                )
            except KeyError:
                bid_text = description if description else ""
        else:
            # Safe template formatting
            try:
                bid_text = template.format(
                    quantity=qty or "0",
                    qty=qty or "0",
                    lf=qty or "0",
                    location=location or "N/A",
                    info=add_info or "",
                    add_info=add_info or "",
                    total=calculated_total_price,
                    total_price=calculated_total_price,
                    unit_price=unit_price or "0",
                    price=unit_price or "0",
                    cause=add_info or ""
                )
            except KeyError:
                template_copy = template
                template_copy = template_copy.replace("{lf}", qty or "0")
                template_copy = template_copy.replace("{qty}", qty or "0")
                template_copy = template_copy.replace("{unit_price}", unit_price or "0")
                template_copy = template_copy.replace("{price}", unit_price or "0")
                template_copy = template_copy.replace("{add_info}", add_info or "")
                
                try:
                    bid_text = template_copy.format(
                        quantity=qty or "0",
                        location=location or "N/A",
                        info=add_info or "",
                        total=calculated_total_price
                    )
                except:
                    bid_text = template_copy.replace("{quantity}", qty or "0") \
                                          .replace("{location}", location or "N/A") \
                                          .replace("{info}", add_info or "") \
                                          .replace("{total}", f"{calculated_total_price:.2f}")
        
        # Add conjunction prefix/suffix if needed
        # Only show conjunction grouping for SELECTED items
        conjunction_prefix = ""
        conjunction_suffix = ""
        if conjunction_key and item.get("selected", False):
            # Find only SELECTED items with the same key (for numbering and message)
            selected_conjuncted_items = [
                i for cat in self.selected_items.values() 
                for i in cat.values() 
                if i.get("conjunction_key") and i["conjunction_key"].get().strip().upper() == conjunction_key and i.get("selected", False)
            ]
            
            # Only show conjunction if there are multiple SELECTED items with the same key
            if len(selected_conjuncted_items) > 1:
                # Sort selected items by key for consistent numbering
                sorted_conjuncted_items = sorted(selected_conjuncted_items, key=lambda i: i.get('instance_info', {}).get('key', ''))
                current_instance_key = item.get('instance_info', {}).get('key', '')
                try:
                    index = next(i for i, sub_item in enumerate(sorted_conjuncted_items) if sub_item.get('instance_info', {}).get('key', '') == current_instance_key)
                    number = index + 1
                    conjunction_prefix = f"{conjunction_key}{number}: "
                    # Use count of SELECTED items for the message
                    conjunction_suffix = f"** {conjunction_key}1 to {conjunction_key}{len(selected_conjuncted_items)} must be approved together **"
                except StopIteration:
                    pass
        
        final_bid_text = f"{conjunction_prefix}{bid_text}\n{conjunction_suffix}".strip()
        
        # Update preview widget directly
        try:
            item["preview_text"].delete("1.0", tk.END)
            item["preview_text"].insert("1.0", final_bid_text)
        except Exception as e:
            print(f"Error updating preview widget: {e}")

    def update_live_preview(self, item, total_price):
        """Update live preview with change detection and caching for performance."""
        if item["preview_text"] is None:
            return
            
        # Check if widget still exists
        try:
            if not item["preview_text"].winfo_exists():
                return
        except:
            return
            
        # For custom items, never auto-update preview if user has edited it
        if item.get("is_custom", False) and item.get("user_edited", False):
            return
            
        # Check if user has manually edited the preview text
        try:
            current_text = item["preview_text"].get("1.0", tk.END).strip()
            if item.get('user_edited', False):
                # User has manually edited, don't overwrite
                return
        except:
            return
            
        # Get current values
        qty = item["qty"].get().strip() or "0"
        location = item["location"].get().strip() or "N/A"
        add_info = item["add_info"].get().strip()
        conjunction_key = item["conjunction_key"].get().strip().upper()
        
        # Note: Hash will be calculated after conjunction processing
        item_key = id(item)
        
        # Format the bid text
        template = item.get("template", "")
        if not template:
            # If template is missing, try to get it from item_data
            # This can happen if item was created before template was loaded
            original_name = item.get("original_name", "")
            category = None
            for cat, items in self.selected_items.items():
                if item in items.values():
                    category = cat
                    break
            if category and category in self.all_items:
                item_data = next((i for i in self.all_items[category] if i['item_name'] == original_name), None)
                if item_data and item_data.get('template'):
                    template = item_data['template']
                    item["template"] = template  # Save it for future use
                else:
                    return  # No template available, can't generate preview
            else:
                return  # Can't find category or item_data
        
        # Initialize calculated_total_price for hash calculation
        calculated_total_price = total_price  # Default to parameter value
        
        # Handle custom items - use only description from add_info (not item name, no price)
        if item.get("is_custom", False):
            description = add_info if add_info else ""
            # Use simple template for custom items (no price, no item name)
            try:
                bid_text = template.format(
                    description=description,
                    quantity=qty,
                    location=location,
                    info=add_info
                )
            except KeyError:
                # Fallback for custom items - just the description
                bid_text = description if description else ""
        else:
            # Extract unit_price for regular items
            unit_price = item.get("unit_price")
            if unit_price:
                unit_price = unit_price.get().strip().replace(",", "") if hasattr(unit_price, "get") else str(unit_price).strip().replace(",", "")
            else:
                unit_price = "0"
            
            # Handle zero or empty quantity
            try:
                qty_float = float(qty) if qty else 0.0
            except ValueError:
                qty_float = 0.0
            
            try:
                unit_price_float = float(unit_price) if unit_price else 0.0
            except ValueError:
                unit_price_float = 0.0
            
            # Recalculate total_price for consistency (use recalculated value in hash)
            calculated_total_price = qty_float * unit_price_float
            
            # Safe template formatting - try format first (faster path)
            try:
                bid_text = template.format(
                    quantity=qty or "0",
                    qty=qty or "0",  # Alias for quantity
                    lf=qty or "0",  # Common placeholder for linear feet
                    location=location or "N/A",
                    info=add_info or "",
                    add_info=add_info or "",  # Alias for info
                    total=calculated_total_price,
                    total_price=calculated_total_price,  # Alias for total
                    unit_price=unit_price or "0",
                    price=unit_price or "0",  # Alias for unit_price
                    cause=add_info or ""  # Use add_info as cause if needed
                )
            except KeyError:
                # Fallback formatting if template has unexpected placeholders
                # Replace common missing placeholders with defaults
                template_copy = template
                template_copy = template_copy.replace("{lf}", qty or "0")
                template_copy = template_copy.replace("{qty}", qty or "0")
                template_copy = template_copy.replace("{unit_price}", unit_price or "0")
                template_copy = template_copy.replace("{price}", unit_price or "0")
                template_copy = template_copy.replace("{add_info}", add_info or "")
                
                try:
                    # Try with just the basic placeholders
                    bid_text = template_copy.format(
                        quantity=qty or "0",
                        location=location or "N/A",
                        info=add_info or "",
                        total=calculated_total_price
                    )
                except:
                    # Ultimate fallback - use raw template with minimal replacement
                    bid_text = template_copy.replace("{quantity}", qty or "0") \
                                          .replace("{location}", location or "N/A") \
                                          .replace("{info}", add_info or "") \
                                          .replace("{total}", f"{calculated_total_price:.2f}")
        
        conjunction_prefix = ""
        conjunction_suffix = ""

        # Only show conjunction grouping for SELECTED items
        if conjunction_key and item.get("selected", False):
            # Find only SELECTED items with the same key (for numbering and message)
            selected_conjuncted_items = [
                i for cat in self.selected_items.values() 
                for i in cat.values() 
                if i.get("conjunction_key") and i["conjunction_key"].get().strip().upper() == conjunction_key and i.get("selected", False)
            ]
            
            if len(selected_conjuncted_items) > 1:
                sorted_conjuncted_items = sorted(selected_conjuncted_items, key=lambda i: i.get('instance_info', {}).get('key', ''))
                current_instance_key = item.get('instance_info', {}).get('key', '')
                
                try:
                    index = next(i for i, sub_item in enumerate(sorted_conjuncted_items) if sub_item.get('instance_info', {}).get('key', '') == current_instance_key)
                    number = index + 1
                    conjunction_prefix = f"{conjunction_key}{number}: "
                    # Use count of SELECTED items for the message
                    conjunction_suffix = f"** {conjunction_key}1 to {conjunction_key}{len(selected_conjuncted_items)} must be approved together **"
                except StopIteration:
                    conjunction_prefix = ""
                    conjunction_suffix = ""
                
        final_bid_text = f"{conjunction_prefix}{bid_text}\n{conjunction_suffix}".strip()
        
        # Include conjunction info in hash for change detection
        item_key = id(item)
        # Use calculated_total_price for regular items, total_price parameter for custom items
        hash_total_price = calculated_total_price if not item.get("is_custom", False) else total_price
        final_hash = hash((
            qty,
            location,
            add_info,
            conjunction_key,
            hash_total_price,
            item.get("is_custom", False),
            conjunction_prefix,
            conjunction_suffix
        ))
        
        # Check cache - skip update if nothing changed
        if item_key in self._preview_cache:
            cached_hash, cached_text = self._preview_cache[item_key]
            if cached_hash == final_hash:
                # Values haven't changed, skip update
                return
        
        # Update the preview text widget with error handling
        try:
            # Only update if text actually changed (avoid unnecessary widget updates)
            current_preview = item["preview_text"].get("1.0", tk.END).strip()
            if current_preview != final_bid_text:
                # Store current cursor position and selection
                current_cursor = item["preview_text"].index(tk.INSERT)
                current_selection = item["preview_text"].tag_ranges(tk.SEL)
                
                item["preview_text"].delete("1.0", tk.END)
                item["preview_text"].insert("1.0", final_bid_text)
                
                # Restore cursor position and selection
                try:
                    item["preview_text"].mark_set(tk.INSERT, current_cursor)
                    if current_selection:
                        item["preview_text"].tag_add(tk.SEL, current_selection[0], current_selection[1])
                except:
                    pass  # Ignore errors restoring cursor/selection
            
            # Update cache with new hash and text
            self._preview_cache[item_key] = (final_hash, final_bid_text)
        except:
            # Widget might be destroyed or invalid
            pass

    def toggle_item(self, category, item_key):
        item = self.selected_items[category][item_key]
        item["selected"] = not item["selected"]
        
        # For custom items, update checkbox variable
        if item.get("is_custom", False):
            if item.get("checkbox_var"):
                item["checkbox_var"].set(item["selected"])
            # The checkbox callback will handle the UI update
            return
        
        if item["button"]:
            item["button"].configure(bg=self.colors['selected'] if item["selected"] else self.colors['white'])
        
        self.update_all_previews()

    def update_all_previews(self):
        for category_items in self.selected_items.values():
            for item_info in category_items.values():
                if item_info["preview_text"]:
                    try:
                        # Check if widgets still exist before updating
                        if item_info["preview_text"].winfo_exists():
                            self.update_total_and_preview(item_info)
                    except:
                        # Widget might be destroyed, skip this item
                        continue
    
    def on_preview_text_change(self, item_info):
        """Handle text changes in the Live Preview and update generated bids if they exist."""
        # Mark this item as user-edited
        item_info['user_edited'] = True
        
        # Save preview text content to the item_info dict for persistence
        if item_info.get("preview_text"):
            try:
                if item_info["preview_text"].winfo_exists():
                    preview_content = item_info["preview_text"].get("1.0", tk.END).strip()
                    item_info["preview_text_content"] = preview_content
            except:
                pass
        
        # Update the generated bids section if it has content
        if hasattr(self, 'output_text') and self.output_text.get("1.0", tk.END).strip():
            self.update_generated_bids_from_preview(item_info)
    
    def update_generated_bids_from_preview(self, edited_item):
        """Update the generated bids section when Live Preview text is edited."""
        if not hasattr(self, 'output_text'):
            return
            
        # Get the current content of the output text
        current_content = self.output_text.get("1.0", tk.END)
        if not current_content.strip():
            return
        
        # Find the item in the generated bids and update it
        item_key = edited_item['instance_info']['key']
        item_name = edited_item['original_name']
        
        # Get the edited preview text
        edited_text = edited_item["preview_text"].get("1.0", tk.END).strip()
        
        # Update the generated bids section
        self.output_text.config(state=tk.NORMAL)
        
        # Find and replace the old bid text with the new one
        lines = current_content.split('\n')
        updated_lines = []
        
        for line in lines:
            if item_name in line and item_key in line:
                # This is the line we need to update
                # Extract the bid number and conjunction info if any
                bid_match = re.match(r'(\d+)\.\s*(.*)', line)
                if bid_match:
                    bid_number = bid_match.group(1)
                    # Create new bid text with the edited preview
                    new_bid_text = f"{bid_number}. {edited_text}"
                    updated_lines.append(new_bid_text)
                else:
                    updated_lines.append(line)
            else:
                updated_lines.append(line)
        
        # Update the output text
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert("1.0", '\n'.join(updated_lines))
        
        self.output_text.config(state=tk.DISABLED)
    
    def save_to_docs1(self):
        """Save to Docs1 format: Table with SL No., Bids, and Photos columns"""
        # This is the original format - table layout
        bid_count = 0
        final_bids = []
        bid_photos = []

        # Count selected items across all categories for verification
        total_selected = 0
        selected_by_category = {}
        for category, cat_items in self.selected_items.items():
            selected_count = sum(1 for item in cat_items.values() if item["selected"])
            if selected_count > 0:
                selected_by_category[category] = selected_count
                total_selected += selected_count

        conjunction_groups = {}
        standalone_bids = []
        for category, category_items in self.selected_items.items():
            for item in category_items.values():
                if item["selected"]:
                    conjunction_key = item["conjunction_key"].get().strip().upper()
                    if conjunction_key:
                        if conjunction_key not in conjunction_groups:
                            conjunction_groups[conjunction_key] = []
                        conjunction_groups[conjunction_key].append(item)
                    else:
                        standalone_bids.append(item)
        
        for key in conjunction_groups:
            conjunction_groups[key].sort(key=lambda x: x['instance_info']['key'])
        
        all_selected_items = []
        for key in sorted(conjunction_groups.keys()):
            all_selected_items.extend(conjunction_groups[key])
        all_selected_items.extend(sorted(standalone_bids, key=lambda x: x['instance_info']['key']))

        if not all_selected_items:
            messagebox.showwarning("No Bids Selected", "Please select some bids before saving to document!")
            return
        
        # Show summary of what will be saved if items from multiple categories
        if len(selected_by_category) > 1:
            category_summary = "\n".join([f"• {cat}: {count} item(s)" for cat, count in selected_by_category.items()])
            result = messagebox.askyesno("Multi-Category Document", 
                f"You are about to save bids from multiple categories:\n\n{category_summary}\n\nTotal: {total_selected} items\n\nContinue?")
            if not result:
                return

        for i, item in enumerate(all_selected_items, 1):
            bid_count += 1
            
            # Use the Live Preview text if available, otherwise use template
            bid_text, _ = self._get_item_bid_data(item)
            
            final_bid_text = f"{bid_count}. {bid_text}"

            final_bids.append(final_bid_text)

            instance_key = item['instance_info']['key']
            category_name = next(cat for cat, items in self.selected_items.items() if instance_key in items)
            photo_key = f"{category_name}_{instance_key}"

            if photo_key in self.item_photos and self.item_photos[photo_key]:
                bid_photos.append(self.item_photos[photo_key])
            else:
                bid_photos.append(None)

        try:
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
            use_docx = True
        except ImportError:
            use_docx = False
            messagebox.showinfo("Info", "python-docx not found. Saving as text file instead.\nTo save as Word document, install: pip install python-docx")

        if use_docx:
            doc = Document()
            
            # Get project information
            project_info = self._get_project_info()
            
            # Title
            wo_number = project_info['wo_number']
            if wo_number != "N/A":
                doc.add_heading(f"Preservation Universe Bid Proposal - WO# {wo_number}", 0)
            else:
                doc.add_heading('Preservation Universe Bid Proposal', 0)
            
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Project Information Section
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            info_table.columns[0].width = Inches(2.0)
            info_table.columns[1].width = Inches(5.0)
            
            # Row 1: Work Order
            info_table.rows[0].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[0].cells[0].paragraphs[0].add_run("Work Order (WO):")
            run1.bold = True
            info_table.rows[0].cells[1].text = project_info['wo_number']
            
            # Row 2: Property Address
            info_table.rows[1].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[1].cells[0].paragraphs[0].add_run("Property Address:")
            run1.bold = True
            info_table.rows[1].cells[1].text = project_info['property_address']
            
            # Row 3: Work Order Type
            info_table.rows[2].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[2].cells[0].paragraphs[0].add_run("Work Order Type:")
            run1.bold = True
            info_table.rows[2].cells[1].text = project_info['wo_type']
            
            # Row 4: Client Code
            info_table.rows[3].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[3].cells[0].paragraphs[0].add_run("Client Code:")
            run1.bold = True
            info_table.rows[3].cells[1].text = project_info['client_code']
            
            doc.add_paragraph()
            
            # Dates and Processor Name Section
            dates_table = doc.add_table(rows=3, cols=2)
            dates_table.style = 'Table Grid'
            dates_table.columns[0].width = Inches(2.0)
            dates_table.columns[1].width = Inches(5.0)
            
            # Row 1: Project Creation Date
            dates_table.rows[0].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[0].cells[0].paragraphs[0].add_run("Project Creation Date:")
            run1.bold = True
            dates_table.rows[0].cells[1].text = project_info['created_at']
            
            # Row 2: Project Last Modification Date
            dates_table.rows[1].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[1].cells[0].paragraphs[0].add_run("Project Last Modification Date:")
            run1.bold = True
            dates_table.rows[1].cells[1].text = project_info['updated_at']
            
            # Row 3: Processor Name
            dates_table.rows[2].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[2].cells[0].paragraphs[0].add_run("Processor Name:")
            run1.bold = True
            processor_name = self.username if hasattr(self, 'username') and self.username else "N/A"
            dates_table.rows[2].cells[1].text = processor_name
            
            doc.add_paragraph()
            
            # Summary Section
            total_items = len(all_selected_items)
            
            # Count occurrences of each item name
            item_count_dict = {}
            for item in all_selected_items:
                item_name = item.get('original_name', 'Unknown Item')
                item_count_dict[item_name] = item_count_dict.get(item_name, 0) + 1
            
            # Create list with counts
            selected_items_with_counts = []
            for item_name, count in sorted(item_count_dict.items()):
                if count > 1:
                    selected_items_with_counts.append(f"{item_name} ({count})")
                else:
                    selected_items_with_counts.append(item_name)
            
            summary_para1 = doc.add_paragraph()
            summary_para1.add_run("Summary: ").bold = True
            summary_para1.add_run(f"Total Bid Count: {total_items}")
            
            summary_para2 = doc.add_paragraph()
            summary_para2.add_run("Items Selected:").bold = True
            for item_with_count in selected_items_with_counts:
                summary_para2.add_run(f"\n• {item_with_count}")
            
            doc.add_paragraph()
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            
            table.columns[0].width = Inches(0.5)   # No.
            table.columns[1].width = Inches(4.0)   # Description
            table.columns[2].width = Inches(1.2)   # Total Price
            table.columns[3].width = Inches(2.0)   # Photo
            
            hdr_cells = table.rows[0].cells
            
            header_color_hex = self.colors['primary_blue'].lstrip('#')
            set_cell_background(hdr_cells[0], header_color_hex)
            set_cell_background(hdr_cells[1], header_color_hex)
            set_cell_background(hdr_cells[2], header_color_hex)
            set_cell_background(hdr_cells[3], header_color_hex)

            for cell, text in zip(hdr_cells, ['No.', 'Description', 'Total Price', 'Photo']):
                cell.text = text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            
            # Create list of items with prices for easier processing
            items_with_prices = []
            for i, (bid_text, photo_data) in enumerate(zip(final_bids, bid_photos)):
                # Get the corresponding item to calculate price
                if i < len(all_selected_items):
                    item = all_selected_items[i]
                    total_price = self._calculate_item_price(item)
                else:
                    total_price = 0.0
                
                # Extract bid text without price and remove number prefix
                # First remove the number prefix like "1. " or "2. "
                bid_text_clean = re.sub(r'^\d+\.\s*', '', bid_text)
                # Then remove price line if it exists
                bid_text_clean = self._extract_bid_text_without_price(bid_text_clean)
                if not bid_text_clean:
                    # If extraction failed, try to remove just the price line from original
                    bid_text_clean = re.sub(r'\nPrice.*?\$.*?\n?', '\n', bid_text)
                    bid_text_clean = re.sub(r'^\d+\.\s*', '', bid_text_clean)
                    bid_text_clean = bid_text_clean.strip()
                    if not bid_text_clean:
                        bid_text_clean = re.sub(r'^\d+\.\s*', '', bid_text)
                
                items_with_prices.append({
                    'number': i + 1,
                    'description': bid_text_clean.strip(),
                    'price': total_price,
                    'photo': photo_data
                })
            
            for item_data in items_with_prices:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item_data['number'])
                row_cells[1].text = item_data['description']
                row_cells[2].text = f"${item_data['price']:.2f}"
                
                if item_data['photo'] and item_data['photo'].get('original'):
                    try:
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                            temp_path = temp_file.name
                            item_data['photo']['original'].save(temp_path)
                        
                        row_cells[3].paragraphs[0].add_run().add_picture(temp_path, width=Inches(1.5))
                        
                        os.remove(temp_path)
                    except Exception as e:
                        print(f"Error adding image: {e}")
                        row_cells[3].text = "Error loading image"
            
            doc.add_paragraph()
            footer = doc.add_paragraph('Generated by Preservation Universe Bid Writer')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer.runs[0]
            footer_run.italic = True
            
            # Generate filename: Bids_{WO}_{date}
            wo_number_for_file = project_info['wo_number'].replace(" ", "_") if project_info['wo_number'] != "N/A" else "N/A"
            date_str = datetime.now().strftime("%d%b")  # Format: 25Dec
            default_filename = f"Bids_{wo_number_for_file}_{date_str}.docx"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save Bids Document (Docs1 Format)"
            )
            
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Success", f"Bids saved successfully to:\n{file_path}")
                
                if messagebox.askyesno("Open File", "Would you like to open the saved document?"):
                    try:
                        if os.name == 'nt':
                            os.startfile(file_path)
                        elif sys.platform == 'darwin':
                            os.system(f'open "{file_path}"')
                        else:
                            os.system(f'xdg-open "{file_path}"')
                    except Exception as e:
                        messagebox.showinfo("File Saved", f"Document saved successfully!\nLocation: {file_path}")
        
        else:
            # Fallback to text file for Docs1
            # Generate filename: Bids_{WO}_{date}
            project_info = self._get_project_info()
            wo_number_for_file = project_info['wo_number'].replace(" ", "_") if project_info['wo_number'] != "N/A" else "N/A"
            date_str = datetime.now().strftime("%d%b")  # Format: 25Dec
            default_filename = f"Bids_{wo_number_for_file}_{date_str}.txt"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text Document", "*.txt"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save Bids Document"
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write("TECHVENGERS BID PROPOSAL\n")
                    file.write("=" * 50 + "\n")
                    file.write(f"Date: {datetime.now().strftime('%B %d, %Y')}\n\n")
                    
                    for bid_text, photo_data in zip(final_bids, bid_photos):
                        file.write(f"{bid_text}\n")
                        if photo_data:
                            file.write("[Photo attached - see Word version for images]\n")
                        file.write("\n")
                        if bid_text != final_bids[-1]:
                            file.write("─" * 50 + "\n\n")
                    
                    file.write("\nGenerated by Preservation Universe Bid Writer\n")
                
                messagebox.showinfo("Success", f"Bids saved successfully to:\n{file_path}")
                
                if messagebox.askyesno("Open File", "Would you like to open the saved document?"):
                    try:
                        if os.name == 'nt':
                            os.startfile(file_path)
                        elif sys.platform == 'darwin':
                            os.system(f'open "{file_path}"')
                        else:
                            os.system(f'xdg-open "{file_path}"')
                    except Exception as e:
                        messagebox.showinfo("File Saved", f"Document saved successfully!\nLocation: {file_path}")
    
    def _calculate_item_price(self, item):
        """Calculate the total price for an item"""
        try:
            qty_str = item["qty"].get().strip().replace(",", "")
            unit_price_str = item["unit_price"].get().strip().replace(",", "")
            qty = float(qty_str) if qty_str else 0.0
            unit_price = float(unit_price_str) if unit_price_str else 0.0
            total_price = round(qty * unit_price, 2)
            return total_price
        except (ValueError, TypeError):
            return 0.0
    
    def _format_date(self, date_str):
        """Format date string to readable format"""
        if not date_str:
            return "N/A"
        try:
            # Handle ISO format: 2024-01-01T12:00:00+00:00 or 2024-01-01
            if 'T' in date_str:
                date_part = date_str.split('T')[0]
            else:
                date_part = date_str[:10]
            dt = datetime.strptime(date_part, '%Y-%m-%d')
            return dt.strftime('%B %d, %Y')
        except Exception:
            return date_str[:10] if len(date_str) >= 10 else date_str
    
    def _get_project_info(self):
        """Get project information including dates from database"""
        wo_number = self.wo_entry.get().strip() if hasattr(self, 'wo_entry') else ""
        property_address = self.address_entry.get().strip() if hasattr(self, 'address_entry') else self.property_address or ""
        client_code = self.client_code_entry.get().strip() if hasattr(self, 'client_code_entry') else ""
        wo_type = self.wo_type_entry.get().strip() if hasattr(self, 'wo_type_entry') else ""
        
        created_at = None
        updated_at = None
        
        # Try to get dates from database
        if self.db and self.user_id and wo_number:
            try:
                bid_data = self.db.load_bid(wo_number, self.user_id, all_users=True)
                if bid_data:
                    created_at = bid_data.get('created_at')
                    updated_at = bid_data.get('updated_at')
            except Exception as e:
                print(f"Error getting dates from database: {e}")
        
        return {
            'wo_number': wo_number or "N/A",
            'property_address': property_address or "N/A",
            'client_code': client_code or "N/A",
            'wo_type': wo_type or "N/A",
            'created_at': self._format_date(created_at),
            'updated_at': self._format_date(updated_at)
        }
    
    def _extract_bid_text_without_price(self, bid_text):
        """Extract bid text without the price line"""
        # Remove lines that contain "Price:" or "Price $"
        lines = bid_text.split('\n')
        cleaned_lines = []
        for line in lines:
            if not (line.strip().startswith('Price:') or line.strip().startswith('Price $')):
                cleaned_lines.append(line)
        return '\n'.join(cleaned_lines).strip()
    
    def save_to_docs2(self):
        """Save to Docs2 format: Bid text -> Price below -> Photo below (vertical layout)"""
        bid_count = 0
        bid_items = []  # List of dicts with bid_text, price, photo_data

        # Count selected items across all categories for verification
        total_selected = 0
        selected_by_category = {}
        for category, cat_items in self.selected_items.items():
            selected_count = sum(1 for item in cat_items.values() if item["selected"])
            if selected_count > 0:
                selected_by_category[category] = selected_count
                total_selected += selected_count

        conjunction_groups = {}
        standalone_bids = []
        for category, category_items in self.selected_items.items():
            for item in category_items.values():
                if item["selected"]:
                    conjunction_key = item["conjunction_key"].get().strip().upper()
                    if conjunction_key:
                        if conjunction_key not in conjunction_groups:
                            conjunction_groups[conjunction_key] = []
                        conjunction_groups[conjunction_key].append(item)
                    else:
                        standalone_bids.append(item)
        
        for key in conjunction_groups:
            conjunction_groups[key].sort(key=lambda x: x['instance_info']['key'])
        
        all_selected_items = []
        for key in sorted(conjunction_groups.keys()):
            all_selected_items.extend(conjunction_groups[key])
        all_selected_items.extend(sorted(standalone_bids, key=lambda x: x['instance_info']['key']))

        if not all_selected_items:
            messagebox.showwarning("No Bids Selected", "Please select some bids before saving to document!")
            return
        
        # Show summary of what will be saved if items from multiple categories
        if len(selected_by_category) > 1:
            category_summary = "\n".join([f"• {cat}: {count} item(s)" for cat, count in selected_by_category.items()])
            result = messagebox.askyesno("Multi-Category Document", 
                f"You are about to save bids from multiple categories:\n\n{category_summary}\n\nTotal: {total_selected} items\n\nContinue?")
            if not result:
                return

        for i, item in enumerate(all_selected_items, 1):
            bid_count += 1
            
            # Get bid text and remove price line if it exists
            bid_text, _ = self._get_item_bid_data(item)
            bid_text_without_price = self._extract_bid_text_without_price(bid_text)
            if not bid_text_without_price:
                bid_text_without_price = bid_text  # Fallback to original if extraction fails
            
            # Calculate price
            total_price = self._calculate_item_price(item)
            
            # Get photo
            instance_key = item['instance_info']['key']
            category_name = next(cat for cat, items in self.selected_items.items() if instance_key in items)
            photo_key = f"{category_name}_{instance_key}"
            photo_data = None
            if photo_key in self.item_photos and self.item_photos[photo_key]:
                photo_data = self.item_photos[photo_key]
            
            bid_items.append({
                'number': bid_count,
                'text': bid_text_without_price,
                'price': total_price,
                'photo': photo_data
            })

        try:
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
            use_docx = True
        except ImportError:
            use_docx = False
            messagebox.showinfo("Info", "python-docx not found. Saving as text file instead.\nTo save as Word document, install: pip install python-docx")

        if use_docx:
            doc = Document()
            
            # Get project information
            project_info = self._get_project_info()
            
            # Title
            wo_number = project_info['wo_number']
            if wo_number != "N/A":
                doc.add_heading(f"Preservation Universe Bid Proposal - WO# {wo_number}", 0)
            else:
                doc.add_heading('Preservation Universe Bid Proposal', 0)
            
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Project Information Section
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            info_table.columns[0].width = Inches(2.0)
            info_table.columns[1].width = Inches(5.0)
            
            # Row 1: Work Order
            info_table.rows[0].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[0].cells[0].paragraphs[0].add_run("Work Order (WO):")
            run1.bold = True
            info_table.rows[0].cells[1].text = project_info['wo_number']
            
            # Row 2: Property Address
            info_table.rows[1].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[1].cells[0].paragraphs[0].add_run("Property Address:")
            run1.bold = True
            info_table.rows[1].cells[1].text = project_info['property_address']
            
            # Row 3: Work Order Type
            info_table.rows[2].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[2].cells[0].paragraphs[0].add_run("Work Order Type:")
            run1.bold = True
            info_table.rows[2].cells[1].text = project_info['wo_type']
            
            # Row 4: Client Code
            info_table.rows[3].cells[0].paragraphs[0].clear()
            run1 = info_table.rows[3].cells[0].paragraphs[0].add_run("Client Code:")
            run1.bold = True
            info_table.rows[3].cells[1].text = project_info['client_code']
            
            doc.add_paragraph()
            
            # Dates and Processor Name Section
            dates_table = doc.add_table(rows=3, cols=2)
            dates_table.style = 'Table Grid'
            dates_table.columns[0].width = Inches(2.0)
            dates_table.columns[1].width = Inches(5.0)
            
            # Row 1: Project Creation Date
            dates_table.rows[0].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[0].cells[0].paragraphs[0].add_run("Project Creation Date:")
            run1.bold = True
            dates_table.rows[0].cells[1].text = project_info['created_at']
            
            # Row 2: Project Last Modification Date
            dates_table.rows[1].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[1].cells[0].paragraphs[0].add_run("Project Last Modification Date:")
            run1.bold = True
            dates_table.rows[1].cells[1].text = project_info['updated_at']
            
            # Row 3: Processor Name
            dates_table.rows[2].cells[0].paragraphs[0].clear()
            run1 = dates_table.rows[2].cells[0].paragraphs[0].add_run("Processor Name:")
            run1.bold = True
            processor_name = self.username if hasattr(self, 'username') and self.username else "N/A"
            dates_table.rows[2].cells[1].text = processor_name
            
            doc.add_paragraph()
            
            # Summary Section
            total_items = len(bid_items)
            
            # Count occurrences of each item name
            item_count_dict = {}
            for item in all_selected_items:
                item_name = item.get('original_name', 'Unknown Item')
                item_count_dict[item_name] = item_count_dict.get(item_name, 0) + 1
            
            # Create list with counts
            selected_items_with_counts = []
            for item_name, count in sorted(item_count_dict.items()):
                if count > 1:
                    selected_items_with_counts.append(f"{item_name} ({count})")
                else:
                    selected_items_with_counts.append(item_name)
            
            summary_para1 = doc.add_paragraph()
            summary_para1.add_run("Summary: ").bold = True
            summary_para1.add_run(f"Total Bid Count: {total_items}")
            
            summary_para2 = doc.add_paragraph()
            summary_para2.add_run("Items Selected:").bold = True
            for item_with_count in selected_items_with_counts:
                summary_para2.add_run(f"\n• {item_with_count}")
            
            doc.add_paragraph()
            
            # Add each bid item with format: Bid text -> Price -> Photo
            for bid_item in bid_items:
                # Add bid number and text
                bid_paragraph = doc.add_paragraph()
                bid_paragraph.add_run(f"{bid_item['number']}. {bid_item['text']}").bold = False
                
                # Add price below bid text
                price_paragraph = doc.add_paragraph()
                price_run = price_paragraph.add_run(f"Price ${bid_item['price']:.2f}")
                price_run.bold = True
                
                # Add photo below price if available
                if bid_item['photo'] and bid_item['photo'].get('original'):
                    try:
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                            temp_path = temp_file.name
                            bid_item['photo']['original'].save(temp_path)
                        
                        photo_paragraph = doc.add_paragraph()
                        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        photo_paragraph.add_run().add_picture(temp_path, width=Inches(4.0))
                        
                        os.remove(temp_path)
                    except Exception as e:
                        print(f"Error adding image: {e}")
                        error_para = doc.add_paragraph()
                        error_para.add_run("Error loading image").italic = True
                
                # Add spacing between items
                doc.add_paragraph()
            
            doc.add_paragraph()
            footer = doc.add_paragraph('Generated by Preservation Universe Bid Writer')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer.runs[0]
            footer_run.italic = True
            
            # Generate filename: Bids_{WO}_{date}
            wo_number_for_file = project_info['wo_number'].replace(" ", "_") if project_info['wo_number'] != "N/A" else "N/A"
            date_str = datetime.now().strftime("%d%b")  # Format: 25Dec
            default_filename = f"Bids_{wo_number_for_file}_{date_str}.docx"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save Bids Document (Docs2 Format)"
            )
            
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Success", f"Bids saved successfully to:\n{file_path}")
                
                if messagebox.askyesno("Open File", "Would you like to open the saved document?"):
                    try:
                        if os.name == 'nt':
                            os.startfile(file_path)
                        elif sys.platform == 'darwin':
                            os.system(f'open "{file_path}"')
                        else:
                            os.system(f'xdg-open "{file_path}"')
                    except Exception as e:
                        messagebox.showinfo("File Saved", f"Document saved successfully!\nLocation: {file_path}")
        
        else:
            # Fallback to text file
            # Generate filename: Bids_{WO}_{date}
            wo_number_for_file = project_info['wo_number'].replace(" ", "_") if project_info['wo_number'] != "N/A" else "N/A"
            date_str = datetime.now().strftime("%d%b")  # Format: 25Dec
            default_filename = f"Bids_{wo_number_for_file}_{date_str}.txt"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text Document", "*.txt"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save Bids Document (Docs2 Format)"
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write("TECHVENGERS BID PROPOSAL\n")
                    file.write("=" * 50 + "\n")
                    file.write(f"Date: {datetime.now().strftime('%B %d, %Y')}\n\n")
                    
                    for bid_item in bid_items:
                        file.write(f"{bid_item['number']}. {bid_item['text']}\n")
                        file.write(f"Price ${bid_item['price']:.2f}\n")
                        if bid_item['photo']:
                            file.write("[Photo attached - see Word version for images]\n")
                        file.write("\n")
                        if bid_item != bid_items[-1]:
                            file.write("─" * 50 + "\n\n")
                    
                    file.write("\nGenerated by Preservation Universe Bid Writer\n")
                
                messagebox.showinfo("Success", f"Bids saved successfully to:\n{file_path}")
                
                if messagebox.askyesno("Open File", "Would you like to open the saved document?"):
                    try:
                        if os.name == 'nt':
                            os.startfile(file_path)
                        elif sys.platform == 'darwin':
                            os.system(f'open "{file_path}"')
                        else:
                            os.system(f'xdg-open "{file_path}"')
                    except Exception as e:
                        messagebox.showinfo("File Saved", f"Document saved successfully!\nLocation: {file_path}")
    
    def generate_bids(self):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete("1.0", tk.END)
        self.output_text.images = [] 
        
        # Check selections across all categories
        total_selected = sum(
            1 for cat_items in self.selected_items.values()
            for item in cat_items.values() if item["selected"]
        )
        if total_selected == 0:
            messagebox.showinfo("No Selections", "No bids selected!")
            self.bid_count_label.config(text="Total Bids: 0")
            self.output_text.config(state=tk.DISABLED)
            return
        
        conjunction_groups = {}
        standalone_bids = []
        
        # Process all categories to collect selected items
        for cat_items in self.selected_items.values():
            for item in cat_items.values():
                if item["selected"]:
                    key = item["conjunction_key"].get().strip().upper()
                    if key:
                        if key not in conjunction_groups:
                            conjunction_groups[key] = []
                        conjunction_groups[key].append(item)
                    else:
                        standalone_bids.append(item)
        
        bid_number = 1
        
        for key in sorted(conjunction_groups.keys()):
            items = conjunction_groups[key]
            sorted_items = sorted(items, key=lambda x: x['instance_info']['key'])
            for item in sorted_items:
                bid_text, photo_key = self._get_item_bid_data(item)
                numbered_bid = f"{bid_number}. {bid_text}"
                self.output_text.insert(tk.END, f"{numbered_bid}\n")
                # Append price line under the bid text
                try:
                    qty_val = float((item["qty"].get() if hasattr(item["qty"], "get") else str(item["qty"])) or 0)
                except Exception:
                    qty_val = 0.0
                try:
                    unit_val = float((item["unit_price"].get() if hasattr(item["unit_price"], "get") else str(item["unit_price"])) or 0)
                except Exception:
                    unit_val = 0.0
                total_val = round(qty_val * unit_val, 2)
                self.output_text.insert(tk.END, f"Price: ${total_val:.2f}\n")
                
                self._insert_photo(photo_key)
                bid_number += 1
        
        for item in sorted(standalone_bids, key=lambda x: x['instance_info']['key']):
            bid_text, photo_key = self._get_item_bid_data(item)
            numbered_bid = f"{bid_number}. {bid_text}"
            self.output_text.insert(tk.END, f"{numbered_bid}\n")
            # Append price line under the bid text
            try:
                qty_val = float((item["qty"].get() if hasattr(item["qty"], "get") else str(item["qty"])) or 0)
            except Exception:
                qty_val = 0.0
            try:
                unit_val = float((item["unit_price"].get() if hasattr(item["unit_price"], "get") else str(item["unit_price"])) or 0)
            except Exception:
                unit_val = 0.0
            total_val = round(qty_val * unit_val, 2)
            self.output_text.insert(tk.END, f"Price: ${total_val:.2f}\n")
            
            self._insert_photo(photo_key)
            bid_number += 1

        if bid_number == 1:
            messagebox.showinfo("No Selections", "No bids selected!")
            
        self.bid_count_label.config(text=f"Total Bids: {bid_number - 1}")
        self.output_text.config(state=tk.DISABLED)

    def open_full_page_view(self):
        """Open all bids in a new full-page window."""
        # Check if there are any selected bids
        total_selected = sum(
            1 for cat_items in self.selected_items.values()
            for item in cat_items.values() if item["selected"]
        )
        if total_selected == 0:
            messagebox.showinfo("No Bids", "No bids selected! Please select bids and generate them first.")
            return
        
        # Create new window
        full_page_window = tk.Toplevel(self.root)
        full_page_window.title("Full Page Bid View - Preservation Universe Bid Writer")
        
        # Make it full screen or large window
        screen_width = full_page_window.winfo_screenwidth()
        screen_height = full_page_window.winfo_screenheight()
        full_page_window.geometry(f"{int(screen_width * 0.9)}x{int(screen_height * 0.9)}")
        full_page_window.configure(bg=self.colors['white'])
        
        # Create header frame
        header_frame = tk.Frame(full_page_window, bg=self.colors['primary_blue'], height=60)
        header_frame.pack(fill='x')
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(header_frame, text="Full Page Bid View", 
                              font=("Arial", 16, "bold"), fg='white', 
                              bg=self.colors['primary_blue'])
        title_label.pack(side="left", padx=20, pady=15)
        
        close_button = tk.Button(header_frame, text="✕ Close", 
                                font=("Arial", 12, "bold"), bg="#dc3545",
                                fg="white", relief="flat", cursor="hand2",
                                activebackground="#c82333",
                                command=full_page_window.destroy)
        close_button.pack(side="right", padx=20, pady=15)
        
        # Create main content frame with scrollbar
        content_frame = tk.Frame(full_page_window, bg=self.colors['white'])
        content_frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        scrollbar = tk.Scrollbar(content_frame)
        scrollbar.pack(side="right", fill="y")
        
        full_page_text = tk.Text(content_frame, font=("Arial", 12), 
                                bg=self.colors['white'], fg=self.colors['text_primary'],
                                wrap=tk.WORD, relief="flat", padx=20, pady=20,
                                yscrollcommand=scrollbar.set)
        full_page_text.pack(fill='both', expand=True)
        scrollbar.config(command=full_page_text.yview)
        
        # Generate bids in the same format as generate_bids method
        full_page_text.config(state=tk.NORMAL)
        full_page_text.delete("1.0", tk.END)
        full_page_text.images = []
        
        conjunction_groups = {}
        standalone_bids = []
        
        # Process all categories to collect selected items
        for cat_items in self.selected_items.values():
            for item in cat_items.values():
                if item["selected"]:
                    key = item["conjunction_key"].get().strip().upper()
                    if key:
                        if key not in conjunction_groups:
                            conjunction_groups[key] = []
                        conjunction_groups[key].append(item)
                    else:
                        standalone_bids.append(item)
        
        bid_number = 1
        
        # Process conjunction groups
        for key in sorted(conjunction_groups.keys()):
            items = conjunction_groups[key]
            sorted_items = sorted(items, key=lambda x: x['instance_info']['key'])
            for item in sorted_items:
                bid_text, photo_key = self._get_item_bid_data(item)
                numbered_bid = f"{bid_number}. {bid_text}"
                full_page_text.insert(tk.END, f"{numbered_bid}\n")
                # Append price line under the bid text
                try:
                    qty_val = float((item["qty"].get() if hasattr(item["qty"], "get") else str(item["qty"])) or 0)
                except Exception:
                    qty_val = 0.0
                try:
                    unit_val = float((item["unit_price"].get() if hasattr(item["unit_price"], "get") else str(item["unit_price"])) or 0)
                except Exception:
                    unit_val = 0.0
                total_val = round(qty_val * unit_val, 2)
                full_page_text.insert(tk.END, f"Price: ${total_val:.2f}\n")
                
                # Insert photo if available
                self._insert_photo_to_text(full_page_text, photo_key)
                
                # Add blank line after each bid for visual separation
                full_page_text.insert(tk.END, "\n")
                bid_number += 1
        
        # Process standalone bids
        for item in sorted(standalone_bids, key=lambda x: x['instance_info']['key']):
            bid_text, photo_key = self._get_item_bid_data(item)
            numbered_bid = f"{bid_number}. {bid_text}"
            full_page_text.insert(tk.END, f"{numbered_bid}\n")
            # Append price line under the bid text
            try:
                qty_val = float((item["qty"].get() if hasattr(item["qty"], "get") else str(item["qty"])) or 0)
            except Exception:
                qty_val = 0.0
            try:
                unit_val = float((item["unit_price"].get() if hasattr(item["unit_price"], "get") else str(item["unit_price"])) or 0)
            except Exception:
                unit_val = 0.0
            total_val = round(qty_val * unit_val, 2)
            full_page_text.insert(tk.END, f"Price: ${total_val:.2f}\n")
            
            # Insert photo if available
            self._insert_photo_to_text(full_page_text, photo_key)
            
            # Add blank line after each bid for visual separation
            full_page_text.insert(tk.END, "\n")
            bid_number += 1
        
        full_page_text.config(state=tk.DISABLED)
        
        # Bind mouse wheel scrolling
        def _on_mousewheel(event):
            try:
                if hasattr(event, 'delta') and event.delta:
                    delta = int(-1 * (event.delta / 120))
                else:
                    delta = -1 if event.num == 4 else 1
                full_page_text.yview_scroll(delta, "units")
            except (AttributeError, TypeError):
                pass
        
        full_page_text.bind("<MouseWheel>", _on_mousewheel)
        full_page_text.bind("<Button-4>", lambda e: full_page_text.yview_scroll(-1, "units"))
        full_page_text.bind("<Button-5>", lambda e: full_page_text.yview_scroll(1, "units"))
        
        # Focus on the new window
        full_page_window.focus_set()
    
    def _insert_photo_to_text(self, text_widget, photo_key):
        """Insert a photo into a text widget (helper for full page view)."""
        if photo_key in self.item_photos and self.item_photos[photo_key]:
            try:
                text_widget.insert(tk.END, "\n")
                image_data = self.item_photos[photo_key]['original'].copy()
                max_width, max_height = 400, 300
                image_data.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                photo_for_output = ImageTk.PhotoImage(image_data)
                
                text_widget.image_create(tk.END, image=photo_for_output)
                
                # Keep reference to prevent garbage collection
                if not hasattr(text_widget, 'images'):
                    text_widget.images = []
                text_widget.images.append(photo_for_output)
                
                text_widget.insert(tk.END, "\n")
            except Exception as e:
                print(f"Error inserting image in full page view: {e}")

    def _get_item_bid_data(self, item):
        """Return the text to use for an item's bid and the associated photo key.
        
        Prefers the current Live Preview text (including any user edits). Falls back
        to generating text from the item's template if the preview is empty or missing.
        """
        # Find the category this item belongs to
        category = None
        for cat, items in self.selected_items.items():
            if item in items.values():
                category = cat
                break
        
        if category is None:
            # Fallback: try to find by instance key
            instance_key = item['instance_info']['key']
            for cat, items in self.selected_items.items():
                if instance_key in items:
                    category = cat
                    break
        
        instance_key = item['instance_info']['key']
        photo_key = f"{category}_{instance_key}" if category else f"unknown_{instance_key}"

        # Prefer whatever is visible in the Live Preview cell (only if widget still exists)
        preview_widget = item.get("preview_text")
        if preview_widget is not None:
            try:
                if preview_widget.winfo_exists():
                    preview_value = preview_widget.get("1.0", tk.END).strip()
                    if preview_value:
                        return preview_value, photo_key
            except Exception:
                # Widget may have been destroyed when switching categories; fall back to template
                pass

        # Fallback to constructing from the template
        qty = item["qty"].get().strip().replace(",", "")
        unit_price = item["unit_price"].get().strip().replace(",", "")
        location = item["location"].get().strip()
        add_info = item["add_info"].get().strip()

        # Handle zero or empty quantity - allow it, calculate price as 0
        try:
            qty_float = float(qty) if qty else 0.0
        except ValueError:
            qty_float = 0.0

        try:
            unit_price_float = float(unit_price) if unit_price else 0.0
        except ValueError:
            unit_price_float = 0.0

        total_price = qty_float * unit_price_float

        # Safe template formatting - handle any missing placeholders
        try:
            bid_text = item["template"].format(
                quantity=qty or "0",
                qty=qty or "0",  # Alias for quantity
                lf=qty or "0",  # Common placeholder for linear feet
                location=location or "N/A",
                info=add_info or "",
                add_info=add_info or "",  # Alias for info
                total=total_price,
                total_price=total_price,  # Alias for total
                unit_price=unit_price or "0",
                price=unit_price or "0"  # Alias for unit_price
            )
        except KeyError as e:
            # If template has unexpected placeholders, use a safer approach
            # Replace common missing placeholders with defaults
            template = item["template"]
            template = template.replace("{lf}", qty or "0")
            template = template.replace("{qty}", qty or "0")
            template = template.replace("{unit_price}", unit_price or "0")
            template = template.replace("{price}", unit_price or "0")
            template = template.replace("{add_info}", add_info or "")
            
            try:
                bid_text = template.format(
                    quantity=qty or "0",
                    location=location or "N/A",
                    info=add_info or "",
                    total=total_price
                )
            except KeyError:
                # Last resort: just use the template as-is with minimal formatting
                bid_text = template.replace("{quantity}", qty or "0") \
                                  .replace("{location}", location or "N/A") \
                                  .replace("{info}", add_info or "") \
                                  .replace("{total}", f"{total_price:.2f}")

        return bid_text, photo_key

    def _insert_photo(self, photo_key):
        if photo_key in self.item_photos and self.item_photos[photo_key]:
            try:
                self.output_text.insert(tk.END, "\n")
                image_data = self.item_photos[photo_key]['original'].copy()
                max_width, max_height = 400, 300
                image_data.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                photo_for_output = ImageTk.PhotoImage(image_data)
                
                self.output_text.image_create(tk.END, image=photo_for_output)
                
                self.output_text.images.append(photo_for_output)
                
                self.output_text.insert(tk.END, "\n\n")
            except Exception as e:
                print(f"Error inserting image: {e}")

    def clear_bids(self):
        """Clears the generated bids section and resets all inputs."""
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete("1.0", tk.END)
        self.output_text.config(state=tk.DISABLED)
        self.bid_count_label.config(text="Total Bids: 0")
        
        self.output_text.images = []
        
        for category_items in self.selected_items.values():
            for item_info in category_items.values():
                item_info["selected"] = False
                if item_info["button"]:
                    item_info["button"].configure(bg=self.colors['white'])
                item_info["qty"].set("0")
                
                original_name = item_info['original_name']
                category_name = next(cat for cat, items in self.selected_items.items() 
                                     for key, info in items.items() if info == item_info)
                
                initial_price = self.get_initial_price(category_name, original_name)
                item_info["unit_price"].set(initial_price)
                
                item_info["location"].set("")
                item_info["add_info"].set("")
                item_info["conjunction_key"].set("")
                item_info["user_edited"] = False
                self.update_total_and_preview(item_info)
    
    def get_initial_price(self, category_name, item_name):
        """Helper to find the initial price from the all_items data structure."""
        if category_name in self.all_items and isinstance(self.all_items[category_name], list):
            item_data = next((item_dict for item_dict in self.all_items[category_name] if item_dict.get('item_name') == item_name), None)
            if item_data and 'unit_price' in item_data:
                return item_data['unit_price']
        
        return "0.00"